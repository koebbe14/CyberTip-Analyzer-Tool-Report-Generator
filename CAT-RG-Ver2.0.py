import json
import requests
import os
import sys
import threading
import webbrowser
import openpyxl
from openpyxl.styles import Font, Alignment
import queue
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ipaddress
from PIL import Image, ImageTk

class CyberTipAnalyzer:
    def load_recent_files(self):
        """Load recent files from JSON file."""
        try:
            if os.path.exists(self.recent_files_file):
                with open(self.recent_files_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return []
        except Exception as e:
            print(f"Error loading recent files: {e}")
            return []

    def load_investigator_info(self):
        """Load investigator info from JSON file or prompt if it doesn't exist."""
        try:
            if os.path.exists(self.investigator_file):
                with open(self.investigator_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.investigator_name = data.get('name', '')
                    self.investigator_title = data.get('title', '')
                if not self.investigator_name or not self.investigator_title:
                    self.get_investigator_info()
            else:
                self.get_investigator_info()
        except Exception as e:
            print(f"Error loading investigator info: {e}")
            self.get_investigator_info()

    def load_credentials(self):
        """Load MaxMind credentials from file if it exists, otherwise prompt the user."""
        try:
            if os.path.exists(self.credentials_file):
                with open(self.credentials_file, 'r', encoding='utf-8') as f:
                    creds = json.load(f)
                    account_id = creds.get('account_id', '')
                    license_key = creds.get('license_key', '')
                    self.maxmind_id_entry.delete(0, tk.END)
                    self.maxmind_id_entry.insert(0, account_id)
                    self.maxmind_key_entry.delete(0, tk.END)
                    self.maxmind_key_entry.insert(0, license_key)
            else:
                dialog = self.show_credentials_dialog()
                if dialog:
                    self.root.wait_window(dialog)
            self.load_arin_credentials()
        except Exception as e:
            messagebox.showwarning("Warning", f"Could not load credentials: {e}")
            dialog = self.show_credentials_dialog()
            if dialog:
                self.root.wait_window(dialog)
            self.load_arin_credentials()

    def load_arin_credentials(self):
        """Load ARIN API key from file if it exists, otherwise prompt the user."""
        try:
            if os.path.exists(self.arin_credentials_file):
                with open(self.arin_credentials_file, 'r', encoding='utf-8') as f:
                    creds = json.load(f)
                    self.arin_api_key = creds.get('api_key', '')
            else:
                dialog = self.show_arin_credentials_dialog()
                if dialog:
                    self.root.wait_window(dialog)
        except Exception as e:
            messagebox.showwarning("Warning", f"Could not load ARIN credentials: {e}")
            dialog = self.show_arin_credentials_dialog()
            if dialog:
                self.root.wait_window(dialog)

    def show_arin_credentials_dialog(self):
        """Display a dialog to enter ARIN API key (optional)."""
        dialog = tk.Toplevel(self.root)
        dialog.title("ARIN API Key")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()

        tk.Label(
            dialog,
            text="Enter your ARIN API key (optional).\n\n"
                 "Adding an ARIN API key increases daily query limits.\n\n"
                 "Daily query limits WITHOUT ARIN API key - 15 per min and 256 per day.\n"
                 "Daily query limits WITH ARIN API Key - 60 per min and 1,000 per day.\n\n"
                 "Get an API key at:"
        ).pack(pady=5)
        arin_link_label = tk.Label(dialog, text="this link", fg="blue", cursor="hand2")
        arin_link_label.pack(pady=5)
        arin_url = "https://account.arin.net/public/account-setup"
        arin_link_label.bind("<Button-1>", lambda e: webbrowser.open_new(arin_url))

        tk.Label(dialog, text="ARIN API Key:").pack(pady=5)
        arin_key_entry = tk.Entry(dialog, width=50)
        arin_key_entry.pack(pady=5)
        arin_key_entry.insert(0, self.arin_api_key)

        def save_and_close():
            self.arin_api_key = arin_key_entry.get().strip()
            if self.arin_api_key:
                self.save_arin_credentials()
                messagebox.showinfo("Success", "ARIN API key saved.")
            dialog.destroy()

        tk.Button(dialog, text="Save", command=save_and_close).pack(pady=10)
        tk.Button(dialog, text="Skip", command=dialog.destroy).pack(pady=5)
        return dialog

    def save_arin_credentials(self):
        """Save ARIN API key to file."""
        creds = {'api_key': self.arin_api_key}
        try:
            with open(self.arin_credentials_file, 'w', encoding='utf-8') as f:
                json.dump(creds, f)
        except Exception as e:
            messagebox.showerror("Error", f"Could not save ARIN credentials: {e}")

    def __init__(self, root):
        self.root = root
        self.root.title("CAT-RG")
        self.root.geometry("800x600")
        self.json_file_path = None
        self.ip_data: Dict = {}
        self.credentials_file = "maxmind_credentials.json"
        self.recent_files_file = "recent_files.json"
        self.investigator_file = "investigator_info.json"
        self.statements_file = "custom_statements.json"
        self.arin_credentials_file = "arin_credentials.json"
        self.arin_api_key = ""
        self.investigator_name = ""
        self.investigator_title = ""
        self.recent_files = self.load_recent_files()
        self.maxmind_id_entry = tk.Entry(self.root)
        self.maxmind_key_entry = tk.Entry(self.root)
        self.selected_statements = set()  # Track selected statements
        self.load_credentials()
        self.load_investigator_info()
        self.load_statements()
        
        if getattr(sys, 'frozen', False):
            base_path = Path(sys.executable).parent
        else:
            base_path = Path(__file__).parent

        logo_path = base_path / "logo.jpg"

        try:
            image = Image.open(logo_path)
            image = image.resize((200, 100), Image.Resampling.LANCZOS)
            self.logo_image = ImageTk.PhotoImage(image)
        except Exception as e:
            print(f"Error loading logo: {e}")
            self.logo_image = None
        
        self.create_menu()
        self.create_widgets()
        self.all_ip_data = {}
        # Initialize selected_statements with all statements
        self.selected_statements = set(self.statements.keys())

    def load_statements(self):
        """Load custom statements from JSON or use defaults."""
        self.default_statements = {
            "meta": """
When Meta responds "Yes" it means the contents of the file were viewed by an employee or contractor at Meta concurrently with or immediately before the file was submitted to NCMEC. When Meta responds "No" it means that while the contents of the file were not reviewed concurrently with or immediately before the file was submitted to NCMEC, historically at least one employee or contractor at Meta viewed a file whose hash matched the hash of the reported content and determined it contained apparent child pornography.

For video files, when Meta responds “Yes” it means the entire contents of the file were viewed by an employee or contractor at Meta concurrently with or immediately before the file was submitted to NCMEC. When Meta responds “No” it means that while the contents of the file were not reviewed concurrently with or immediately before the file was submitted to NCMEC, historically at least one employee or contractor at Meta viewed a file and determined it contained apparent child pornography, and that file's hash matched a violating portion or the entirety of the reported content.
""",
            "ip_intro": """
These following IP addresses were reported in the Cybertip. Each IP address was queried through the American Registry for Internet Numbers (ARIN) and Maxmind.com. ARIN is responsible for managing and distributing Internet number resources, like IP addresses, in North America. It is one of the five Regional Internet Registries (RIRs) worldwide, working under the global Internet Assigned Numbers Authority (IANA). ARIN maintains a public database (WHOIS) which tracks who holds what IP address. Maxmind is a company that provides a webservice IP geolocation tool. It should be noted that the estimated geographical location obtained from Maxmind’s geolocation database are not always accurate, particularly when resolving IP addresses utilized by cellular providers, as a mobile user’s location is constantly changing. The exact location of where an IP addresses geographically resolves to, along with the subscriber details, can only be obtained through legal process served to the provider.
""",
            "bingimage": """
"BingImage" (referred to as Visual Search) is a service of Microsoft's Bing search engine that provides similar images to an image provided by the user. This image can be provided either via upload or as a URL. The date/time provided indicates the time at which the image was received and evaluated by the BingImage service.
""",
            "xcorp": """
\n\n X retains different types of information for different time periods. Given X's real-time nature, some information may only be stored for a very brief period of time. \n\nFor accounts reported to NCMEC, X provides a copy of the preserved files, within the CyberTip report in the form of a .zip file which may be uploaded in multiple parts. \n\nAll times reported by X are in UTC. \n\nThe incident date/time is the timestamp from the most recent reported Post; however, if a Post is not reported, then the incident date/time will represent the account creation timestamp. \n\nX logs IPs in connection with user authentications to X (i.e., sessions, which may span multiple days) rather than individual Post postings; as a result, X is unable to provide insight into which IP address a specific Post was posted from. While X does not capture IPs for individual Posts, X provided a log of IPs for the timeframe relevant to the report. \n\nAll IP addresses in this report are associated with log-ins to the specific user account identified in the report.\n
"""
        }
        try:
            if os.path.exists(self.statements_file):
                with open(self.statements_file, 'r', encoding='utf-8') as f:
                    self.statements = json.load(f)
            else:
                self.statements = self.default_statements.copy()
        except Exception as e:
            print(f"Error loading statements: {e}")
            self.statements = self.default_statements.copy()

    def save_statements(self):
        """Save custom statements to JSON."""
        try:
            with open(self.statements_file, 'w', encoding='utf-8') as f:
                json.dump(self.statements, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error", f"Could not save custom statements: {e}")

    def show_customize_statements_dialog(self):
        """Display a dialog to customize report statements."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Customize Report Statements")
        dialog.geometry("1000x800")
        dialog.transient(self.root)
        dialog.grab_set()

        notebook = ttk.Notebook(dialog)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Tabs for default statements
        for key, default_text in self.default_statements.items():
            frame = ttk.Frame(notebook)
            notebook.add(frame, text=key.capitalize())
            tk.Label(frame, text=f"Edit {key} statement:").pack(pady=5)
            text_widget = tk.Text(frame, width=80, height=20, wrap='word')
            text_widget.pack(pady=5)
            text_widget.insert(tk.END, self.statements.get(key, default_text))

            def save_tab(k=key, tw=text_widget):
                self.statements[k] = tw.get(1.0, tk.END).strip()

            tk.Button(frame, text="Save This Statement", command=save_tab).pack(pady=5)

        # Tab for adding new statements
        add_frame = ttk.Frame(notebook)
        notebook.add(add_frame, text="Add New")
        tk.Label(add_frame, text="Select where to insert the statement:").pack(pady=5)
        placement_options = [
            "At Beginning of Report",
            "Before Incident Summary",
            "After Incident Summary",
            "Before Suspect Information",
            "After Suspect Information",
            "Before Evidence Summary",
            "After Evidence Summary",
            "Before IP Address Analysis",
            "After IP Address Analysis",
            "At End of Report"
        ]
        placement_var = tk.StringVar(value=placement_options[0])
        placement_dropdown = ttk.Combobox(add_frame, textvariable=placement_var, values=placement_options, state="readonly")
        placement_dropdown.pack(pady=5)

        tk.Label(add_frame, text="Select a template (optional):").pack(pady=5)
        template_options = [
            "None",
            "Legal Disclaimer",
            "Investigator Note"
        ]
        template_var = tk.StringVar(value=template_options[0])
        template_dropdown = ttk.Combobox(add_frame, textvariable=template_var, values=template_options, state="readonly")
        template_dropdown.pack(pady=5)

        tk.Label(add_frame, text="Statement Name (e.g., 'warning'):").pack(pady=5)
        key_entry = tk.Entry(add_frame, width=50)
        key_entry.pack(pady=5)

        tk.Label(add_frame, text="Build Condition (leave as 'None' for no condition):").pack(pady=5)
        condition_frame = ttk.Frame(add_frame)
        condition_frame.pack(pady=5, fill='x')

        tk.Label(condition_frame, text="Field:").pack(side=tk.LEFT, padx=5)
        field_var = tk.StringVar(value="None")
        field_dropdown = ttk.Combobox(condition_frame, textvariable=field_var, values=["None", "esp_name"], state="readonly", width=15)
        field_dropdown.pack(side=tk.LEFT, padx=5)

        tk.Label(condition_frame, text="Operator:").pack(side=tk.LEFT, padx=5)
        operator_var = tk.StringVar(value="==")
        operator_dropdown = ttk.Combobox(condition_frame, textvariable=operator_var, values=["==", "in"], state="readonly", width=10)
        operator_dropdown.pack(side=tk.LEFT, padx=5)

        tk.Label(condition_frame, text="Value:").pack(side=tk.LEFT, padx=5)
        esp_values = ["Custom", "Discord", "Dropbox", "Facebook", "Google", "Imgur", "Instagram", "Kik", "MeetMe", "Microsoft", "Reddit", "Roblox", "Snapchat", "Sony", "Synchronoss", "TikTok", "WhatsApp", "X (Twitter)", "Yahoo"]
        value_frame = ttk.Frame(condition_frame)
        value_frame.pack(side=tk.LEFT, padx=5)
        value_dropdown = ttk.Combobox(value_frame, values=esp_values, state="readonly", width=20)
        value_dropdown.pack()
        value_listbox = tk.Listbox(value_frame, selectmode=tk.MULTIPLE, width=20, height=5, exportselection=False)
        for value in esp_values[:-1]:  # Exclude "Custom"
            value_listbox.insert(tk.END, value)
        custom_value_entry = tk.Entry(value_frame, width=20, state="disabled")
        custom_value_entry.pack()

        condition_label = tk.Label(add_frame, text="Condition: None", wraplength=700)
        condition_label.pack(pady=5, anchor='w')

        def update_condition_label(*args):
            field = field_var.get()
            operator = operator_var.get()
            if field == "None":
                condition_label.config(text="Condition: None")
                operator_dropdown.config(state="disabled")
                value_dropdown.pack_forget()
                value_listbox.pack_forget()
                custom_value_entry.pack_forget()
                return
            operator_dropdown.config(state="normal")
            if operator == "==":
                value_dropdown.pack()
                value_listbox.pack_forget()
                custom_value_entry.pack()
                value = value_dropdown.get()
                if value == "Custom":
                    custom_value_entry.config(state="normal")
                    value = custom_value_entry.get().strip() or "[Enter custom value]"
                else:
                    custom_value_entry.config(state="disabled")
                    custom_value_entry.delete(0, tk.END)
                condition = f"{field} {operator} \"{value}\""
            else:  # operator == "in"
                value_dropdown.pack_forget()
                value_listbox.pack()
                custom_value_entry.pack()
                selected_indices = value_listbox.curselection()
                values = [value_listbox.get(i) for i in selected_indices]
                if not values and custom_value_entry.get().strip():
                    values = [custom_value_entry.get().strip()]
                if not values:
                    condition = f"{field} {operator} []"
                else:
                    condition = f"{field} {operator} [{', '.join(f'\"{v}\"' for v in values)}]"
            condition_label.config(text=f"Condition: {condition}")

        field_var.trace("w", update_condition_label)
        operator_var.trace("w", update_condition_label)
        value_dropdown.bind("<<ComboboxSelected>>", update_condition_label)
        value_listbox.bind("<<ListboxSelect>>", update_condition_label)
        custom_value_entry.bind("<KeyRelease>", update_condition_label)

        tk.Label(add_frame, text="Examples: Set Field to 'esp_name', Operator to '==', Value to 'Facebook' for esp_name == \"Facebook\"; or Operator 'in', select multiple Values for esp_name in [\"Facebook\", \"Instagram\"].").pack(pady=5, anchor='w')
        tk.Label(add_frame, text="Tip: Choose a placement and unique name. Templates pre-fill common text. Conditions limit when the statement appears (e.g., only for specific ESPs).").pack(pady=5, anchor='w')

        tk.Label(add_frame, text="Statement Text:").pack(pady=5)
        add_text = tk.Text(add_frame, width=80, height=10, wrap='word')
        add_text.pack(pady=5)

        def apply_template(event=None):
            template = template_var.get()
            add_text.delete(1.0, tk.END)
            if template == "Legal Disclaimer":
                add_text.insert(tk.END, "This report is for official use only and may contain sensitive information protected by law.")
            elif template == "Investigator Note":
                add_text.insert(tk.END, "Investigator's additional observations: [Enter details here].")
            key_entry.delete(0, tk.END)
            if template != "None":
                key_entry.insert(0, template.lower().replace(" ", "_"))

        template_dropdown.bind("<<ComboboxSelected>>", apply_template)

        def add_new():
            placement = placement_var.get()
            name = key_entry.get().strip()
            field = field_var.get()
            operator = operator_var.get()
            condition = ""
            if field != "None":
                if operator == "==":
                    value = value_dropdown.get()
                    if value == "Custom":
                        value = custom_value_entry.get().strip()
                        if not value:
                            messagebox.showwarning("Warning", "Custom value must be non-empty.")
                            return
                    condition = f"{field} {operator} \"{value}\""
                else:  # operator == "in"
                    selected_indices = value_listbox.curselection()
                    values = [value_listbox.get(i) for i in selected_indices]
                    if custom_value_entry.get().strip():
                        values.append(custom_value_entry.get().strip())
                    if not values:
                        messagebox.showwarning("Warning", "Select at least one value or enter a custom value.")
                        return
                    condition = f"{field} {operator} [{', '.join(f'\"{v}\"' for v in values)}]"
            if not name:
                messagebox.showwarning("Warning", "Statement name must be non-empty.")
                return
            prefix_map = {
                "At Beginning of Report": "at_beginning:",
                "Before Incident Summary": "before_incident:",
                "After Incident Summary": "after_incident:",
                "Before Suspect Information": "before_suspect:",
                "After Suspect Information": "after_suspect:",
                "Before Evidence Summary": "before_evidence:",
                "After Evidence Summary": "after_evidence:",
                "Before IP Address Analysis": "before_ip:",
                "After IP Address Analysis": "after_ip:",
                "At End of Report": ""
            }
            prefix = prefix_map[placement]
            full_key = f"{prefix}{name}" if prefix else name
            if full_key in self.statements:
                messagebox.showwarning("Warning", f"Statement name '{name}' already exists for this placement.")
                return
            text = add_text.get(1.0, tk.END).strip()
            if not text:
                messagebox.showwarning("Warning", "Statement text must be non-empty.")
                return
            self.statements[full_key] = {'text': text, 'condition': condition}
            messagebox.showinfo("Success", f"Added new statement: {full_key}")
            key_entry.delete(0, tk.END)
            field_var.set("None")
            operator_var.set("==")
            value_dropdown.set(esp_values[0])
            value_listbox.selection_clear(0, tk.END)
            custom_value_entry.delete(0, tk.END)
            add_text.delete(1.0, tk.END)
            template_var.set("None")
            update_condition_label()

        tk.Button(add_frame, text="Add New Statement", command=add_new).pack(pady=5)

        # Tab for managing custom statements
        manage_frame = ttk.Frame(notebook)
        notebook.add(manage_frame, text="Manage Custom Statements")
        tk.Label(manage_frame, text="Search statements (by key or content):").pack(pady=5)
        search_var = tk.StringVar()
        search_entry = tk.Entry(manage_frame, textvariable=search_var, width=50)
        search_entry.pack(pady=5)

        listbox = tk.Listbox(manage_frame, width=80, height=10)
        listbox.pack(pady=5)
        scrollbar = tk.Scrollbar(manage_frame, orient=tk.VERTICAL)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=listbox.yview)

        tk.Label(manage_frame, text="Selected Statement Text:").pack(pady=5)
        edit_text = tk.Text(manage_frame, width=80, height=8, wrap='word')
        edit_text.pack(pady=5)

        tk.Label(manage_frame, text="Build Condition:").pack(pady=5)
        edit_condition_frame = ttk.Frame(manage_frame)
        edit_condition_frame.pack(pady=5, fill='x')

        tk.Label(edit_condition_frame, text="Field:").pack(side=tk.LEFT, padx=5)
        edit_field_var = tk.StringVar(value="None")
        edit_field_dropdown = ttk.Combobox(edit_condition_frame, textvariable=edit_field_var, values=["None", "esp_name"], state="readonly", width=15)
        edit_field_dropdown.pack(side=tk.LEFT, padx=5)

        tk.Label(edit_condition_frame, text="Operator:").pack(side=tk.LEFT, padx=5)
        edit_operator_var = tk.StringVar(value="==")
        edit_operator_dropdown = ttk.Combobox(edit_condition_frame, textvariable=edit_operator_var, values=["==", "in"], state="readonly", width=10)
        edit_operator_dropdown.pack(side=tk.LEFT, padx=5)

        tk.Label(edit_condition_frame, text="Value:").pack(side=tk.LEFT, padx=5)
        edit_value_frame = ttk.Frame(edit_condition_frame)
        edit_value_frame.pack(side=tk.LEFT, padx=5)
        edit_value_dropdown = ttk.Combobox(edit_value_frame, values=esp_values, state="readonly", width=20)
        edit_value_dropdown.pack()
        edit_value_listbox = tk.Listbox(edit_value_frame, selectmode=tk.MULTIPLE, width=20, height=5, exportselection=False)
        for value in esp_values[:-1]:  # Exclude "Custom"
            edit_value_listbox.insert(tk.END, value)
        edit_custom_value_entry = tk.Entry(edit_value_frame, width=20, state="disabled")
        edit_custom_value_entry.pack()

        edit_condition_label = tk.Label(manage_frame, text="Condition: None", wraplength=700)
        edit_condition_label.pack(pady=5, anchor='w')

        def update_edit_condition_label(*args):
            field = edit_field_var.get()
            operator = edit_operator_var.get()
            if field == "None":
                edit_condition_label.config(text="Condition: None")
                edit_operator_dropdown.config(state="disabled")
                edit_value_dropdown.pack_forget()
                edit_value_listbox.pack_forget()
                edit_custom_value_entry.pack_forget()
                return
            edit_operator_dropdown.config(state="normal")
            if operator == "==":
                edit_value_dropdown.pack()
                edit_value_listbox.pack_forget()
                edit_custom_value_entry.pack()
                value = edit_value_dropdown.get()
                if value == "Custom":
                    edit_custom_value_entry.config(state="normal")
                    value = edit_custom_value_entry.get().strip() or "[Enter custom value]"
                else:
                    edit_custom_value_entry.config(state="disabled")
                    edit_custom_value_entry.delete(0, tk.END)
                condition = f"{field} {operator} \"{value}\""
            else:  # operator == "in"
                edit_value_dropdown.pack_forget()
                edit_value_listbox.pack()
                edit_custom_value_entry.pack()
                selected_indices = edit_value_listbox.curselection()
                values = [edit_value_listbox.get(i) for i in selected_indices]
                if edit_custom_value_entry.get().strip():
                    values.append(edit_custom_value_entry.get().strip())
                if not values:
                    condition = f"{field} {operator} []"
                else:
                    condition = f"{field} {operator} [{', '.join(f'\"{v}\"' for v in values)}]"
            edit_condition_label.config(text=f"Condition: {condition}")

        edit_field_var.trace("w", update_edit_condition_label)
        edit_operator_var.trace("w", update_edit_condition_label)
        edit_value_dropdown.bind("<<ComboboxSelected>>", update_edit_condition_label)
        edit_value_listbox.bind("<<ListboxSelect>>", update_edit_condition_label)
        edit_custom_value_entry.bind("<KeyRelease>", update_edit_condition_label)

        def update_listbox(search_query=""):
            listbox.delete(0, tk.END)
            custom_statements = {k: v for k, v in self.statements.items() if k not in self.default_statements}
            for key in sorted(custom_statements.keys()):
                if search_query.lower() in key.lower() or (isinstance(custom_statements[key], dict) and search_query.lower() in custom_statements[key]['text'].lower()):
                    prefix_map = {
                        "at_beginning:": "At Beginning of Report",
                        "before_incident:": "Before Incident Summary",
                        "after_incident:": "After Incident Summary",
                        "before_suspect:": "Before Suspect Information",
                        "after_suspect:": "After Suspect Information",
                        "before_evidence:": "Before Evidence Summary",
                        "after_evidence:": "After Evidence Summary",
                        "before_ip:": "Before IP Address Analysis",
                        "after_ip:": "After IP Address Analysis",
                        "": "At End of Report"
                    }
                    placement = next((v for p, v in prefix_map.items() if key.startswith(p)), "At End of Report")
                    listbox.insert(tk.END, f"{key} ({placement})")

        def on_select(event):
            selection = listbox.curselection()
            if selection:
                key = listbox.get(selection[0]).split(' (')[0]
                edit_text.delete(1.0, tk.END)
                edit_field_var.set("None")
                edit_operator_var.set("==")
                edit_value_dropdown.set(esp_values[0])
                edit_value_listbox.selection_clear(0, tk.END)
                edit_custom_value_entry.delete(0, tk.END)
                stmt = self.statements.get(key)
                if isinstance(stmt, dict):
                    edit_text.insert(tk.END, stmt['text'])
                    condition = stmt.get('condition', '')
                    if condition:
                        parts = condition.split()
                        if len(parts) >= 3:
                            edit_field_var.set(parts[0])
                            edit_operator_var.set(parts[1])
                            values = ' '.join(parts[2:]).strip('[]').split(', ')
                            values = [v.strip('\"') for v in values if v.strip('\"')]
                            if len(values) == 1 and values[0] in esp_values:
                                edit_value_dropdown.set(values[0])
                            elif values:
                                for value in values:
                                    if value in esp_values[:-1]:
                                        index = esp_values.index(value)
                                        edit_value_listbox.select_set(index)
                                    else:
                                        edit_value_dropdown.set("Custom")
                                        edit_custom_value_entry.config(state="normal")
                                        edit_custom_value_entry.insert(0, value)
                else:
                    edit_text.insert(tk.END, stmt)
                update_edit_condition_label()

        def save_edit():
            selection = listbox.curselection()
            if selection:
                key = listbox.get(selection[0]).split(' (')[0]
                new_text = edit_text.get(1.0, tk.END).strip()
                if not new_text:
                    messagebox.showwarning("Warning", "Statement text must be non-empty.")
                    return
                field = edit_field_var.get()
                operator = edit_operator_var.get()
                condition = ""
                if field != "None":
                    if operator == "==":
                        value = edit_value_dropdown.get()
                        if value == "Custom":
                            value = edit_custom_value_entry.get().strip()
                            if not value:
                                messagebox.showwarning("Warning", "Custom value must be non-empty.")
                                return
                        condition = f"{field} {operator} \"{value}\""
                    else:  # operator == "in"
                        selected_indices = edit_value_listbox.curselection()
                        values = [edit_value_listbox.get(i) for i in selected_indices]
                        if edit_custom_value_entry.get().strip():
                            values.append(edit_custom_value_entry.get().strip())
                        if not values:
                            messagebox.showwarning("Warning", "Select at least one value or enter a custom value.")
                            return
                        condition = f"{field} {operator} [{', '.join(f'\"{v}\"' for v in values)}]"
                self.statements[key] = {'text': new_text, 'condition': condition}
                messagebox.showinfo("Success", f"Updated statement: {key}")
                update_listbox(search_var.get())

        def delete_statement():
            selection = listbox.curselection()
            if selection:
                key = listbox.get(selection[0]).split(' (')[0]
                if messagebox.askyesno("Confirm", f"Delete statement '{key}'?"):
                    del self.statements[key]
                    messagebox.showinfo("Success", f"Deleted statement: {key}")
                    update_listbox(search_var.get())
                    edit_text.delete(1.0, tk.END)
                    edit_custom_value_entry.delete(0, tk.END)

        def drag_start(event):
            listbox.drag_start_index = listbox.nearest(event.y)

        def drag_motion(event):
            pass  # Optional: Add visual feedback if desired

        def drop(event):
            end_index = listbox.nearest(event.y)
            if listbox.drag_start_index != end_index:
                custom_statements = {k: v for k, v in self.statements.items() if k not in self.default_statements}
                keys = list(custom_statements.keys())
                start_key = keys[listbox.drag_start_index]
                keys.pop(listbox.drag_start_index)
                keys.insert(end_index, start_key)
                new_statements = {k: custom_statements[k] for k in keys}
                for key in self.default_statements:
                    new_statements[key] = self.statements[key]
                self.statements = new_statements
                update_listbox(search_var.get())

        listbox.bind('<<ListboxSelect>>', on_select)
        listbox.bind('<Button-1>', drag_start)
        listbox.bind('<B1-Motion>', drag_motion)
        listbox.bind('<ButtonRelease-1>', drop)
        search_var.trace("w", lambda *args: update_listbox(search_var.get()))
        update_listbox()

        tk.Button(manage_frame, text="Save Edited Statement", command=save_edit).pack(pady=5)
        tk.Button(manage_frame, text="Delete Selected Statement", command=delete_statement).pack(pady=5)

        def close_and_save():
            self.save_statements()
            dialog.destroy()

        tk.Button(dialog, text="Save All and Close", command=close_and_save).pack(pady=10)
        tk.Button(dialog, text="Cancel", command=dialog.destroy).pack(pady=5)

    def show_choose_statements_dialog(self):
        """Display a dialog to choose statements for inclusion in the report."""
        if not self.json_file_path:
            messagebox.showwarning("Warning", "Please select a JSON file first.")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Choose Statements for Report")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        dialog.grab_set()

        tk.Label(dialog, text="Select statements to include in the report (all checked by default):").pack(pady=5)
        frame = ttk.Frame(dialog)
        frame.pack(fill='both', expand=True, padx=10, pady=10)
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        check_vars = {}
        for key in sorted(self.statements.keys()):
            var = tk.BooleanVar(value=key in self.selected_statements)
            check_vars[key] = var
            prefix_map = {
                "at_beginning:": "At Beginning of Report",
                "before_incident:": "Before Incident Summary",
                "after_incident:": "After Incident Summary",
                "before_suspect:": "Before Suspect Information",
                "after_suspect:": "After Suspect Information",
                "before_evidence:": "Before Evidence Summary",
                "after_evidence:": "After Evidence Summary",
                "before_ip:": "Before IP Address Analysis",
                "after_ip:": "After IP Address Analysis",
                "": "At End of Report"
            }
            placement = next((v for p, v in prefix_map.items() if key.startswith(p)), "Default Statement" if key in self.default_statements else "At End of Report")
            condition = self.statements[key].get('condition', '') if isinstance(self.statements[key], dict) else ''
            ttk.Checkbutton(scrollable_frame, text=f"{key} ({placement}) [Condition: {condition if condition else 'None'}]", variable=var).pack(anchor='w', pady=2)

        def save_selection():
            self.selected_statements = {key for key, var in check_vars.items() if var.get()}
            messagebox.showinfo("Success", "Statement selection saved.")
            dialog.destroy()

        def preview_report():
            data = self.load_json()
            if data is None:
                messagebox.showwarning("Warning", "Failed to load JSON for preview.")
                return
            selected = {key for key, var in check_vars.items() if var.get()}
            temp_statements = {k: v for k, v in self.statements.items() if k in selected}
            original_statements = self.statements
            original_selected = self.selected_statements
            self.statements = temp_statements
            self.selected_statements = selected
            preview_report = self.generate_police_report(data)[:1000]  # Limit to 1000 chars for preview
            self.statements = original_statements
            self.selected_statements = original_selected

            preview_dialog = tk.Toplevel(dialog)
            preview_dialog.title("Report Preview")
            preview_dialog.geometry("600x400")
            preview_dialog.transient(dialog)
            preview_dialog.grab_set()
            tk.Label(preview_dialog, text="Report Preview (first 1000 characters):").pack(pady=5)
            text_area = scrolledtext.ScrolledText(preview_dialog, width=80, height=20, wrap='word')
            text_area.pack(pady=5)
            text_area.insert(tk.END, preview_report)
            text_area.config(state='disabled')
            tk.Button(preview_dialog, text="Close", command=preview_dialog.destroy).pack(pady=5)

        tk.Button(dialog, text="Preview Report", command=preview_report).pack(pady=5)
        tk.Button(dialog, text="Save Selection", command=save_selection).pack(pady=5)
        tk.Button(dialog, text="Cancel", command=dialog.destroy).pack(pady=5)

    def save_investigator_info(self):
        """Save investigator info to JSON file."""
        data = {
            'name': self.investigator_name,
            'title': self.investigator_title
        }
        try:
            with open(self.investigator_file, 'w', encoding='utf-8') as f:
                json.dump(data, f)
        except Exception as e:
            messagebox.showerror("Error", f"Could not save investigator info: {e}")

    def get_investigator_info(self):
        """Prompt the user to input their name and title."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Investigator Information")
        dialog.geometry("300x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text="Enter Your Name:").pack(pady=5)
        name_entry = tk.Entry(dialog, width=30)
        name_entry.pack(pady=5)
        name_entry.insert(0, self.investigator_name)
        
        tk.Label(dialog, text="Enter Your Title:").pack(pady=5)
        title_entry = tk.Entry(dialog, width=30)
        title_entry.pack(pady=5)
        title_entry.insert(0, self.investigator_title)
        
        def save_info():
            self.investigator_name = name_entry.get().strip()
            self.investigator_title = title_entry.get().strip()
            if not self.investigator_name or not self.investigator_title:
                messagebox.showwarning("Warning", "Please enter both name and title.")
            else:
                self.save_investigator_info()
                dialog.destroy()
        
        tk.Button(dialog, text="Save", command=save_info).pack(pady=10)
        dialog.protocol("WM_DELETE_WINDOW", lambda: self.root.quit())
        self.root.wait_window(dialog)

    def show_credentials_dialog(self):
        """Display a dialog to update MaxMind credentials (optional)."""
        dialog = tk.Toplevel(self.root)
        dialog.title("MaxMind Credentials")
        dialog.geometry("400x350")
        dialog.transient(self.root)
        dialog.grab_set()

        tk.Label(
            dialog,
            text="For IP Geo Lookup, enter your MaxMind credentials below (optional).\n\nCreate new account at "
        ).pack(pady=5)
        signup_label = tk.Label(dialog, text="this link", fg="blue", cursor="hand2")
        signup_label.pack(pady=5)
        signup_url = "https://www.maxmind.com/en/geolite2/signup?utm_source=kb&utm_medium=kb-link&utm_campaign=kb-create-account"
        signup_label.bind("<Button-1>", lambda e: webbrowser.open_new(signup_url))

        tk.Label(dialog, text="MaxMind Account ID:").pack(pady=5)
        maxmind_id_entry = tk.Entry(dialog, width=30)
        maxmind_id_entry.pack(pady=5)
        maxmind_id_entry.insert(0, self.maxmind_id_entry.get())

        tk.Label(dialog, text="MaxMind License Key:").pack(pady=5)
        maxmind_key_entry = tk.Entry(dialog, width=30, show="*")
        maxmind_key_entry.pack(pady=5)
        maxmind_key_entry.insert(0, self.maxmind_key_entry.get())

        def save_and_close():
            account_id = maxmind_id_entry.get().strip()
            license_key = maxmind_key_entry.get().strip()
            self.maxmind_id_entry.delete(0, tk.END)
            self.maxmind_id_entry.insert(0, account_id)
            self.maxmind_key_entry.delete(0, tk.END)
            self.maxmind_key_entry.insert(0, license_key)
            self.save_credentials()
            dialog.destroy()
            if account_id and license_key:
                messagebox.showinfo("Success", "MaxMind credentials updated successfully.")
            else:
                messagebox.showinfo("Info", "MaxMind credentials skipped - geolocation will not be available.")

        tk.Button(dialog, text="Save", command=save_and_close).pack(pady=10)
        tk.Button(dialog, text="Skip (No Geolocation)", command=dialog.destroy).pack(pady=5)
        return dialog

    def save_credentials(self):
        """Save MaxMind credentials to file."""
        creds = {
            'account_id': self.maxmind_id_entry.get(),
            'license_key': self.maxmind_key_entry.get()
        }
        try:
            with open(self.credentials_file, 'w', encoding='utf-8') as f:
                json.dump(creds, f)
        except Exception as e:
            messagebox.showerror("Error", f"Could not save credentials: {e}")

    def update_investigator_info(self):
        """Allow user to update investigator information via menu."""
        self.get_investigator_info()

    def create_menu(self):
        """Create a menu bar with File and Help options."""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        self.file_menu = tk.Menu(menubar, tearoff=0)
        self.file_menu.add_command(label="Open File", command=self.open_file, accelerator="Ctrl+O")
        self.recent_menu = tk.Menu(self.file_menu, tearoff=0)
        self.file_menu.add_cascade(label="Recent Files", menu=self.recent_menu)
        self.file_menu.add_command(label="Save Report As", command=self.save_report_as, accelerator="Ctrl+S")
        self.file_menu.add_command(label="Export as Text", command=self.export_as_text)
        self.file_menu.add_command(label="Export IP Data to Excel", command=self.export_ip_data_to_excel)
        self.file_menu.add_command(label="Export Evidence to Excel", command=self.export_evidence_to_excel)
        self.file_menu.add_command(label="Clear Output", command=self.clear_output)
        self.file_menu.add_command(label="New Analysis", command=self.new_analysis, accelerator="Ctrl+N")
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Exit", command=self.exit_app)
        menubar.add_cascade(label="File", menu=self.file_menu)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="About", command=self.show_about)
        help_menu.add_command(label="Help", command=self.show_help)
        help_menu.add_command(label="Update MaxMind Credentials", command=self.show_credentials_dialog)
        help_menu.add_command(label="Update Investigator Info", command=self.update_investigator_info)
        help_menu.add_command(label="Update ARIN API Key", command=self.show_arin_credentials_dialog)
        help_menu.add_command(label="Customize Statements", command=self.show_customize_statements_dialog)
        help_menu.add_command(label="Check for Updates", command=self.check_for_updates)
        menubar.add_cascade(label="Help", menu=help_menu)
        
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-s>", lambda e: self.save_report_as())
        self.root.bind("<Control-n>", lambda e: self.new_analysis())
        
        self.update_recent_files_menu()

    def show_about(self):
        """Display an About dialog."""
        messagebox.showinfo("About", "CyberTip Analysis Tool & Report Generator\nVersion 2.0\nDeveloped by Patrick Koebbe with Grok AI\n\n"
                                    "This tool analyzes CyberTipline JSON reports and generates formatted DOCX reports.")

    def show_help(self):
        """Display a Help dialog."""
        messagebox.showinfo("Help", "To use the CyberTip Analysis Tool & Report Generator:\n\n"
                                    "1. Select a JSON file via 'File > Open File' or the 'Browse' button.\n\n"
                                    "2. Click 'Choose Statements' to select which statements to include in the report.\n\n"
                                    "3. Click 'Analyze Report' or use 'File > New Analysis' to process and generate a report.\n\n"
                                    "4. Save the report with 'File > Save Report As' (Ctrl+S) or 'File > Export as Text'.\n\n"
                                    "5. Use 'File > Recent Files' to quickly reopen recent files.\n\n"
                                    "6. Use 'File > Clear Output' to clear the program's textbox field.\n\n"
                                    "7. Use 'File > New Analysis' to initiate a new CyberTip analysis without closing the program.\n\n"
                                    "8. Update MaxMind credentials via 'Help > Update MaxMind Credentials'.\n\n"
                                    "9. Customize report statements via 'Help > Customize Statements' to add, edit, or delete text.\n\n"
                                    "HOT KEYS\n"
                                    "- Ctrl+O: Open new .json file\n"
                                    "- Ctrl+S: Save report as\n"
                                    "- Ctrl+N: Start new analysis\n\n"
                                    "Supported ESPs (others will work but need to be verified):\n"
                                    "- Discord\n- Dropbox\n- Facebook\n- Google\n- Imgur\n- Instagram\n- Kik\n- MeetMe\n"
                                    "- Microsoft\n- Reddit\n- Roblox\n- Snapchat\n- Sony\n- Synchronoss\n- TikTok\n- WhatsApp\n- X (Twitter)\n- Yahoo\n\n"
                                    "For support, contact Patrick Koebbe - Patrick.Koebbe@gmail.com")

    def create_widgets(self):
        title_label = tk.Label(self.root, text="CAT-RG", font=("Arial", 18, "bold"))
        title_label.pack(pady=5)
        
        title_label = tk.Label(self.root, text="CyberTip Analysis Tool & Report Generator", font=("Arial", 11, "bold", "underline"))
        title_label.pack(pady=5)

        if self.logo_image:
            logo_label = tk.Label(self.root, image=self.logo_image)
            logo_label.pack(pady=10)
        
        tk.Label(self.root, text="CyberTip .json File:").pack(pady=5)
        self.file_entry = tk.Entry(self.root, width=50)
        self.file_entry.pack(pady=5)
        tk.Button(self.root, text="Browse to .json", command=self.open_file).pack(pady=5)
        
        self.choose_statements_button = tk.Button(self.root, text="Choose Statements", command=self.show_choose_statements_dialog, state="disabled")
        self.choose_statements_button.pack(pady=5)
        
        tk.Button(self.root, text="Analyze CyberTip", command=self.analyze_report).pack(pady=5)
        
        self.status_label = tk.Label(self.root, text="Ready", font=("Arial", 10), fg="blue")
        self.status_label.pack(pady=5)
     
        self.output_text = scrolledtext.ScrolledText(self.root, width=90, height=30)
        self.output_text.pack(pady=10)

    def open_file(self):
        """Open a JSON file and add to recent files."""
        self.json_file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if self.json_file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, self.json_file_path)
            if self.json_file_path not in self.recent_files:
                self.recent_files.insert(0, self.json_file_path)
                self.recent_files = self.recent_files[:5]
                self.save_recent_files()
            self.update_recent_files_menu()
            self.choose_statements_button.config(state="normal")

    def save_recent_files(self):
        """Save recent files to JSON file."""
        try:
            with open(self.recent_files_file, 'w', encoding='utf-8') as f:
                json.dump(self.recent_files, f)
        except Exception as e:
            print(f"Error saving recent files: {e}")

    def update_recent_files_menu(self):
        """Update the Recent Files submenu."""
        self.recent_menu.delete(0, tk.END)
        if not self.recent_files:
            self.recent_menu.add_command(label="No recent files", state="disabled")
        else:
            for file_path in self.recent_files:
                self.recent_menu.add_command(label=file_path, command=lambda f=file_path: self.load_recent_file(f))

    def load_recent_file(self, file_path):
        """Load a file from the recent files list."""
        if os.path.exists(file_path):
            self.json_file_path = file_path
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, self.json_file_path)
            self.choose_statements_button.config(state="normal")
        else:
            messagebox.showwarning("Warning", f"File not found: {file_path}")
            self.recent_files.remove(file_path)
            self.save_recent_files()
            self.update_recent_files_menu()

    def save_report_as(self):
        """Save the current report as a DOCX file."""
        current_text = self.output_text.get(1.0, tk.END).strip()
        if not current_text:
            messagebox.showwarning("Warning", "No report available to save.")
            return
        police_report, ip_report = current_text.split("IP ADDRESS ANALYSIS:", 1)
        police_report += "IP ADDRESS ANALYSIS:"
        ip_report = "IP ADDRESS ANALYSIS:" + ip_report
        try:
            filename = self.save_report(police_report, ip_report)
            messagebox.showinfo("Success", f"Report saved as {filename}")
        except ValueError as e:
            messagebox.showwarning("Warning", str(e))

    def export_as_text(self):
        """Export the current report as a text file."""
        current_text = self.output_text.get(1.0, tk.END).strip()
        if not current_text:
            messagebox.showwarning("Warning", "No report available to export.")
            return
        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt")],
            initialfile=f"cybertip_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        )
        if filename:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(current_text)
            messagebox.showinfo("Success", f"Report exported as {filename}")

    def clear_output(self):
        """Clear the output text and reset the file entry."""
        self.output_text.delete(1.0, tk.END)
        self.file_entry.delete(0, tk.END)
        self.json_file_path = None
        self.ip_data.clear()
        self.all_ip_data.clear()
        self.choose_statements_button.config(state="disabled")
        self.selected_statements = set(self.statements.keys())

    def new_analysis(self):
        """Clear output, open a new file, and analyze it."""
        self.clear_output()
        self.open_file()
        if self.json_file_path:
            self.analyze_report()

    def exit_app(self):
        """Exit the application."""
        self.save_recent_files()
        self.root.destroy()

    def load_json(self) -> dict:
        try:
            with open(self.json_file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            messagebox.showerror("Error", f"Error loading JSON file: {e}")
            return None

    def is_valid_ip(self, ip: str) -> bool:
        """Check if the provided string is a valid IPv4 or IPv6 address."""
        try:
            ipaddress.ip_address(ip)
            return True
        except ValueError:
            return False

    def extract_ip_addresses(self, data: dict) -> None:
        """Extract all valid IPv4 or IPv6 addresses from the JSON data, avoiding duplicates."""
        self.ip_data.clear()
        self.all_ip_data.clear()
        seen_entries = set()
        temp_entries = []
        unique_ips = set()
        
        def collect_entry(ip, entry):
            unique_key = (ip, entry['datetime'], entry.get('port'), entry['event'])
            if unique_key not in seen_entries:
                seen_entries.add(unique_key)
                temp_entries.append((ip, entry))
                unique_ips.add(ip)

        self.status_label.config(text="Extracting IPs from persons...")
        self.root.update_idletasks()
        reported_persons = data.get('reportedInformation', {}).get('reportedPeople', {}).get('reportedPersons', [])
        for person in reported_persons:
            source_info = person.get('sourceInformation')
            if source_info:
                for capture in source_info.get('sourceCaptures', []):
                    ip = capture.get('value')
                    if ip and self.is_valid_ip(ip):
                        entry = {
                            'datetime': capture.get('dateTime'),
                            'port': capture.get('port'),
                            'event': capture.get('eventName')
                        }
                        collect_entry(ip, entry)

        self.status_label.config(text="Extracting IPs from files...")
        self.root.update_idletasks()
        uploaded_files = data.get('reportedInformation', {}).get('uploadedFiles', {}).get('uploadedFiles', [])
        for file in uploaded_files:
            source_info = file.get('sourceInformation')
            if source_info:
                for capture in source_info.get('sourceCaptures', []):
                    ip = capture.get('value')
                    if ip and self.is_valid_ip(ip):
                        entry = {
                            'datetime': capture.get('dateTime'),
                            'port': capture.get('port'),
                            'event': capture.get('eventName')
                        }
                        collect_entry(ip, entry)
        
        for ip, entry in temp_entries:
            if ip not in self.all_ip_data:
                self.all_ip_data[ip] = []
            self.all_ip_data[ip].append(entry)
        
        max_ips = 50
        num_unique_ips = len(unique_ips)
        print(f"DEBUG: Found {num_unique_ips} unique IPs")
        if num_unique_ips > max_ips:
            self.status_label.config(text="IP limit exceeded, prompting user...")
            self.root.update_idletasks()
            process_all = messagebox.askyesno(
                "IP Limit Exceeded",
                f"{num_unique_ips} unique IPs found (more than {max_ips}).\n"
                "Query all IPs for geolocation/WHOIS? (Yes: All - may be slow; No: First 50 - faster)\n"
                "All IPs will be listed in the report regardless."
            )
            if not process_all:
                selected_ips = list(unique_ips)[:max_ips]
                temp_entries = [(ip, entry) for ip, entry in temp_entries if ip in selected_ips]
                self.status_label.config(text=f"Querying first {max_ips} IPs (all {num_unique_ips} will be listed)")
            else:
                self.status_label.config(text=f"Querying all {num_unique_ips} IPs... this may take a while!")
        else:
            self.status_label.config(text=f"Extracted {num_unique_ips} IPs for querying")
        
        self.root.update_idletasks()
        
        self.ip_data.clear()
        for ip, entry in temp_entries:
            if ip not in self.ip_data:
                self.ip_data[ip] = []
            self.ip_data[ip].append(entry)
        
        print(f"DEBUG: Querying {len(self.ip_data)} IPs")

    def query_maxmind(self, ip: str) -> dict:
        account_id = self.maxmind_id_entry.get()
        license_key = self.maxmind_key_entry.get()
        if not account_id or not license_key:
            return {"error": "MaxMind credentials not provided"}
        try:
            url = f"https://geolite.info/geoip/v2.1/city/{ip}"
            response = requests.get(url, auth=(account_id, license_key), timeout=10)
            if response.status_code == 200:
                self.save_credentials()
                return response.json()
            return {"error": f"MaxMind request failed with status {response.status_code}"}
        except requests.Timeout:
            return {"error": "MaxMind query timed out after 10 seconds"}
        except Exception as e:
            return {"error": str(e)}

    def query_arin(self, ip: str) -> dict:
        try:
            base_url = f"https://whois.arin.net/rest/ip/{ip}.json"
            url = f"{base_url}?apikey={self.arin_api_key}" if self.arin_api_key else base_url
            headers = {'Accept': 'application/json'}
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                return response.json()
            return {"error": f"ARIN request failed with status {response.status_code}"}
        except requests.Timeout:
            return {"error": "ARIN query timed out after 10 seconds"}
        except Exception as e:
            return {"error": str(e)}

    def evaluate_condition(self, condition: str, data: dict) -> bool:
        """Evaluate a condition string against report data."""
        if not condition:
            return True
        try:
            esp_name = data.get('reportedInformation', {}).get('reportingEsp', {}).get('espName', 'N/A')
            if '==' in condition:
                # Handle == as a case-insensitive substring match
                field, value = condition.split('==', 1)
                field = field.strip()
                value = value.strip().strip('"')
                if field == 'esp_name':
                    return value.lower() in esp_name.lower()
                return eval(condition, {"esp_name": esp_name})
            else:
                # Handle 'in' and other operators as before
                return eval(condition, {"esp_name": esp_name})
        except Exception as e:
            print(f"Error evaluating condition '{condition}': {e}")
            return False

    def get_custom_statements(self, prefix: str, data: dict) -> str:
        """Retrieve custom statements for a given prefix, respecting conditions."""
        custom_statements = {k: v for k, v in self.statements.items() if k not in self.default_statements and k.startswith(prefix) and k in self.selected_statements}
        statements = []
        for key, value in sorted(custom_statements.items()):
            if isinstance(value, dict):
                if self.evaluate_condition(value.get('condition', ''), data):
                    sub_key = key[len(prefix):].strip()
                    if sub_key:
                        statements.append(f"{sub_key.upper()}: {value['text']}")
                    else:
                        statements.append(value['text'])
            else:
                sub_key = key[len(prefix):].strip()
                if sub_key:
                    statements.append(f"{sub_key.upper()}: {value}")
                else:
                    statements.append(value)
        return "\n\n" + "\n\n".join(statements) + "\n\n" if statements else ""

    def get_end_custom_statements(self, data: dict) -> str:
        """Retrieve custom statements without specific prefixes."""
        prefixes = ["at_beginning:", "before_incident:", "after_incident:", "before_suspect:", "after_suspect:", "before_evidence:", "after_evidence:", "before_ip:", "after_ip:"]
        custom_statements = {k: v for k, v in self.statements.items() if k not in self.default_statements and not any(k.startswith(p) for p in prefixes) and k in self.selected_statements}
        statements = []
        for key, value in sorted(custom_statements.items()):
            if isinstance(value, dict):
                if self.evaluate_condition(value.get('condition', ''), data):
                    statements.append(f"{key.upper()}: {value['text']}")
            else:
                statements.append(f"{key.upper()}: {value}")
        if statements:
            return "\n\nCUSTOM STATEMENTS:\n" + "\n\n".join(statements) + "\n\n"
        return ""

    def get_end_custom_statements(self, data: dict) -> str:
        """Retrieve custom statements without specific prefixes."""
        prefixes = ["at_beginning:", "before_incident:", "after_incident:", "before_suspect:", "after_suspect:", "before_evidence:", "after_evidence:", "before_ip:", "after_ip:"]
        custom_statements = {k: v for k, v in self.statements.items() if k not in self.default_statements and not any(k.startswith(p) for p in prefixes) and k in self.selected_statements}
        statements = []
        for key, value in sorted(custom_statements.items()):
            if isinstance(value, dict):
                if self.evaluate_condition(value.get('condition', ''), data):
                    statements.append(f"{key.upper()}: {value['text']}")
            else:
                statements.append(f"{key.upper()}: {value}")
        if statements:
            return "\n\nCUSTOM STATEMENTS:\n" + "\n\n".join(statements)
        return ""

    def generate_police_report(self, data: dict) -> str:
        execution_date = datetime.now().strftime('%m-%d-%Y')
        cybertip_number = data.get('reportId', 'N/A')
        date_received = data.get('dateReceived', 'N/A')
        if date_received != 'N/A':
            try:
                date_received = datetime.strptime(date_received, '%Y-%m-%dT%H:%M:%SZ').strftime('%m/%d/%Y')
            except ValueError:
                date_received = 'N/A'
        intro = f"On {execution_date}, I, {self.investigator_title} {self.investigator_name}, reviewed Cybertip #{cybertip_number}, which was received by the National Center for Missing and Exploited Children (NCMEC) on {date_received}. I observed the following information regarding this CyberTip:\n\n"
        
        incident = data.get('reportedInformation', {}).get('incidentSummary', {})
        reporter = data.get('reportedInformation', {}).get('reportingEsp', {})
        esp_name = reporter.get('espName', 'N/A')
        incident_details = ""
        if incident.get('incidentType') is not None and incident.get('incidentType') != 'N/A':
            incident_details += f"Incident Type: {incident.get('incidentType')}\n"
        if incident.get('incidentDateTime') is not None and incident.get('incidentDateTime') != 'N/A':
            try:
                incident_datetime = datetime.strptime(incident.get('incidentDateTime'), '%Y-%m-%dT%H:%M:%SZ')
                formatted_datetime = incident_datetime.strftime('%m/%d/%Y %H:%M:%S UTC')
                incident_details += f"Incident Date/Time: {formatted_datetime}\n"
            except ValueError:
                incident_details += f"Incident Date/Time: {incident.get('incidentDateTime')}\n"
        
        if esp_name is not None and esp_name != 'N/A':
            reported_by = esp_name
            if "Microsoft" in esp_name:
                last_name = reporter.get('lastName', 'N/A')
                if last_name == "Microsoft BingImage":
                    reported_by += ", BingImage"
            incident_details += f"Reported By: {reported_by}\n"
        
        incident_desc = incident.get('incidentDateTimeDescription')
        if incident_desc is not None and incident_desc != 'N/A':
            incident_details += f"Incident Date/Time Description: {incident_desc}\n"

        if "Reddit" in esp_name:
            reported_infos = data.get('reportedInformation', {}).get('additionalInformations', [])
            if reported_infos:
                for info in reported_infos:
                    info_text = info.get('value', 'N/A')
                    if info_text != 'N/A':
                        incident_details += f"Additional Information: {info_text}\n"

        if "Microsoft" in esp_name and "BingImage" in str(reporter) and "bingimage" in self.selected_statements:
            bingimage_statement = self.statements.get("bingimage", self.default_statements["bingimage"])
            incident_details += f"{bingimage_statement}\n\n"

        if "X Corp" in esp_name and "xcorp" in self.selected_statements:
            xcorp_statement = self.statements.get("xcorp", self.default_statements["xcorp"])
            incident_details += f"{xcorp_statement}\n"

        email_incidents = data.get('reportedInformation', {}).get('incidentDetails', {}).get('emailIncident', [])
        email_data_by_msg_id = {}
        if email_incidents:
            for email in email_incidents:
                contents = email.get('contents', [{}])[0].get('value', '')
                sent_date, from_email, to_email, msg_id = None, None, None, None
                for line in contents.split('\n'):
                    if line.startswith('Sent Date:'): sent_date = line[len('Sent Date:'):].strip()
                    elif line.startswith('From:'): from_email = line[len('From:'):].strip()
                    elif line.startswith('To:'): to_email = line[len('To:'):].strip()
                    elif line.startswith('X-Ymail-Msg-Id:'): msg_id = line[len('X-Ymail-Msg-Id:'):].strip()
                if msg_id:
                    email_data_by_msg_id[msg_id] = {'sent_date': sent_date, 'from': from_email, 'to': to_email}
        
        report = self.get_custom_statements("at_beginning:", data) + intro + self.get_custom_statements("before_incident:", data) + "INCIDENT SUMMARY:\n" + incident_details + self.get_custom_statements("after_incident:", data) + "\n"

        suspect_details = ""
        persons = data.get('reportedInformation', {}).get('reportedPeople', {}).get('reportedPersons', [])
        if persons:
            suspect_details += "SUSPECT INFORMATION:\n"
            for person in persons:
                person_fields = {
                    "First Name": person.get('firstName'), "Middle Name": person.get('middleName'), "Last Name": person.get('lastName'),
                    "Preferred Name": person.get('preferredName'), "Gender": person.get('gender'), "Preferred Pronouns": person.get('preferredPronouns'),
                    "Date of Birth": person.get('dateOfBirth'), "Approximate Age": person.get('approximateAge'), "Physical Description": person.get('physicalDescription'),
                    "Vehicle Description": person.get('vehicleDescription'), "Vehicle Tag Number": person.get('vehicleTagNumber'), "Occupation": person.get('occupation'),
                    "ESP Service": person.get('espService'), "ESP User ID": person.get('espUserId'), "IP Address": person.get('ipAddress'),
                    "Relationship to Reporter": person.get('relationshipToReporter'), "Relationship to Child Victim": person.get('relationshipToChildVictim'),
                    "Access to Child Victim": person.get('accessToChildVictim'), "Access to Children": person.get('accessToChildren'),
                    "Access to Firearms": person.get('accessToFirearms'), "Convicted Sex Offender": person.get('convictedSexOffender'),
                    "Aware of Report": person.get('awareOfReport'), "Gang Affiliation": person.get('gangAffiliation')
                }
                
                for field_name, field_value in person_fields.items():
                    if field_value is not None and field_value != 'N/A':
                        suspect_details += f"{field_name}: {field_value}\n"

                screen_name = person.get('screenName', {}) if person.get('screenName') is not None else {}
                if screen_name.get('value') is not None and screen_name.get('value') != 'N/A':
                    suspect_details += f"Screen Name: {screen_name.get('value')}\n"

                emails_data = person.get('emails')
                emails = emails_data.get('emails', []) if emails_data is not None else []
                if emails:
                    for email in emails:
                        if email.get('value') is not None and email.get('value') != 'N/A':
                            suspect_details += f"Email: {email.get('value')} "
                            if email.get('value') == "georgehawley454@gmail.com" and "MeetMe" in esp_name:
                                suspect_details += "(Note: This email was voluntarily provided by the user and may not be verified by MeetMe)\n"
                            elif email.get('verified') is not None and email.get('verified') != 'N/A':
                                suspect_details += f"(Verified: {email.get('verified')})\n"
                            else:
                                suspect_details += "\n"

                addresses_data = person.get('addresses')
                if addresses_data is not None:
                    addresses = addresses_data.get('addresses', [])
                    if addresses:
                        for addr in addresses:
                            addr_str = ""
                            if addr.get('street1') is not None and addr.get('street1') != 'N/A': addr_str += f"{addr.get('street1')} "
                            if addr.get('street2') is not None and addr.get('street2') != 'N/A': addr_str += f"{addr.get('street2')} "
                            if addr.get('city') is not None and addr.get('city') != 'N/A': addr_str += f"{addr.get('city')}, "
                            if addr.get('state') is not None and addr.get('state') != 'N/A': addr_str += f"{addr.get('state')} "
                            if addr.get('postalCode') is not None and addr.get('postalCode') != 'N/A': addr_str += f"{addr.get('postalCode')} "
                            if addr.get('country') is not None and addr.get('country') != 'N/A': addr_str += f"{addr.get('country')}"
                            if addr_str.strip():
                                suspect_details += f"Address: {addr_str.strip()}\n"

                phones_data = person.get('phones')
                if phones_data is not None:
                    phones = phones_data.get('phones', [])
                    if phones:
                        for phone in phones:
                            if phone.get('value') is not None and phone.get('value') != 'N/A':
                                suspect_details += f"Phone: {phone.get('value')} "
                                if phone.get('verified') is not None and phone.get('verified') != 'N/A':
                                    suspect_details += f"(Verified: {phone.get('verified')})\n"
                                else:
                                    suspect_details += "\n"

                if person.get('additionalContactInformation') is not None and person.get('additionalContactInformation') != 'N/A':
                    suspect_details += f"Additional Contact Information: {person.get('additionalContactInformation')}\n"

                languages = person.get('languages')
                if languages is not None and languages != 'N/A':
                    suspect_details += f"Languages: {', '.join(languages)}\n"

                races = person.get('races')
                if races is not None and races != 'N/A':
                    suspect_details += f"Races: {', '.join(races)}\n"

                disabilities = person.get('disabilities')
                if disabilities is not None and disabilities != 'N/A':
                    suspect_details += f"Disabilities: {', '.join(disabilities)}\n"

                if person.get('additionalDisabilityInformation') is not None and person.get('additionalDisabilityInformation') != 'N/A':
                    suspect_details += f"Additional Disability Information: {person.get('additionalDisabilityInformation')}\n"

                if "TikTok" in esp_name and person.get('sourceInformation'):
                    source_info = person.get('sourceInformation', {})
                    if 'sourceCaptures' in source_info:
                        for capture in source_info.get('sourceCaptures', []):
                            if capture.get('captureType') == "IP Address":
                                ip_value = capture.get('value', 'N/A')
                                datetime_value = capture.get('dateTime', 'N/A')
                                event_value = capture.get('eventName', 'N/A')
                                if ip_value != 'N/A':
                                    suspect_details += f"IP Address (Login): {ip_value}\n"
                                    suspect_details += f"Login Date/Time: {datetime_value}\n"
                                    suspect_details += f"Event: {event_value}\n"

                if "Imgur" in esp_name:
                    additional_infos = data.get('reportedInformation', {}).get('additionalInformations', [])
                    if additional_infos:
                        for info in additional_infos:
                            value = info.get('value', 'N/A')
                            if value != 'N/A':
                                suspect_details += f"Additional Information from ESP:\n{value}\n"

                if "X Corp" in esp_name:
                    additional_info = person.get('additionalInformations', [{}])[0].get('value', '')
                    source_info = person.get('sourceInformation', {})
                    profile_url = None
                    
                    if source_info and 'sourceCaptures' in source_info:
                        for capture in source_info.get('sourceCaptures', []):
                            if capture.get('captureType') == 'Profile URL':
                                profile_url = capture.get('value')
                                break

                    full_name = location = description = None
                    for line in additional_info.split('\n'):
                        if line.startswith('Full Name:'): full_name = line[len('Full Name:'):].strip()
                        elif line.startswith('Location:'): location = line[len('Location:'):].strip()
                        elif line.startswith('Description:'): description = line[len('Description:'):].strip()

                    if full_name and full_name != 'N/A': suspect_details += f"Full Name: {full_name}\n"
                    if location and location != 'N/A': suspect_details += f"Location: {location}\n"
                    if description and description != 'N/A': suspect_details += f"Description: {description}\n"
                    if profile_url and profile_url != 'N/A': suspect_details += f"Profile URL: {profile_url}\n"

                if "MeetMe" in esp_name:
                    additional_infos = data.get('reportedInformation', {}).get('additionalInformations', [])
                    if additional_infos:
                        for info in additional_infos:
                            value = info.get('value', 'N/A')
                            if value != 'N/A' and "Registration details from Suspect’s MeetMe profile" in value:
                                lines = value.split('\n')
                                meetme_profile = {}
                                for line in lines:
                                    if line.startswith('MeetMe Profile Name:'): meetme_profile['Profile Name'] = line[len('MeetMe Profile Name:'):].strip()
                                    elif line.startswith('MeetMe UserID:'): meetme_profile['UserID'] = line[len('MeetMe UserID:'):].strip()
                                    elif line.startswith('DOB:'): meetme_profile['DOB'] = line[len('DOB:'):].strip()
                                    elif line.startswith('Age:'): meetme_profile['Age'] = line[len('Age:'):].strip()
                                    elif line.startswith('Zip:'): meetme_profile['Zip'] = line[len('Zip:'):].strip()
                                    elif line.startswith('City:'): meetme_profile['City'] = line[len('City:'):].strip()
                                    elif line.startswith('State:'): meetme_profile['State'] = line[len('State:'):].strip()
                                    elif line.startswith('Email:'): meetme_profile['Email'] = line[len('Email:'):].strip()
                                    elif line.startswith('Date Joined meetme.com:'): meetme_profile['Date Joined'] = line[len('Date Joined meetme.com:'):].strip()
                                    elif line.startswith('Registration IP:'): meetme_profile['Registration IP'] = line[len('Registration IP:'):].strip()
                                    elif line.startswith('Phone number used to verify account:'): meetme_profile['Phone'] = line[len('Phone number used to verify account:'):].strip()
                                    elif line.startswith('Recent GPS Data:'): meetme_profile['GPS'] = line[len('Recent GPS Data: Lat./Long.:'):].strip()

                                suspect_details += "Registration details from Suspect’s MeetMe profile (provided by visitor, and is NOT verified):\n"
                                if 'Profile Name' in meetme_profile: suspect_details += f"MeetMe Profile Name: {meetme_profile['Profile Name']}\n"
                                if 'UserID' in meetme_profile: suspect_details += f"MeetMe UserID: {meetme_profile['UserID']}\n"
                                if 'DOB' in meetme_profile: suspect_details += f"Date of Birth: {meetme_profile['DOB']}\n"
                                if 'Age' in meetme_profile: suspect_details += f"Approximate Age: {meetme_profile['Age']}\n"
                                if 'Zip' in meetme_profile or 'City' in meetme_profile or 'State' in meetme_profile:
                                    addr_str = ""
                                    if 'City' in meetme_profile: addr_str += f"{meetme_profile['City']}, "
                                    if 'State' in meetme_profile: addr_str += f"{meetme_profile['State']} "
                                    if 'Zip' in meetme_profile: addr_str += f"{meetme_profile['Zip']}"
                                    suspect_details += f"Address: {addr_str.strip()}\n"
                                if 'Email' in meetme_profile: 
                                    suspect_details += f"Email: {meetme_profile['Email']} "
                                    suspect_details += "(Note: This email was voluntarily provided by the user and may not be verified by MeetMe)\n"
                                if 'Date Joined' in meetme_profile: suspect_details += f"Date Joined MeetMe: {meetme_profile['Date Joined']}\n"
                                if 'Registration IP' in meetme_profile: suspect_details += f"Registration IP: {meetme_profile['Registration IP']}\n"
                                if 'Phone' in meetme_profile and meetme_profile['Phone'] != 'N/A': suspect_details += f"Phone: {meetme_profile['Phone']}\n"
                                if 'GPS' in meetme_profile: 
                                    suspect_details += f"Recent GPS Data: Lat./Long.: {meetme_profile['GPS']}\n"
                                    suspect_details += "Note: This GPS data is used for business purposes and may or may not be indicative of a user’s true geographic location.\n"

                if "MeetMe" in esp_name and person.get('sourceInformation'):
                    source_info = person.get('sourceInformation', {})
                    if 'sourceCaptures' in source_info:
                        for capture in source_info.get('sourceCaptures', []):
                            if capture.get('captureType') == "IP Address":
                                ip_value = capture.get('value', 'N/A')
                                if ip_value != 'N/A':
                                    suspect_details += f"IP Address (Login): {ip_value}\n"

                suspect_details += "\n"

        report += self.get_custom_statements("before_suspect:", data) + suspect_details + self.get_custom_statements("after_suspect:", data) + "\n"

        evidence_details = ""
        files = data.get('reportedInformation', {}).get('uploadedFiles', {}).get('uploadedFiles', [])
        webpage_incidents = data.get('reportedInformation', {}).get('incidentDetails', {}).get('webpageIncident', [])
        if files or ("Reddit" in esp_name and webpage_incidents):
            evidence_details += "EVIDENCE SUMMARY:\n"
            
            if "X Corp" in esp_name and webpage_incidents:
                evidence_details += "Webpage/URL Information:\n"
                for i, webpage in enumerate(webpage_incidents, 1):
                    url = None
                    webpage_source_info = webpage.get('sourceInformation', {})
                    if webpage_source_info and 'sourceCaptures' in webpage_source_info:
                        for capture in webpage_source_info.get('sourceCaptures', []):
                            if capture.get('captureType') == 'URL':
                                url = capture.get('value')
                                break

                    type_info = text = None
                    additional_infos = webpage.get('additionalInformations', [])
                    if additional_infos:
                        info_text = additional_infos[0].get('value', '')
                        for line in info_text.split('\n'):
                            if line.startswith('Type:'): type_info = line[len('Type:'):].strip()
                            elif line.startswith('Text:'): text = line[len('Text:'):].strip()

                    evidence_details += f"  Webpage {i}:\n"
                    if url and url != 'N/A': evidence_details += f"    URL: {url}\n"
                    if type_info and type_info != 'N/A': evidence_details += f"    Type: {type_info}\n"
                    if text and text != 'N/A': evidence_details += f"    Text: {text}\n"
                evidence_details += "\n"

            if "Reddit" in esp_name and webpage_incidents:
                evidence_details += "Reddit Chat Information:\n"
                for i, webpage in enumerate(webpage_incidents, 1):
                    additional_infos = webpage.get('additionalInformations', [])
                    if additional_infos:
                        for info in additional_infos:
                            info_text = info.get('value', 'N/A')
                            if info_text != 'N/A':
                                evidence_details += f"Additional Information: {info_text}\n"
                evidence_details += "\n"

            meta_statement = self.statements.get("meta", self.default_statements["meta"]) if "meta" in self.selected_statements else ""
            file_number = 1
            total_files = len(files)
            for index, file in enumerate(files):
                evidence_details += f"FILE NUMBER {file_number}:\n\n"
                if file.get('filename') is not None and file.get('filename') != 'N/A':
                    evidence_details += f"File Name: {file.get('filename')}\n"
                
                if "Facebook" in esp_name or "Instagram" in esp_name:
                    submittal_id = file.get('submittalId', 'N/A')
                    if submittal_id != 'N/A':
                        evidence_details += f"NCMEC Identifier: {submittal_id}\n"

                verification_hash = file.get('verificationHash', 'N/A')
                if verification_hash != 'N/A':
                    evidence_details += f"MD5 Hash: {verification_hash}\n"

                original_filename = file.get('originalFilename')
                additional_info = file.get('additionalInformations', [])
                msg_id = None
                if additional_info:
                    for info in additional_info:
                        value = info.get('value', '')
                        if value.startswith('Message ID:'):
                            msg_id = value[len('Message ID:'):].strip()
                            break
                
                if msg_id and msg_id in email_data_by_msg_id:
                    email_data = email_data_by_msg_id[msg_id]
                    if email_data.get('sent_date') is not None and email_data.get('sent_date') != 'N/A':
                        evidence_details += f"Sent Date: {email_data.get('sent_date')}\n"
                    if email_data.get('from') is not None and email_data.get('from') != 'N/A':
                        evidence_details += f"Emailed From: {email_data.get('from')}\n"
                    if email_data.get('to') is not None and email_data.get('to') != 'N/A':
                        evidence_details += f"Emailed To: {email_data.get('to')}\n"
                if original_filename is not None and original_filename != 'N/A':
                    evidence_details += f"Original Filename: {original_filename}\n"
                
                if additional_info and esp_name in ["Facebook", "Instagram, Inc."]:
                    raw_description = additional_info[0].get('value', 'N/A')
                    if raw_description and raw_description != 'N/A':
                        lines = raw_description.split('\n')
                        filtered_lines = []
                        for line in lines:
                            line = line.strip()
                            if any(phrase in line for phrase in [
                                "With respect to the section \"Was File Reviewed by Company?",
                                "When Meta responds \"Yes\"",
                                "When Meta responds \"No\"",
                                "File's unique ESP Identifier:",
                                "Messenger Thread ID:",
                                "Uploaded "
                            ]):
                                continue
                            if line and line != "The content can be found in this report.":
                                filtered_lines.append(line)
                        description = '\n'.join(filtered_lines).strip()
                        if description:
                            evidence_details += f"Additional Information from ESP:\n{description}\n"
                elif "MeetMe" in esp_name:
                    additional_infos = data.get('reportedInformation', {}).get('additionalInformations', [])
                    if additional_infos:
                        for info in additional_infos:
                            value = info.get('value', 'N/A')
                            if value != 'N/A':
                                lines = value.split('\n')
                                messages = []
                                current_message = {}
                                previous_subject = None
                                capture_messages = False
                                for line in lines:
                                    if "Complete private message correspondence" in line:
                                        capture_messages = True
                                        continue
                                    if not capture_messages or line.strip().startswith("2024-"):
                                        if capture_messages and line.strip().startswith("2024-"):
                                            break
                                        continue
                                    if line.strip():
                                        if line.startswith('To:'):
                                            if current_message:
                                                messages.append(current_message)
                                            current_message = {'To': line[len('To:'):].strip()}
                                        elif line.startswith('From:'):
                                            current_message['From'] = line[len('From:'):].strip()
                                        elif line.startswith('Sent:'):
                                            current_message['Sent'] = line[len('Sent:'):].strip()
                                        elif line.startswith('Subject:'):
                                            current_message['Subject'] = line[len('Subject:'):].strip()
                                        elif line.startswith('Message:'):
                                            current_message['Message'] = line[len('Message:'):].strip()
                                if current_message:
                                    messages.append(current_message)
                                
                                if messages:
                                    evidence_details += "Private Message Correspondence:\n"
                                    for msg in messages:
                                        evidence_details += f"- To: {msg.get('To', 'N/A')}\n"
                                        evidence_details += f"- From: {msg.get('From', 'N/A')}\n"
                                        evidence_details += f"- Sent: {msg.get('Sent', 'N/A')}\n"
                                        current_subject = msg.get('Subject', 'N/A')
                                        if current_subject != previous_subject and current_subject != 'N/A':
                                            evidence_details += f"- Subject: {current_subject}\n"
                                            previous_subject = current_subject
                                        evidence_details += f"- Message: {msg.get('Message', 'N/A')}\n\n"
                elif additional_info:
                    raw_description = additional_info[0].get('value', 'N/A')
                    if esp_name == "WhatsApp Inc.":
                        paragraphs = raw_description.split('\n\n')
                        filtered_paragraphs = [p.replace('\n', ' ') for p in paragraphs if '{"report_surface":' not in p]
                        description = '\n\n'.join(filtered_paragraphs).strip()
                    else:
                        lines = raw_description.split('\n')
                        filtered_lines = [line for line in lines if "The content can be found in this report" not in line]
                        description = '\n'.join(filtered_lines).strip()
                    if description and description != 'N/A':
                        evidence_details += f"Additional Information from ESP:\n{description}\n"
                
                viewed_by_esp = file.get('viewedByEsp')
                if viewed_by_esp is True:
                    evidence_details += "Viewed by ESP: Yes\n"
                    if esp_name in ["Instagram, Inc.", "Facebook"] and "meta" in self.selected_statements:
                        evidence_details += f"{meta_statement}\n\n"
                elif viewed_by_esp is False:
                    evidence_details += "Viewed by ESP: No\n"
                    if esp_name in ["Instagram, Inc.", "Facebook"] and "meta" in self.selected_statements:
                        evidence_details += f"{meta_statement}\n\n"
                else:
                    evidence_details += "Viewed by ESP: Unknown\n\n"

                source_info = file.get('sourceInformation')
                esp_metadata = file.get('espMetadata')
                if esp_name == "Dropbox, Inc." and esp_metadata is not None and 'metadatas' in esp_metadata:
                    upload_time = 'N/A'
                    for metadata in esp_metadata.get('metadatas', []):
                        if metadata.get('name') == "Upload Date and Time":
                            upload_time = metadata.get('value', 'N/A')
                            break
                    if upload_time != 'N/A':
                        try:
                            upload_datetime = datetime.strptime(upload_time, '%Y-%m-%dT%H:%M:%SZ')
                            formatted_upload_time = upload_datetime.strftime('%m/%d/%Y %H:%M:%S UTC')
                            evidence_details += f"Upload Date/Time: {formatted_upload_time}\n"
                        except ValueError:
                            evidence_details += f"Upload Date/Time: {upload_time}\n"
                elif "Imgur" in esp_name and source_info is not None:
                    upload_time = 'N/A'
                    for capture in source_info.get('sourceCaptures', []):
                        if (capture.get('captureType') == "IP Address" and 
                            capture.get('eventName') == "UPLOAD"):
                            upload_time = capture.get('dateTime', 'N/A')
                            break
                    if upload_time != 'N/A':
                        try:
                            upload_datetime = datetime.strptime(upload_time, '%Y-%m-%dT%H:%M:%SZ')
                            formatted_upload_time = upload_datetime.strftime('%m/%d/%Y %H:%M:%S UTC')
                            evidence_details += f"Upload Date/Time: {formatted_upload_time}\n"
                        except ValueError:
                            evidence_details += f"Upload Date/Time: {upload_time}\n"
                else:
                    upload_time = 'N/A'
                    if source_info is not None:
                        source_captures = source_info.get('sourceCaptures', [])
                        if source_captures:
                            upload_time = source_captures[0].get('dateTime') or 'N/A'
                    if upload_time != 'N/A':
                        try:
                            upload_datetime = datetime.strptime(upload_time, '%Y-%m-%dT%H:%M:%SZ')
                            formatted_upload_time = upload_datetime.strftime('%m/%d/%Y %H:%M:%S UTC')
                            evidence_details += f"Upload Date/Time: {formatted_upload_time}\n"
                        except ValueError:
                            evidence_details += f"Upload Date/Time: {upload_time}\n"

                ncmec_tags = file.get('ncmecTags')
                tags = ', '.join([tag.get('value') for tag in ncmec_tags.get('groups', [{}])[0].get('tags', [])]) if ncmec_tags is not None else 'N/A'
                if tags != 'N/A':
                    evidence_details += f"NCMEC Tags: {tags}\n"
                
                if source_info is not None and 'sourceCaptures' in source_info:
                    captures = source_info.get('sourceCaptures', [])
                    if captures:
                        evidence_details += "Upload IP Address:\n"
                        if "X Corp" in esp_name:
                            evidence_details += "X does not capture IPs for individual Posts, but information from the log of IPs provided by X for the timeframe relevant to the upload date/time will be documented below.\n"
                        elif "Discord" in esp_name:
                            for capture in captures:
                                if capture.get('captureType') == "IP Address":
                                    value = capture.get('value', 'N/A')
                                    if value != 'N/A':
                                        evidence_details += f"{value}\n"
                        else:
                            for capture in captures:
                                if capture.get('captureType') == "IP Address":
                                    value = capture.get('value', 'N/A')
                                    if value != 'N/A':
                                        evidence_details += f"{value}\n"
                
                evidence_details += "\nThis file was viewed by the reporting Investigator\n\n"
                evidence_details += "Investigator's Description:\n"
                evidence_details += "\n"
                
                if index < total_files - 1:
                    evidence_details += "="*50 + "\n\n"
                
                file_number += 1

        report += self.get_custom_statements("before_evidence:", data) + evidence_details + self.get_custom_statements("after_evidence:", data) 

        meetme_ip_section = ""
        if "MeetMe" in esp_name:
            meetme_ip_section += "IP ADDRESS ANALYSIS:\n"
            additional_infos = data.get('reportedInformation', {}).get('additionalInformations', [])
            ip_logs = []
            if additional_infos:
                for info in additional_infos:
                    value = info.get('value', 'N/A')
                    if value != 'N/A':
                        lines = value.split('\n')
                        for line in lines:
                            if line.strip() and line.startswith('2024-'):
                                parts = line.split()
                                if len(parts) >= 3:
                                    timestamp = f"{parts[0]} {parts[1]}"
                                    ip = parts[2]
                                    device = parts[3] if len(parts) > 3 and parts[3].startswith('(') else 'N/A'
                                    ip_logs.append({'timestamp': timestamp, 'ip': ip, 'device': device.strip('()')})
            
            if ip_logs:
                meetme_ip_section += "The following IP addresses and login timestamps were extracted from the MeetMe login history, with geolocation and ownership details:\n\n"
                for log in ip_logs:
                    meetme_ip_section += f"Login Date/Time: {log['timestamp']}\n"
                    meetme_ip_section += f"IP Address: {log['ip']}\n"
                    if log['device'] != 'N/A':
                        meetme_ip_section += f"Device: {log['device']}\n"

                    maxmind_data = self.query_maxmind(log['ip'])
                    meetme_ip_section += "MaxMind Geolocation Data:\n"
                    if "error" in maxmind_data:
                        meetme_ip_section += f"  Error: {maxmind_data['error']}\n"
                    else:
                        try:
                            meetme_ip_section += f"  Country: {maxmind_data.get('country', {}).get('names', {}).get('en', 'N/A')}\n"
                            meetme_ip_section += f"  City: {maxmind_data.get('city', {}).get('names', {}).get('en', 'N/A')}\n"
                        except:
                            meetme_ip_section += "  Error parsing MaxMind data\n"

                    arin_data = self.query_arin(log['ip'])
                    meetme_ip_section += "ARIN WHOIS Data:\n"
                    if "error" in arin_data:
                        meetme_ip_section += f"  Error: {arin_data['error']}\n"
                    else:
                        try:
                            net = arin_data.get('net', {})
                            meetme_ip_section += f"  Organization: {net.get('orgRef', {}).get('@name', 'N/A')}\n"
                        except:
                            meetme_ip_section += "  Error parsing ARIN data\n"
                    
                    meetme_ip_section += "\n"
            else:
                meetme_ip_section += "No IP address login history found in the provided data.\n"
            meetme_ip_section += "\n"

        report += self.get_custom_statements("before_ip:", data) + meetme_ip_section + self.get_custom_statements("after_ip:", data) + self.get_end_custom_statements(data)

        return report

    def generate_ip_report(self) -> str:
        """Generate IP address analysis report with all unique IPs, querying only selected ones."""
        report = "IP ADDRESS ANALYSIS:\n"
        total_unique_ips = len(self.all_ip_data)
        report += f"Total Unique IP Addresses: {total_unique_ips}\n\n"

        queried_ips = set(self.ip_data.keys())
        total_queried = len(queried_ips)
        query_count = 0
        self.status_label.config(text=f"Querying {total_queried} of {total_unique_ips} IPs...")
        self.root.update_idletasks()

        for ip, occurrences in self.all_ip_data.items():
            report += f"IP Address: {ip}\n"
            report += "Occurrences:\n"
            for occ in occurrences:
                datetime_str = occ['datetime'] if occ['datetime'] else 'N/A'
                if datetime_str != 'N/A':
                    try:
                        datetime_obj = datetime.strptime(datetime_str, '%Y-%m-%dT%H:%M:%SZ')
                        formatted_datetime = datetime_obj.strftime('%m/%d/%Y %H:%M:%S UTC')
                        report += f"      - Date/Time: {formatted_datetime}\n"
                    except ValueError:
                        report += f"      - Date/Time: {datetime_str}\n"
                else:
                    report += f"      - Date/Time: N/A\n"

                if 'port' in occ and occ['port'] is not None:
                    report += f"        Port: {occ['port']}\n"
                report += f"        IP Event: {occ['event'] if occ['event'] else 'N/A'}\n"

            if ip in queried_ips:
                query_count += 1
                self.status_label.config(text=f"Queried {query_count} of {total_queried} IPs...")
                self.root.update_idletasks()
                report += "\nMaxMind Geolocation Data:\n"
                maxmind_data = self.query_maxmind(ip)
                if "error" in maxmind_data:
                    report += f"        Error: {maxmind_data['error']}\n"
                else:
                    try:
                        report += f"        Country: {maxmind_data.get('country', {}).get('names', {}).get('en', 'N/A')}\n"
                        report += f"        City: {maxmind_data.get('city', {}).get('names', {}).get('en', 'N/A')}\n"
                    except:
                        report += "        Error parsing MaxMind data\n"

                report += "\nARIN WHOIS Data:\n"
                arin_data = self.query_arin(ip)
                if "error" in arin_data:
                    report += f"        Error: {arin_data['error']}\n"
                else:
                    try:
                        net = arin_data.get('net', {})
                        report += f"        Organization: {net.get('orgRef', {}).get('@name', 'N/A')}\n"
                    except:
                        report += "        Error parsing ARIN data\n"
            else:
                report += "\nMaxMind Geolocation Data:\n        Not queried (IP limit applied)\n"
                report += "\nARIN WHOIS Data:\n        Not queried (IP limit applied)\n"

            report += "\n" + "="*50 + "\n\n"

        self.status_label.config(text=f"Finished querying {query_count} of {total_queried} IPs")
        self.root.update_idletasks()
        return report

    def export_ip_data_to_excel(self):
        """Export IP address analysis data to an Excel file."""
        if not self.all_ip_data:
            messagebox.showwarning("Warning", "No IP data available to export.")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            initialfile=f"cybertip_ip_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        if not filename:
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "IP Address Analysis"

        headers = ["IP Address", "Date/Time", "Port", "IP Event", "MaxMind Country", "MaxMind City", "ARIN Organization"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        row = 2
        for ip, occurrences in self.all_ip_data.items():
            for occ in occurrences:
                ws.cell(row=row, column=1).value = ip
                datetime_str = occ['datetime'] if occ['datetime'] else 'N/A'
                if datetime_str != 'N/A':
                    try:
                        datetime_obj = datetime.strptime(datetime_str, '%Y-%m-%dT%H:%M:%SZ')
                        formatted_datetime = datetime_obj.strftime('%m/%d/%Y %H:%M:%S UTC')
                        ws.cell(row=row, column=2).value = formatted_datetime
                    except ValueError:
                        ws.cell(row=row, column=2).value = datetime_str
                else:
                    ws.cell(row=row, column=2).value = 'N/A'
                ws.cell(row=row, column=3).value = occ['port'] if occ['port'] else 'N/A'
                ws.cell(row=row, column=4).value = occ['event'] if occ['event'] else 'N/A'

                if ip in self.ip_data:
                    maxmind_data = self.query_maxmind(ip)
                    if "error" in maxmind_data:
                        ws.cell(row=row, column=5).value = f"Error: {maxmind_data['error']}"
                        ws.cell(row=row, column=6).value = "N/A"
                    else:
                        try:
                            ws.cell(row=row, column=5).value = maxmind_data.get('country', {}).get('names', {}).get('en', 'N/A')
                            ws.cell(row=row, column=6).value = maxmind_data.get('city', {}).get('names', {}).get('en', 'N/A')
                        except:
                            ws.cell(row=row, column=5).value = "Error parsing MaxMind data"
                            ws.cell(row=row, column=6).value = "N/A"

                    arin_data = self.query_arin(ip)
                    if "error" in arin_data:
                        ws.cell(row=row, column=7).value = f"Error: {arin_data['error']}"
                    else:
                        try:
                            net = arin_data.get('net', {})
                            ws.cell(row=row, column=7).value = net.get('orgRef', {}).get('@name', 'N/A')
                        except:
                            ws.cell(row=row, column=7).value = "Error parsing ARIN data"
                else:
                    ws.cell(row=row, column=5).value = "Not queried (IP limit applied)"
                    ws.cell(row=row, column=6).value = "Not queried (IP limit applied)"
                    ws.cell(row=row, column=7).value = "Not queried (IP limit applied)"
                row += 1

        # Adjust column widths
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = max_length + 2
            ws.column_dimensions[column].width = adjusted_width

        try:
            wb.save(filename)
            messagebox.showinfo("Success", f"IP data exported as {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export IP data: {e}")

    def export_evidence_to_excel(self):
        """Export evidence summary data to an Excel file."""
        if not self.json_file_path:
            messagebox.showwarning("Warning", "No JSON file loaded to export evidence data.")
            return

        data = self.load_json()
        if data is None:
            messagebox.showwarning("Warning", "Failed to load JSON for evidence export.")
            return

        files = data.get('reportedInformation', {}).get('uploadedFiles', {}).get('uploadedFiles', [])
        webpage_incidents = data.get('reportedInformation', {}).get('incidentDetails', {}).get('webpageIncident', [])
        esp_name = data.get('reportedInformation', {}).get('reportingEsp', {}).get('espName', 'N/A')

        if not files and not ("Reddit" in esp_name and webpage_incidents):
            messagebox.showwarning("Warning", "No evidence data available to export.")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            initialfile=f"cybertip_evidence_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        if not filename:
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Evidence Summary"

        headers = [
            "File Number", "File Name", "NCMEC Identifier", "MD5 Hash", "Sent Date",
            "Emailed From", "Emailed To", "Original Filename", "Additional Information",
            "Viewed by ESP", "Upload Date/Time", "NCMEC Tags", "Upload IP Address",
            "Webpage Number", "URL", "Type", "Text"
        ]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        row = 2
        email_data_by_msg_id = {}
        email_incidents = data.get('reportedInformation', {}).get('incidentDetails', {}).get('emailIncident', [])
        if email_incidents:
            for email in email_incidents:
                contents = email.get('contents', [{}])[0].get('value', '')
                sent_date, from_email, to_email, msg_id = None, None, None, None
                for line in contents.split('\n'):
                    if line.startswith('Sent Date:'): sent_date = line[len('Sent Date:'):].strip()
                    elif line.startswith('From:'): from_email = line[len('From:'):].strip()
                    elif line.startswith('To:'): to_email = line[len('To:'):].strip()
                    elif line.startswith('X-Ymail-Msg-Id:'): msg_id = line[len('X-Ymail-Msg-Id:'):].strip()
                if msg_id:
                    email_data_by_msg_id[msg_id] = {'sent_date': sent_date, 'from': from_email, 'to': to_email}

        if "X Corp" in esp_name and webpage_incidents:
            for i, webpage in enumerate(webpage_incidents, 1):
                ws.cell(row=row, column=14).value = f"Webpage {i}"
                url = None
                webpage_source_info = webpage.get('sourceInformation', {})
                if webpage_source_info and 'sourceCaptures' in webpage_source_info:
                    for capture in webpage_source_info.get('sourceCaptures', []):
                        if capture.get('captureType') == 'URL':
                            url = capture.get('value')
                            break
                ws.cell(row=row, column=15).value = url if url and url != 'N/A' else 'N/A'

                type_info = text = None
                additional_infos = webpage.get('additionalInformations', [])
                if additional_infos:
                    info_text = additional_infos[0].get('value', '')
                    for line in info_text.split('\n'):
                        if line.startswith('Type:'): type_info = line[len('Type:'):].strip()
                        elif line.startswith('Text:'): text = line[len('Text:'):].strip()
                ws.cell(row=row, column=16).value = type_info if type_info and type_info != 'N/A' else 'N/A'
                ws.cell(row=row, column=17).value = text if text and text != 'N/A' else 'N/A'
                row += 1

        if "Reddit" in esp_name and webpage_incidents:
            for i, webpage in enumerate(webpage_incidents, 1):
                ws.cell(row=row, column=14).value = f"Webpage {i}"
                additional_infos = webpage.get('additionalInformations', [])
                if additional_infos:
                    info_text = additional_infos[0].get('value', 'N/A')
                    ws.cell(row=row, column=9).value = info_text if info_text != 'N/A' else 'N/A'
                row += 1

        file_number = 1
        for file in files:
            ws.cell(row=row, column=1).value = file_number
            ws.cell(row=row, column=2).value = file.get('filename', 'N/A')

            if "Facebook" in esp_name or "Instagram" in esp_name:
                ws.cell(row=row, column=3).value = file.get('submittalId', 'N/A')
            else:
                ws.cell(row=row, column=3).value = 'N/A'

            ws.cell(row=row, column=4).value = file.get('verificationHash', 'N/A')

            additional_info = file.get('additionalInformations', [])
            msg_id = None
            if additional_info:
                for info in additional_info:
                    value = info.get('value', '')
                    if value.startswith('Message ID:'):
                        msg_id = value[len('Message ID:'):].strip()
                        break

            if msg_id and msg_id in email_data_by_msg_id:
                email_data = email_data_by_msg_id[msg_id]
                ws.cell(row=row, column=5).value = email_data.get('sent_date', 'N/A')
                ws.cell(row=row, column=6).value = email_data.get('from', 'N/A')
                ws.cell(row=row, column=7).value = email_data.get('to', 'N/A')
            else:
                ws.cell(row=row, column=5).value = 'N/A'
                ws.cell(row=row, column=6).value = 'N/A'
                ws.cell(row=row, column=7).value = 'N/A'

            ws.cell(row=row, column=8).value = file.get('originalFilename', 'N/A')

            if additional_info and esp_name in ["Facebook", "Instagram, Inc."]:
                raw_description = additional_info[0].get('value', 'N/A')
                if raw_description and raw_description != 'N/A':
                    lines = raw_description.split('\n')
                    filtered_lines = []
                    for line in lines:
                        line = line.strip()
                        if any(phrase in line for phrase in [
                            "With respect to the section \"Was File Reviewed by Company?",
                            "When Meta responds \"Yes\"",
                            "When Meta responds \"No\"",
                            "File's unique ESP Identifier:",
                            "Messenger Thread ID:",
                            "Uploaded "
                        ]):
                            continue
                        if line and line != "The content can be found in this report.":
                            filtered_lines.append(line)
                    description = '\n'.join(filtered_lines).strip()
                    ws.cell(row=row, column=9).value = description if description else 'N/A'
            elif additional_info:
                raw_description = additional_info[0].get('value', 'N/A')
                if esp_name == "WhatsApp Inc.":
                    paragraphs = raw_description.split('\n\n')
                    filtered_paragraphs = [p.replace('\n', ' ') for p in paragraphs if '{"report_surface":' not in p]
                    description = '\n\n'.join(filtered_paragraphs).strip()
                else:
                    lines = raw_description.split('\n')
                    filtered_lines = [line for line in lines if "The content can be found in this report" not in line]
                    description = '\n'.join(filtered_lines).strip()
                ws.cell(row=row, column=9).value = description if description and description != 'N/A' else 'N/A'
            else:
                ws.cell(row=row, column=9).value = 'N/A'

            viewed_by_esp = file.get('viewedByEsp')
            ws.cell(row=row, column=10).value = "Yes" if viewed_by_esp is True else "No" if viewed_by_esp is False else "Unknown"

            source_info = file.get('sourceInformation')
            esp_metadata = file.get('espMetadata')
            upload_time = 'N/A'
            if esp_name == "Dropbox, Inc." and esp_metadata is not None and 'metadatas' in esp_metadata:
                for metadata in esp_metadata.get('metadatas', []):
                    if metadata.get('name') == "Upload Date and Time":
                        upload_time = metadata.get('value', 'N/A')
                        break
            elif "Imgur" in esp_name and source_info is not None:
                for capture in source_info.get('sourceCaptures', []):
                    if (capture.get('captureType') == "IP Address" and 
                        capture.get('eventName') == "UPLOAD"):
                        upload_time = capture.get('dateTime', 'N/A')
                        break
            elif source_info is not None:
                source_captures = source_info.get('sourceCaptures', [])
                if source_captures:
                    upload_time = source_captures[0].get('dateTime') or 'N/A'

            if upload_time != 'N/A':
                try:
                    upload_datetime = datetime.strptime(upload_time, '%Y-%m-%dT%H:%M:%SZ')
                    formatted_upload_time = upload_datetime.strftime('%m/%d/%Y %H:%M:%S UTC')
                    ws.cell(row=row, column=11).value = formatted_upload_time
                except ValueError:
                    ws.cell(row=row, column=11).value = upload_time
            else:
                ws.cell(row=row, column=11).value = 'N/A'

            ncmec_tags = file.get('ncmecTags')
            tags = ', '.join([tag.get('value') for tag in ncmec_tags.get('groups', [{}])[0].get('tags', [])]) if ncmec_tags is not None else 'N/A'
            ws.cell(row=row, column=12).value = tags

            ip_addresses = []
            if source_info is not None and 'sourceCaptures' in source_info:
                captures = source_info.get('sourceCaptures', [])
                if "X Corp" in esp_name:
                    ip_addresses.append("X does not capture IPs for individual Posts")
                elif "Discord" in esp_name:
                    for capture in captures:
                        if capture.get('captureType') == "IP Address":
                            value = capture.get('value', 'N/A')
                            if value != 'N/A':
                                ip_addresses.append(value)
                else:
                    for capture in captures:
                        if capture.get('captureType') == "IP Address":
                            value = capture.get('value', 'N/A')
                            if value != 'N/A':
                                ip_addresses.append(value)
            ws.cell(row=row, column=13).value = '; '.join(ip_addresses) if ip_addresses else 'N/A'

            row += 1
            file_number += 1

        # Adjust column widths
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width

        try:
            wb.save(filename)
            messagebox.showinfo("Success", f"Evidence summary exported as {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export evidence summary: {e}")

    def save_report(self, police_report: str, ip_report: str) -> str:
        full_report = police_report + ip_report
        
        default_filename = f"cybertip_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=default_filename,
            title="Save CyberTipline Report As"
        )
        
        if not filename:
            raise ValueError("Save operation cancelled by user")

        doc = Document()
        
        headers = [
            "INCIDENT SUMMARY:",
            "SUSPECT INFORMATION:",
            "EVIDENCE SUMMARY:",
            "IP ADDRESS ANALYSIS:"
        ]
        labels = [
            "Report Date:",
            "Report ID:",
            "Date Received:",
            "Incident Type:",
            "Incident Date/Time:",
            "Reported By:",
            "Incident Date/Time Description:",
            "Chat Service or Client:",
            "Chat Room Name:",
            "Screen Name:",
            "User ID:",
            "Email:",
            "Service:",
            "File Number",
            "File Name:",
            "Additional Information from ESP:",
            "Viewed by ESP:",
            "Upload Date/Time:",
            "NCMEC Tags:",
            "Total Unique IP Addresses:",
            "IP Address:",
            "Upload Information:",
            "Investigator's Description:",
            "First Name:",
            "Middle Name:",
            "Last Name:",
            "Date of Birth:",
            "Approximate Age:",
            "Address:",
            "Phone:",
            "Verified Email:",
            "ESP Service:",
            "ESP User ID:",
            "Email Verification Date:",
            "Gender:",
            "Upload IP Address:",
            "Emailed From:",
            "Emailed To:",
            "Sent Date:",
            "Original Filename:",
            "Full Name:",
            "Location:",
            "Description:",
            "Profile URL:",
            "Webpage/URL Information:",
            "Private Message Correspondence:",
            "IP Address (Login):",
            "MeetMe Profile Name:",
            "MeetMe UserID:",
            "Date Joined MeetMe:",
            "Registration IP:",
            "Recent GPS Data:",
            "IP Log for MeetMe UserID:",
            "Additional Information:",
            "Login Date/Time:",
            "Event:",
            "NCMEC Identifier:",
            "MD5 Hash:"
        ]
        custom_headers = []
        for key in self.selected_statements:
            if key not in self.default_statements:
                prefix_map = {
                    "at_beginning:": "At Beginning of Report",
                    "before_incident:": "Before Incident Summary",
                    "after_incident:": "After Incident Summary",
                    "before_suspect:": "Before Suspect Information",
                    "after_suspect:": "After Suspect Information",
                    "before_evidence:": "Before Evidence Summary",
                    "after_evidence:": "After Evidence Summary",
                    "before_ip:": "Before IP Address Analysis",
                    "after_ip:": "After IP Address Analysis",
                    "": "Custom Statements"
                }
                for prefix, section in prefix_map.items():
                    if key.startswith(prefix):
                        sub_key = key[len(prefix):].strip()
                        if sub_key:
                            custom_headers.append(f"{sub_key.upper()}:")
                        break
        
        ip_intro_statement = self.statements.get("ip_intro", self.default_statements["ip_intro"]) if "ip_intro" in self.selected_statements else ""
        
        sections = full_report.split('\n\n')
        for section in sections:
            if not section.strip():
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = 0
            p.paragraph_format.line_spacing = 1.0
            lines = section.split('\n')
            
            for i, line in enumerate(lines):
                stripped_line = line.strip()
                if not stripped_line:
                    continue
                
                if stripped_line in headers or stripped_line in custom_headers:
                    run = p.add_run(stripped_line + '\n')
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.name = 'Times New Roman'
                    run_element = run._r
                    rPr = run_element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), 'D3D3D3')
                    rPr.append(shd)
                    if stripped_line == "IP ADDRESS ANALYSIS:" and ip_intro_statement:
                        run = p.add_run('\n')
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        run = p.add_run(ip_intro_statement + '\n\n')
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                else:
                    content = line.lstrip()
                    leading_spaces = line[:len(line) - len(content)]
                    
                    if content == self.statements.get("meta", self.default_statements["meta"]) and "meta" in self.selected_statements:
                        p_meta = doc.add_paragraph()
                        p_meta.paragraph_format.left_indent = Inches(0.5)
                        p_meta.paragraph_format.space_after = 0
                        p_meta.paragraph_format.line_spacing = 1.0
                        run = p_meta.add_run(content)
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                    elif any(content.startswith(label) for label in labels):
                        for label in labels:
                            if content.startswith(label):
                                value = content[len(label):].strip()
                                run = p.add_run(leading_spaces)
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                run = p.add_run(label)
                                run.bold = True
                                run.underline = True
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                if label == "Investigator's Description:":
                                    run_element = run._r
                                    rPr = run_element.get_or_add_rPr()
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:fill'), 'FF0000')
                                    rPr.append(shd)
                                run = p.add_run("  ")
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                run = p.add_run(value)
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                break
                    else:
                        run = p.add_run(line)
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        if content == "This file was viewed by the reporting Investigator":
                            run.italic = True
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                
                if i < len(lines) - 1 and lines[i + 1].strip():
                    run = p.add_run('\n')
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
            
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = 12
        
        doc.save(filename)
        return filename

    def analyze_report(self):
        if not self.json_file_path:
            messagebox.showerror("Error", "Please select a JSON file")
            return

        self.status_label.config(text="Starting analysis...")
        self.root.update_idletasks()

        def run_analysis(q):
            try:
                data = self.load_json()
                if data is None:
                    q.put(("error", "Failed to load JSON"))
                    return

                self.extract_ip_addresses(data)
                police_report = self.generate_police_report(data)
                ip_report = self.generate_ip_report()

                q.put(("success", police_report, ip_report))
            except Exception as e:
                q.put(("error", str(e)))

        q = queue.Queue()
        thread = threading.Thread(target=run_analysis, args=(q,))
        thread.start()

        def check_thread():
            if thread.is_alive():
                self.root.after(100, check_thread)
            else:
                result = q.get()
                if result[0] == "error":
                    messagebox.showerror("Error", result[1])
                    self.status_label.config(text="Analysis failed")
                else:
                    police_report, ip_report = result[1], result[2]
                    self.output_text.delete(1.0, tk.END)
                    self.output_text.insert(tk.END, police_report + ip_report)
                    try:
                        filename = self.save_report(police_report, ip_report)
                        messagebox.showinfo("Success", f"Report saved as {filename}")
                        self.status_label.config(text="Analysis complete")
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to save: {e}")
                        self.status_label.config(text="Save failed")

        check_thread()
        
    def check_for_updates(self):
        """Check GitHub for updates."""
        current_version = "2.0"
        repo_api_url = "https://api.github.com/repos/koebbe14/CyberTip-Analyzer-Tool-Report-Generator/releases/latest"
        repo_download_url = "https://github.com/koebbe14/CyberTip-Analyzer-Tool-Report-Generator/releases/latest"
        
        def version_tuple(v):
            return tuple(map(int, v.split(".")))

        try:
            response = requests.get(repo_api_url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                latest_version = data.get("tag_name", "0.0").lstrip("v")
                if version_tuple(latest_version) > version_tuple(current_version):
                    if messagebox.askyesno("Update Available", f"A new version {latest_version} is available.\n\nWould you like to open the download page?"):
                        webbrowser.open(repo_download_url)
                else:
                    messagebox.showinfo("Up to Date", "You are running the latest version.")
            else:
                messagebox.showerror("Error", f"Failed to check for updates: HTTP {response.status_code}")
        except Exception as e:
            messagebox.showerror("Error", f"Error checking for updates: {e}")

def main():
    root = tk.Tk()
    app = CyberTipAnalyzer(root)
    root.mainloop()

if __name__ == "__main__":
    main()