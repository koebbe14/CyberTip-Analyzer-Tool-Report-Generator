"""DOCX report formatting and styling with per-statement formatting support."""

from __future__ import annotations

import re
from typing import Dict, Optional, Set

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from catrg.models.statements import StatementManager, DEFAULT_STATEMENTS, get_formatting
from catrg.core.report_generator import ReportTemplate, DEFAULT_SECTION_NAMES
from catrg.utils.logger import get_logger

log = get_logger(__name__)

SECTION_HEADERS = [
    "INCIDENT SUMMARY:",
    "SUSPECT INFORMATION:",
    "EVIDENCE SUMMARY:",
    "IP ADDRESS ANALYSIS:",
    "SUMMARY STATISTICS:",
]

FIELD_LABELS = [
    "Report Date:", "Report ID:", "Date Received:",
    "Incident Type:", "Incident Date/Time:", "Reported By:",
    "Incident Date/Time Description:", "Chat Service or Client:",
    "Chat Room Name:", "Screen Name:", "User ID:", "Email:",
    "Service:", "File Number", "File Name:",
    "Additional Information from ESP:", "Viewed by ESP:",
    "Upload Date/Time:", "NCMEC Tags:", "Total Unique IP Addresses:",
    "IP Address:", "Upload Information:", "Investigator's Description:",
    "First Name:", "Middle Name:", "Last Name:", "Date of Birth:",
    "Approximate Age:", "Address:", "Phone:", "Verified Email:",
    "ESP Service:", "ESP User ID:", "Email Verification Date:",
    "Gender:", "Upload IP Address:", "Emailed From:", "Emailed To:",
    "Sent Date:", "Original Filename:", "Full Name:", "Location:",
    "Description:", "Profile URL:", "Webpage/URL Information:",
    "Private Message Correspondence:", "IP Address (Login):",
    "MeetMe Profile Name:", "MeetMe UserID:", "Date Joined MeetMe:",
    "Registration IP:", "Recent GPS Data:", "IP Log for MeetMe UserID:",
    "Additional Information:", "Login Date/Time:", "Event:",
    "NCMEC Identifier:", "MD5 Hash:", "Preferred Name:", "Preferred Pronouns:",
    "Physical Description:", "Vehicle Description:", "Vehicle Tag Number:",
    "Occupation:", "Relationship to Reporter:", "Relationship to Child Victim:",
    "Access to Child Victim:", "Access to Children:", "Access to Firearms:",
    "Convicted Sex Offender:", "Aware of Report:", "Gang Affiliation:",
    "Additional Contact Information:", "Languages:", "Races:",
    "Disabilities:", "Additional Disability Information:",
    "CyberTip ID:", "ESP:", "Total Suspect Records:", "Total Evidence Files:",
    "Total Unique IPs:", "Countries Identified:", "ISPs/Orgs Identified:",
]


def _build_all_headers(stmts: StatementManager, tmpl: Optional[ReportTemplate] = None) -> list[str]:
    """Build the full set of section header strings for DOCX bold styling."""
    headers = list(SECTION_HEADERS)

    if tmpl:
        for section_id, custom_name in tmpl.section_names.items():
            header_str = f"{custom_name}:"
            if header_str not in headers:
                headers.append(header_str)
        for cs in tmpl.custom_sections:
            cs_header = f"{cs.get('name', '').upper()}:"
            if cs_header and cs_header not in headers:
                headers.append(cs_header)

    for key in stmts.selected:
        if key not in DEFAULT_STATEMENTS:
            for prefix in ("at_beginning:", "after_intro:", "before_incident:", "after_incident:",
                           "before_suspect:", "after_suspect:", "before_evidence:",
                           "after_evidence:", "before_ip:", "after_ip:",
                           "after_all_sections:", ""):
                if key.startswith(prefix):
                    sub = key[len(prefix):].strip()
                    if sub:
                        headers.append(f"{sub.upper()}:")
                    break
    return headers


HIGHLIGHT_MAP = {
    "yellow": "yellow",
    "green": "green",
    "cyan": "cyan",
    "magenta": "magenta",
    "red": "red",
    "blue": "blue",
    "darkYellow": "darkYellow",
}


def _build_stmt_format_map(stmts: StatementManager) -> Dict[str, dict]:
    """Build a mapping of statement text fragments to their formatting metadata.

    Only includes custom statements that have non-default formatting.
    """
    fmt_map: Dict[str, dict] = {}
    for key, val in stmts.statements.items():
        if key in DEFAULT_STATEMENTS:
            continue
        if not isinstance(val, dict):
            continue
        fmt = val.get("formatting")
        if not fmt:
            continue
        is_default = (
            fmt.get("font_size", 12) == 12
            and not fmt.get("bold", False)
            and not fmt.get("italic", False)
            and fmt.get("indent", 0.0) == 0.0
            and not fmt.get("highlight", "")
        )
        if is_default:
            continue
        text = val.get("text", "").strip()
        if text:
            first_line = text.split("\n")[0][:80]
            fmt_map[first_line] = fmt
    return fmt_map


def save_docx(
    full_report: str,
    filename: str,
    stmts: StatementManager,
    ip_intro_statement: str = "",
    meta_statement: str = "",
    template: Optional[ReportTemplate] = None,
) -> str:
    """Write *full_report* to a styled DOCX file and return the filename."""
    doc = Document()
    all_headers = _build_all_headers(stmts, template)
    stmt_fmt_map = _build_stmt_format_map(stmts)

    sections = full_report.split("\n\n")
    for section in sections:
        if not section.strip():
            continue

        stmt_fmt = _match_stmt_format(section.strip(), stmt_fmt_map)

        if stmt_fmt and stmt_fmt.get("indent", 0.0) > 0:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(stmt_fmt["indent"])
        else:
            p = doc.add_paragraph()
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0
        lines = section.split("\n")

        is_cybertip_separator = bool(re.match(r"^={10,}", lines[0].strip())) if lines else False
        if is_cybertip_separator:
            doc.add_page_break()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = 0
            p.paragraph_format.line_spacing = 1.0

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue

            is_cybertip_header = bool(re.match(r"^CYBERTIP #\d+", stripped))

            if is_cybertip_header:
                run = p.add_run(stripped + "\n")
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = "Times New Roman"
                _apply_shading(run, "B0C4DE")

            elif stripped in all_headers:
                run = p.add_run(stripped + "\n")
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = "Times New Roman"
                _apply_shading(run, "D3D3D3")

                if "IP ADDRESS ANALYSIS" in stripped and ip_intro_statement:
                    run = p.add_run("\n")
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"
                    run = p.add_run(ip_intro_statement + "\n\n")
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"

            elif stripped == meta_statement.strip() and meta_statement:
                p_meta = doc.add_paragraph()
                p_meta.paragraph_format.left_indent = Inches(0.5)
                p_meta.paragraph_format.space_after = 0
                p_meta.paragraph_format.line_spacing = 1.0
                run = p_meta.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

            elif _starts_with_label(stripped):
                _format_label_line(p, line, stripped)

            elif stmt_fmt:
                run = p.add_run(line)
                run.font.size = Pt(stmt_fmt.get("font_size", 12))
                run.font.name = "Times New Roman"
                if stmt_fmt.get("bold"):
                    run.bold = True
                if stmt_fmt.get("italic"):
                    run.italic = True
                hl = stmt_fmt.get("highlight", "")
                if hl and hl in HIGHLIGHT_MAP:
                    run.font.highlight_color = _docx_highlight(hl)

            else:
                run = p.add_run(line)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                if stripped == "This file was viewed by the reporting Investigator":
                    run.italic = True
                    run.font.size = Pt(10)

            if i < len(lines) - 1 and lines[i + 1].strip():
                run = p.add_run("\n")
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = 12

    doc.save(filename)
    return filename


def _match_stmt_format(text: str, fmt_map: Dict[str, dict]) -> Optional[dict]:
    """Check if *text* starts with any known custom statement first-line."""
    for first_line, fmt in fmt_map.items():
        if first_line in text:
            return fmt
    return None


def _docx_highlight(color_name: str):
    """Map a highlight color name to a python-docx highlight color enum value."""
    from docx.enum.text import WD_COLOR_INDEX
    mapping = {
        "yellow": WD_COLOR_INDEX.YELLOW,
        "green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "cyan": WD_COLOR_INDEX.TURQUOISE,
        "magenta": WD_COLOR_INDEX.PINK,
        "red": WD_COLOR_INDEX.RED,
        "blue": WD_COLOR_INDEX.BLUE,
    }
    return mapping.get(color_name)


def _starts_with_label(text: str) -> bool:
    return any(text.startswith(label) for label in FIELD_LABELS)


def _format_label_line(p, line: str, content: str) -> None:
    leading = line[: len(line) - len(content)]

    for label in FIELD_LABELS:
        if content.startswith(label):
            value = content[len(label):].strip()

            run = p.add_run(leading)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

            run = p.add_run(label)
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

            if label == "Investigator's Description:":
                _apply_shading(run, "FF0000")

            run = p.add_run("  ")
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

            run = p.add_run(value)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            return


def _apply_shading(run, color_hex: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    rPr.append(shd)
