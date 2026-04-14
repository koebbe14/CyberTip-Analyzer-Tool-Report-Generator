"""All dialog windows for CAT-RG.

Credentials, investigator info, statement customisation,
statement chooser, report-template editor, and report comparison.
"""

from __future__ import annotations

import os
import re
import webbrowser
from collections import defaultdict
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, simpledialog, ttk
from typing import TYPE_CHECKING, Callable, Dict, FrozenSet, List, Optional, Set, Tuple

from catrg.models.config import ConfigManager
from catrg.models.statements import (
    StatementManager,
    DEFAULT_STATEMENTS,
    DEFAULT_FORMATTING,
    PLACEMENT_PREFIXES,
    PREFIX_TO_PLACEMENT,
    get_all_placement_prefixes,
)
from catrg.core.parser import load_json, extract_comparison_data
from catrg.core.report_generator import (
    ReportTemplate,
    DEFAULT_SECTIONS,
    DEFAULT_SECTION_NAMES,
)
from catrg.utils.logger import get_logger

if TYPE_CHECKING:
    from catrg.gui.main_window import CyberTipAnalyzer

log = get_logger(__name__)

ESP_VALUES = [
    "Custom", "Discord", "Dropbox", "Facebook", "Google", "Imgur",
    "Instagram", "Kik", "MeetMe", "Microsoft", "Reddit", "Roblox",
    "Snapchat", "Sony", "Synchronoss", "TikTok", "WhatsApp",
    "X (Twitter)", "Yahoo",
]

_ESP_LIST = [
    "Discord", "Dropbox", "Facebook", "Google", "Imgur",
    "Instagram", "Kik", "MeetMe", "Microsoft", "Reddit", "Roblox",
    "Snapchat", "Sony", "Synchronoss", "TikTok", "WhatsApp",
    "X (Twitter)", "Yahoo",
]


# ── Investigator Info ─────────────────────────────────────────────

def show_investigator_dialog(root: tk.Tk, config: ConfigManager) -> None:
    dialog = tk.Toplevel(root)
    dialog.title("Investigator Information")
    dialog.geometry("300x200")
    dialog.transient(root)
    dialog.grab_set()

    tk.Label(dialog, text="Enter Your Name:").pack(pady=5)
    name_entry = tk.Entry(dialog, width=30)
    name_entry.pack(pady=5)
    name_entry.insert(0, config.investigator.name)

    tk.Label(dialog, text="Enter Your Title:").pack(pady=5)
    title_entry = tk.Entry(dialog, width=30)
    title_entry.pack(pady=5)
    title_entry.insert(0, config.investigator.title)

    def save():
        n = name_entry.get().strip()
        t = title_entry.get().strip()
        if not n or not t:
            messagebox.showwarning("Warning", "Please enter both name and title.")
            return
        config.investigator.name = n
        config.investigator.title = t
        config.save_investigator()
        dialog.destroy()

    tk.Button(dialog, text="Save", command=save).pack(pady=10)
    dialog.protocol("WM_DELETE_WINDOW", lambda: root.quit())
    root.wait_window(dialog)


# ── MaxMind Credentials ──────────────────────────────────────────

def show_maxmind_dialog(root: tk.Tk, config: ConfigManager) -> Optional[tk.Toplevel]:
    dialog = tk.Toplevel(root)
    dialog.title("MaxMind Credentials")
    dialog.geometry("400x350")
    dialog.transient(root)
    dialog.grab_set()

    tk.Label(dialog, text="For IP Geo Lookup, enter your MaxMind credentials below (optional).\n\nCreate new account at ").pack(pady=5)
    link = tk.Label(dialog, text="this link", fg="blue", cursor="hand2")
    link.pack(pady=5)
    link.bind("<Button-1>", lambda e: webbrowser.open_new(
        "https://www.maxmind.com/en/geolite2/signup?utm_source=kb&utm_medium=kb-link&utm_campaign=kb-create-account"
    ))

    tk.Label(dialog, text="MaxMind Account ID:").pack(pady=5)
    id_entry = tk.Entry(dialog, width=30)
    id_entry.pack(pady=5)
    id_entry.insert(0, config.maxmind.account_id)

    tk.Label(dialog, text="MaxMind License Key:").pack(pady=5)
    key_entry = tk.Entry(dialog, width=30, show="*")
    key_entry.pack(pady=5)
    key_entry.insert(0, config.maxmind.license_key)

    def save_and_close():
        config.maxmind.account_id = id_entry.get().strip()
        config.maxmind.license_key = key_entry.get().strip()
        config.save_maxmind()
        dialog.destroy()
        if config.maxmind.is_configured:
            messagebox.showinfo("Success", "MaxMind credentials updated successfully.")
        else:
            messagebox.showinfo("Info", "MaxMind credentials skipped -- geolocation will not be available.")

    tk.Button(dialog, text="Save", command=save_and_close).pack(pady=10)
    tk.Button(dialog, text="Skip (No Geolocation)", command=dialog.destroy).pack(pady=5)
    return dialog


# ── ARIN Credentials ─────────────────────────────────────────────

def show_arin_dialog(root: tk.Tk, config: ConfigManager) -> Optional[tk.Toplevel]:
    dialog = tk.Toplevel(root)
    dialog.title("ARIN API Key")
    dialog.geometry("400x300")
    dialog.transient(root)
    dialog.grab_set()

    tk.Label(
        dialog,
        text="Enter your ARIN API key (optional).\n\n"
             "Adding an ARIN API key increases daily query limits.\n\n"
             "Daily query limits WITHOUT ARIN API key - 15 per min and 256 per day.\n"
             "Daily query limits WITH ARIN API Key - 60 per min and 1,000 per day.\n\n"
             "Get an API key at:",
    ).pack(pady=5)
    link = tk.Label(dialog, text="this link", fg="blue", cursor="hand2")
    link.pack(pady=5)
    link.bind("<Button-1>", lambda e: webbrowser.open_new("https://account.arin.net/public/account-setup"))

    tk.Label(dialog, text="ARIN API Key:").pack(pady=5)
    key_entry = tk.Entry(dialog, width=50)
    key_entry.pack(pady=5)
    key_entry.insert(0, config.arin.api_key)

    def save_and_close():
        config.arin.api_key = key_entry.get().strip()
        if config.arin.is_configured:
            config.save_arin()
            messagebox.showinfo("Success", "ARIN API key saved.")
        dialog.destroy()

    tk.Button(dialog, text="Save", command=save_and_close).pack(pady=10)
    tk.Button(dialog, text="Skip", command=dialog.destroy).pack(pady=5)
    return dialog


# ── Customise Statements ─────────────────────────────────────────

# Display names for built-in statement keys (inner tabs under "Standard report statements").
_STANDARD_STATEMENT_TAB_LABELS: Dict[str, str] = {
    "intro": "Introduction",
    "meta": "Meta disclaimer",
    "ip_intro": "IP analysis intro",
    "bingimage": "Bing Image",
    "xcorp": "X (Twitter)",
}


def show_customize_statements_dialog(root: tk.Tk, stmts: StatementManager,
                                      template=None) -> None:
    dialog = tk.Toplevel(root)
    dialog.title("Customize Report Statements")
    dialog.geometry("1000x800")
    dialog.transient(root)
    dialog.grab_set()

    notebook = ttk.Notebook(dialog)
    notebook.pack(fill="both", expand=True, padx=10, pady=10)

    # ── Outer tab 1: all built-in statements in a nested notebook ──
    standard_outer = ttk.Frame(notebook)
    notebook.add(standard_outer, text="Standard report statements")

    intro = ttk.Frame(standard_outer)
    intro.pack(fill="x", padx=12, pady=(10, 4))
    tk.Label(
        intro,
        text="These are the default report paragraphs (introduction, provider notes, IP intro, etc.). "
             "They are separate from custom statements you add elsewhere.",
        font=("Arial", 9),
        fg="#444",
        wraplength=900,
        justify=tk.LEFT,
    ).pack(anchor="w")

    inner_nb = ttk.Notebook(standard_outer)
    inner_nb.pack(fill="both", expand=True, padx=8, pady=(4, 10))

    for key, default_text in DEFAULT_STATEMENTS.items():
        frame = ttk.Frame(inner_nb)
        tab_lbl = _STANDARD_STATEMENT_TAB_LABELS.get(
            key, key.replace("_", " ").strip().title()
        )
        inner_nb.add(frame, text=tab_lbl)
        tk.Label(frame, text=f"Edit {tab_lbl} text:").pack(pady=5)
        if key == "intro":
            tk.Label(frame, text="Tip: Use [CURRENT_DATE], [INVESTIGATOR_NAME], [INVESTIGATOR_TITLE], "
                     "[CYBERTIP_NUMBER], [REPORT_DATE_RECEIVED] placeholders.",
                     fg="gray", wraplength=700).pack(pady=2)
        tw = tk.Text(frame, width=80, height=20, wrap="word")
        tw.pack(pady=5)
        current_val = stmts.statements.get(key, default_text)
        display_text = current_val if isinstance(current_val, str) else current_val.get("text", str(default_text))
        tw.insert(tk.END, display_text)

        def save_tab(k=key, widget=tw):
            stmts.statements[k] = widget.get(1.0, tk.END).strip()

        def reset_tab(k=key, widget=tw, dt=default_text):
            widget.delete(1.0, tk.END)
            widget.insert(tk.END, dt)

        btn_row = ttk.Frame(frame)
        btn_row.pack(pady=5)
        tk.Button(btn_row, text="Save This Statement", command=save_tab).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_row, text="Reset to Default", command=reset_tab).pack(side=tk.LEFT, padx=5)

    _manage_statement_refresh_hooks: List[Callable[[], None]] = []
    _build_add_new_tab(
        notebook,
        stmts,
        template=template,
        manage_refresh_hooks=_manage_statement_refresh_hooks,
        tab_title="Add custom statement",
    )
    _build_manage_tab(
        notebook,
        stmts,
        template=template,
        manage_refresh_hooks=_manage_statement_refresh_hooks,
        tab_title="Manage custom statements",
    )

    bottom = ttk.Frame(dialog)
    bottom.pack(fill="x", padx=10, pady=5)

    def close_and_save():
        stmts.save()
        dialog.destroy()

    def _export_stmts():
        path = filedialog.asksaveasfilename(
            title="Export Custom Statements", defaultextension=".json",
            filetypes=[("JSON files", "*.json")],
        )
        if path:
            try:
                stmts.export_to_file(path)
                messagebox.showinfo("Exported", f"Statements exported to:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {e}")

    def _import_stmts():
        path = filedialog.askopenfilename(
            title="Import Statements", filetypes=[("JSON files", "*.json")],
        )
        if not path:
            return
        try:
            count = stmts.import_from_file(path, mode="merge")
            for hook in _manage_statement_refresh_hooks:
                hook()
            messagebox.showinfo("Imported", f"{count} statement(s) imported.")
        except Exception as e:
            messagebox.showerror("Error", f"Import failed: {e}")

    ttk.Button(bottom, text="Export Statements", command=_export_stmts).pack(side=tk.LEFT, padx=5)
    ttk.Button(bottom, text="Import Statements", command=_import_stmts).pack(side=tk.LEFT, padx=5)
    ttk.Button(bottom, text="Save All & Close", command=close_and_save).pack(side=tk.RIGHT, padx=5)
    ttk.Button(bottom, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)


# ── Statement Templates ──────────────────────────────────────────

STATEMENT_TEMPLATES: Dict[str, Dict[str, str]] = {
    "Legal Disclaimer": {
        "name": "legal_disclaimer",
        "text": "This report is for official use only and may contain sensitive "
                "information protected by law. Unauthorized disclosure, copying, or "
                "distribution of this report or its contents is strictly prohibited.",
        "desc": "Standard law-enforcement confidentiality notice.",
    },
    "Investigator Note": {
        "name": "investigator_note",
        "text": "Investigator's additional observations:\n\n[Enter details here]",
        "desc": "Placeholder for free-form investigator observations.",
    },
    "Chain of Custody Notice": {
        "name": "chain_of_custody",
        "text": "This report and all associated digital evidence have been handled in "
                "accordance with department chain-of-custody procedures. All files were "
                "downloaded directly from the NCMEC CyberTipline portal and stored on "
                "a forensic workstation.",
        "desc": "Documents how evidence was received and stored.",
    },
    "Evidence Handling Note": {
        "name": "evidence_handling",
        "text": "Digital evidence referenced in this report has been preserved in its "
                "original form. Hash values (MD5) listed for each file can be used to "
                "verify that evidence has not been altered.",
        "desc": "Explains evidence integrity verification.",
    },
    "ESP Response Disclaimer": {
        "name": "esp_response_disclaimer",
        "text": "Information provided by the Electronic Service Provider (ESP) in this "
                "CyberTip has not been independently verified by the investigating agency "
                "unless otherwise noted. Subscriber information and account details are "
                "as reported by the ESP.",
        "desc": "Notes that ESP-provided data is unverified.",
    },
    "Jurisdiction Statement": {
        "name": "jurisdiction_statement",
        "text": "Based on the IP address geolocation and/or subscriber information "
                "contained in this CyberTip, the reported activity appears to have "
                "occurred within the jurisdiction of [AGENCY_NAME].",
        "desc": "Establishes jurisdictional basis for the investigation.",
    },
    "Sensitivity Warning": {
        "name": "sensitivity_warning",
        "text": "WARNING: This report contains references to child sexual abuse material "
                "(CSAM). The content described herein is disturbing in nature. This report "
                "should only be reviewed by authorized personnel.",
        "desc": "Content sensitivity warning for report readers.",
    },
    "Case Assignment Note": {
        "name": "case_assignment",
        "text": "This CyberTip has been assigned to [INVESTIGATOR_NAME] for follow-up "
                "investigation. Case number: [CASE_NUMBER]. Date assigned: [CURRENT_DATE].",
        "desc": "Records case assignment and tracking information.",
    },
    "Duplicate from Existing...": {
        "name": "",
        "text": "",
        "desc": "Copy text from an existing custom statement.",
    },
}

PLACEMENT_DESCRIPTIONS: Dict[str, str] = {
    "At Beginning of Report": "Before the investigator introduction paragraph (absolute top of the narrative).",
    "After introduction paragraph": "Right after the opening paragraph, before Incident Summary and other main sections.",
    "Before Incident Summary": "Immediately above the INCIDENT SUMMARY section.",
    "After Incident Summary": "Immediately after incident details (type, date, ESP, etc.).",
    "Before Suspect Information": "Immediately above the SUSPECT INFORMATION section.",
    "After Suspect Information": "Immediately after all suspect / person details.",
    "Before Evidence Summary": "Immediately above the EVIDENCE SUMMARY section.",
    "After Evidence Summary": "Immediately after all evidence file listings.",
    "Before IP Address Analysis": "Immediately above the IP ADDRESS ANALYSIS section.",
    "After IP Address Analysis": "Immediately after all IP lookup results.",
    "After main sections (before final custom block)": "After Incident, Suspect, Evidence, IP, and custom template sections—before the closing CUSTOM STATEMENTS area.",
    "At End of Report": "Inside the final CUSTOM STATEMENTS block at the very end of the report.",
}


def _placement_description(label: str) -> str:
    """Help text for a placement combobox label (built-in or custom template section)."""
    if label in PLACEMENT_DESCRIPTIONS:
        return PLACEMENT_DESCRIPTIONS[label]
    if label.startswith("Before "):
        return (
            f'Immediately before the "{label[7:]}" section in the report (from your report template).'
        )
    if label.startswith("After "):
        return (
            f'Immediately after the "{label[6:]}" section in the report (from your report template).'
        )
    return ""

HIGHLIGHT_OPTIONS = ["(none)", "yellow", "green", "cyan", "magenta", "red", "blue"]


# ── Add New Tab ───────────────────────────────────────────────────

def _build_add_new_tab(
    notebook: ttk.Notebook,
    stmts: StatementManager,
    template=None,
    manage_refresh_hooks: Optional[List[Callable[[], None]]] = None,
    tab_title: str = "Add custom statement",
) -> None:
    outer = ttk.Frame(notebook)
    notebook.add(outer, text=tab_title)
    canvas = tk.Canvas(outer, highlightthickness=0)
    vsb = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
    add_frame = ttk.Frame(canvas)
    add_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=add_frame, anchor="nw")
    canvas.configure(yscrollcommand=vsb.set)
    canvas.pack(side="left", fill="both", expand=True)
    vsb.pack(side="right", fill="y")

    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    canvas.bind_all("<MouseWheel>", _on_mousewheel, add="+")

    # ── Step 1: Starting Point ────────────────────────────────────
    step1 = ttk.LabelFrame(add_frame, text="  Step 1 -- Choose a Starting Point  ", padding=10)
    step1.pack(fill="x", padx=10, pady=(10, 5))

    template_names = ["None (blank)"] + list(STATEMENT_TEMPLATES.keys())
    tmpl_var = tk.StringVar(value=template_names[0])

    tk.Label(step1, text="Start with a template to pre-fill common text, or begin from scratch:").pack(anchor="w")
    tmpl_frame = ttk.Frame(step1)
    tmpl_frame.pack(fill="x", pady=5)
    tk.Label(tmpl_frame, text="Template:").pack(side=tk.LEFT, padx=(0, 5))
    tmpl_cb = ttk.Combobox(tmpl_frame, textvariable=tmpl_var, values=template_names,
                           state="readonly", width=35)
    tmpl_cb.pack(side=tk.LEFT)

    tmpl_desc_var = tk.StringVar(value="")
    tmpl_desc_label = tk.Label(step1, textvariable=tmpl_desc_var, fg="gray",
                               wraplength=700, justify=tk.LEFT)
    tmpl_desc_label.pack(anchor="w", pady=(2, 0))

    # ── Step 2: Statement Identity ────────────────────────────────
    step2 = ttk.LabelFrame(add_frame, text="  Step 2 -- Name & Placement  ", padding=10)
    step2.pack(fill="x", padx=10, pady=5)

    name_row = ttk.Frame(step2)
    name_row.pack(fill="x", pady=3)
    tk.Label(name_row, text="Statement Name:").pack(side=tk.LEFT, padx=(0, 5))
    key_entry = tk.Entry(name_row, width=40)
    key_entry.pack(side=tk.LEFT)
    tk.Label(name_row, text="(unique identifier, e.g. 'chain_of_custody')", fg="gray").pack(side=tk.LEFT, padx=5)

    place_row = ttk.Frame(step2)
    place_row.pack(fill="x", pady=3)
    tk.Label(place_row, text="Insert Location:").pack(side=tk.LEFT, padx=(0, 5))
    _add_all_prefixes = get_all_placement_prefixes(template)
    placement_opts = list(_add_all_prefixes.keys())
    placement_var = tk.StringVar(value="At End of Report")
    place_cb = ttk.Combobox(place_row, textvariable=placement_var, values=placement_opts,
                            state="readonly", width=35)
    place_cb.pack(side=tk.LEFT)

    place_desc_var = tk.StringVar(value=_placement_description("At End of Report"))
    place_desc_label = tk.Label(step2, textvariable=place_desc_var, fg="gray",
                                wraplength=700, justify=tk.LEFT)
    place_desc_label.pack(anchor="w", pady=(2, 0))
    tk.Label(
        step2,
        text="If you use the same placement for more than one statement, they are ordered alphabetically by statement name.",
        fg="gray",
        font=("Arial", 8),
        wraplength=700,
        justify=tk.LEFT,
    ).pack(anchor="w", pady=(4, 0))

    def _on_placement_change(*_args):
        place_desc_var.set(_placement_description(placement_var.get()))
    placement_var.trace_add("write", _on_placement_change)

    # ── Step 3: Conditional Display ───────────────────────────────
    step3 = ttk.LabelFrame(add_frame, text="  Step 3 -- When Should This Statement Appear? (Optional)  ", padding=10)
    step3.pack(fill="x", padx=10, pady=5)

    cond_mode = tk.StringVar(value="always")
    tk.Radiobutton(step3, text="Include in ALL reports (no condition)",
                   variable=cond_mode, value="always").pack(anchor="w")
    tk.Radiobutton(step3, text="Only include when the report is from specific ESP(s):",
                   variable=cond_mode, value="specific").pack(anchor="w", pady=(5, 0))

    esp_check_frame = ttk.Frame(step3)
    esp_check_frame.pack(fill="x", padx=25, pady=5)
    esp_vars: Dict[str, tk.BooleanVar] = {}

    col_count = 4
    for i, esp in enumerate(_ESP_LIST):
        var = tk.BooleanVar(value=False)
        esp_vars[esp] = var
        r, c = divmod(i, col_count)
        cb = ttk.Checkbutton(esp_check_frame, text=esp, variable=var)
        cb.grid(row=r, column=c, sticky="w", padx=5, pady=1)

    custom_esp_row = ttk.Frame(step3)
    custom_esp_row.pack(fill="x", padx=25, pady=(0, 5))
    tk.Label(custom_esp_row, text="Other ESP (type name):").pack(side=tk.LEFT, padx=(0, 5))
    custom_esp_entry = tk.Entry(custom_esp_row, width=25)
    custom_esp_entry.pack(side=tk.LEFT)

    tk.Radiobutton(step3, text="Only include when data condition(s) are met (all selected must match):",
                   variable=cond_mode, value="data").pack(anchor="w", pady=(10, 0))

    _data_conditions = [
        ("has_evidence", "Report has evidence files"),
        ("has_ips", "Report has IP addresses"),
        ("is_multi_tip", "Analyzing multiple CyberTips"),
        ("suspect_count > 1", "More than one suspect"),
    ]
    data_cond_frame = ttk.Frame(step3)
    data_cond_frame.pack(fill="x", padx=25, pady=5)
    data_cond_vars: Dict[str, tk.BooleanVar] = {}
    for _dc_val, _dc_label in _data_conditions:
        _dv = tk.BooleanVar(value=False)
        data_cond_vars[_dc_val] = _dv
        ttk.Checkbutton(data_cond_frame, text=_dc_label, variable=_dv).pack(anchor="w")

    cond_summary_var = tk.StringVar(value="Condition: None (appears in all reports)")
    cond_summary = tk.Label(step3, textvariable=cond_summary_var, font=("Arial", 9, "italic"), fg="blue")
    cond_summary.pack(anchor="w", pady=(5, 0))

    def _update_cond_summary(*_args):
        mode = cond_mode.get()
        if mode == "always":
            cond_summary_var.set("Condition: None (appears in all reports)")
            for w in esp_check_frame.winfo_children():
                w.configure(state="disabled")
            custom_esp_entry.configure(state="disabled")
            for w in data_cond_frame.winfo_children():
                w.configure(state="disabled")
        elif mode == "specific":
            for w in esp_check_frame.winfo_children():
                w.configure(state="normal")
            custom_esp_entry.configure(state="normal")
            for w in data_cond_frame.winfo_children():
                w.configure(state="disabled")
            selected = [name for name, var in esp_vars.items() if var.get()]
            custom = custom_esp_entry.get().strip()
            if custom:
                selected.append(custom)
            if not selected:
                cond_summary_var.set("Condition: (select at least one ESP above)")
            elif len(selected) == 1:
                cond_summary_var.set(f'Condition: Only when ESP is "{selected[0]}"')
            else:
                cond_summary_var.set(f"Condition: Only when ESP is one of: {', '.join(selected)}")
        elif mode == "data":
            for w in esp_check_frame.winfo_children():
                w.configure(state="disabled")
            custom_esp_entry.configure(state="disabled")
            for w in data_cond_frame.winfo_children():
                w.configure(state="normal")
            picked = [lab for val, lab in _data_conditions if data_cond_vars[val].get()]
            if not picked:
                cond_summary_var.set("Condition: (select at least one data condition above)")
            elif len(picked) == 1:
                cond_summary_var.set(f"Condition: {picked[0]}")
            else:
                cond_summary_var.set("Condition: All of — " + "; ".join(picked))

    cond_mode.trace_add("write", _update_cond_summary)
    for var in esp_vars.values():
        var.trace_add("write", _update_cond_summary)
    custom_esp_entry.bind("<KeyRelease>", lambda e: _update_cond_summary())
    for _dv in data_cond_vars.values():
        _dv.trace_add("write", _update_cond_summary)
    _update_cond_summary()

    def _get_condition() -> str:
        mode = cond_mode.get()
        if mode == "always":
            return ""
        if mode == "data":
            parts = [val for val, _ in _data_conditions if data_cond_vars[val].get()]
            if not parts:
                return ""
            return "&&".join(parts)
        selected = [name for name, var in esp_vars.items() if var.get()]
        custom = custom_esp_entry.get().strip()
        if custom:
            selected.append(custom)
        if not selected:
            return ""
        if len(selected) == 1:
            return f'esp_name == "{selected[0]}"'
        quoted = ", ".join('"' + v + '"' for v in selected)
        return f"esp_name in [{quoted}]"

    # ── Step 4: Statement Text ────────────────────────────────────
    step4 = ttk.LabelFrame(add_frame, text="  Step 4 -- Write Your Statement  ", padding=10)
    step4.pack(fill="x", padx=10, pady=5)

    toolbar = ttk.Frame(step4)
    toolbar.pack(fill="x", pady=(0, 5))
    tk.Label(toolbar, text="Insert:").pack(side=tk.LEFT, padx=(0, 5))

    add_text = tk.Text(step4, width=90, height=12, wrap="word", font=("Consolas", 10))
    add_text.pack(fill="x", pady=5)

    def _insert_snippet(snippet: str):
        add_text.insert(tk.INSERT, snippet)
        add_text.focus_set()

    snippets = [
        ("Today's Date", "[CURRENT_DATE]"),
        ("Investigator", "[INVESTIGATOR_NAME]"),
        ("Title", "[INVESTIGATOR_TITLE]"),
        ("Case #", "[CASE_NUMBER]"),
        ("Agency", "[AGENCY_NAME]"),
        ("CyberTip #", "[CYBERTIP_NUMBER]"),
        ("ESP", "[ESP_NAME]"),
        ("Suspect Name", "[SUSPECT_NAME]"),
        ("Suspect Email", "[SUSPECT_EMAIL]"),
        ("Suspect Phone", "[SUSPECT_PHONE]"),
        ("Screen Name", "[SUSPECT_SCREEN_NAME]"),
        ("Incident Date", "[INCIDENT_DATE]"),
        ("Incident Type", "[INCIDENT_TYPE]"),
        ("Evidence Count", "[EVIDENCE_COUNT]"),
    ]
    for label, snip in snippets:
        ttk.Button(toolbar, text=label,
                   command=lambda s=snip: _insert_snippet(s)).pack(side=tk.LEFT, padx=2)

    char_count_var = tk.StringVar(value="0 characters")
    tk.Label(step4, textvariable=char_count_var, fg="gray").pack(anchor="e")

    def _update_char_count(*_args):
        length = len(add_text.get(1.0, tk.END).strip())
        char_count_var.set(f"{length} character{'s' if length != 1 else ''}")
    add_text.bind("<KeyRelease>", _update_char_count)

    # ── Step 5: Formatting (Optional) ─────────────────────────────
    step5 = ttk.LabelFrame(add_frame, text="  Step 5 -- DOCX Formatting (Optional)  ", padding=10)
    step5.pack(fill="x", padx=10, pady=5)

    fmt_row1 = ttk.Frame(step5)
    fmt_row1.pack(fill="x", pady=3)

    tk.Label(fmt_row1, text="Font Size:").pack(side=tk.LEFT, padx=(0, 3))
    font_size_var = tk.IntVar(value=12)
    tk.Spinbox(fmt_row1, from_=8, to=28, textvariable=font_size_var, width=4).pack(side=tk.LEFT, padx=(0, 15))

    bold_var = tk.BooleanVar(value=False)
    ttk.Checkbutton(fmt_row1, text="Bold", variable=bold_var).pack(side=tk.LEFT, padx=5)

    italic_var = tk.BooleanVar(value=False)
    ttk.Checkbutton(fmt_row1, text="Italic", variable=italic_var).pack(side=tk.LEFT, padx=5)

    fmt_row2 = ttk.Frame(step5)
    fmt_row2.pack(fill="x", pady=3)

    tk.Label(fmt_row2, text="Left Indent (inches):").pack(side=tk.LEFT, padx=(0, 3))
    indent_var = tk.DoubleVar(value=0.0)
    tk.Spinbox(fmt_row2, from_=0.0, to=3.0, increment=0.25, textvariable=indent_var, width=5).pack(side=tk.LEFT, padx=(0, 15))

    tk.Label(fmt_row2, text="Highlight:").pack(side=tk.LEFT, padx=(0, 3))
    highlight_var = tk.StringVar(value="(none)")
    ttk.Combobox(fmt_row2, textvariable=highlight_var, values=HIGHLIGHT_OPTIONS,
                 state="readonly", width=12).pack(side=tk.LEFT)

    def _get_formatting() -> dict:
        return {
            "font_size": font_size_var.get(),
            "bold": bold_var.get(),
            "italic": italic_var.get(),
            "indent": indent_var.get(),
            "highlight": "" if highlight_var.get() == "(none)" else highlight_var.get(),
        }

    # ── Live Preview ──────────────────────────────────────────────
    step6 = ttk.LabelFrame(add_frame, text="  Preview  ", padding=10)
    step6.pack(fill="x", padx=10, pady=5)

    preview_text = tk.Text(step6, width=90, height=6, wrap="word", state="disabled",
                           bg="#f5f5f5", font=("Times New Roman", 10))
    preview_text.pack(fill="x")

    def _refresh_preview(*_args):
        name = key_entry.get().strip() or "(unnamed)"
        placement = placement_var.get()
        body = add_text.get(1.0, tk.END).strip() or "(no text entered)"
        cond = cond_summary_var.get()
        fmt = _get_formatting()
        fmt_desc_parts = []
        if fmt["bold"]:
            fmt_desc_parts.append("Bold")
        if fmt["italic"]:
            fmt_desc_parts.append("Italic")
        fmt_desc_parts.append(f"{fmt['font_size']}pt")
        if fmt["indent"] > 0:
            fmt_desc_parts.append(f"Indent {fmt['indent']}in")
        if fmt["highlight"]:
            fmt_desc_parts.append(f"Highlight: {fmt['highlight']}")
        fmt_desc = ", ".join(fmt_desc_parts)

        preview_text.configure(state="normal")
        preview_text.delete(1.0, tk.END)
        preview_text.insert(tk.END, f"--- {placement} ---\n\n")
        preview_text.insert(tk.END, f"{name.upper()}: {body}\n\n")
        preview_text.insert(tk.END, f"[{cond}]  |  Format: {fmt_desc}")
        preview_text.configure(state="disabled")

    add_text.bind("<KeyRelease>", lambda e: (_update_char_count(), _refresh_preview()))
    placement_var.trace_add("write", _refresh_preview)
    cond_summary_var.trace_add("write", _refresh_preview)
    key_entry.bind("<KeyRelease>", lambda e: _refresh_preview())
    for fv in (bold_var, italic_var):
        fv.trace_add("write", _refresh_preview)
    font_size_var.trace_add("write", _refresh_preview)
    highlight_var.trace_add("write", _refresh_preview)

    # ── Template application logic ────────────────────────────────
    def _on_template_change(event=None):
        chosen = tmpl_var.get()
        if chosen == "None (blank)":
            tmpl_desc_var.set("")
            return

        tmpl = STATEMENT_TEMPLATES.get(chosen)
        if not tmpl:
            return

        tmpl_desc_var.set(tmpl["desc"])

        if chosen == "Duplicate from Existing...":
            _show_duplicate_picker()
            return

        key_entry.delete(0, tk.END)
        key_entry.insert(0, tmpl["name"])
        add_text.delete(1.0, tk.END)
        add_text.insert(tk.END, tmpl["text"])
        _update_char_count()
        _refresh_preview()

    tmpl_cb.bind("<<ComboboxSelected>>", _on_template_change)

    def _show_duplicate_picker():
        custom = {k: v for k, v in stmts.statements.items() if k not in DEFAULT_STATEMENTS}
        if not custom:
            messagebox.showinfo("Info", "No custom statements to duplicate from.")
            tmpl_var.set(template_names[0])
            return

        picker = tk.Toplevel(add_frame.winfo_toplevel())
        picker.title("Duplicate From Existing Statement")
        picker.geometry("500x350")
        picker.transient(add_frame.winfo_toplevel())
        picker.grab_set()

        tk.Label(picker, text="Select a statement to copy:").pack(pady=5)
        lb = tk.Listbox(picker, width=60, height=12)
        lb.pack(fill="both", expand=True, padx=10, pady=5)
        keys = sorted(custom.keys())
        for k in keys:
            placement = stmts.get_placement_label(k)
            lb.insert(tk.END, f"{k}  ({placement})")

        def _do_duplicate():
            sel = lb.curselection()
            if not sel:
                messagebox.showwarning("Warning", "Select a statement to duplicate.")
                return
            orig_key = keys[sel[0]]
            val = stmts.statements[orig_key]
            text = val["text"] if isinstance(val, dict) else val
            key_entry.delete(0, tk.END)
            key_entry.insert(0, f"{orig_key}_copy")
            add_text.delete(1.0, tk.END)
            add_text.insert(tk.END, text)
            if isinstance(val, dict) and val.get("condition"):
                _apply_condition_from_string(val["condition"])
            if isinstance(val, dict) and val.get("formatting"):
                fmt = val["formatting"]
                font_size_var.set(fmt.get("font_size", 12))
                bold_var.set(fmt.get("bold", False))
                italic_var.set(fmt.get("italic", False))
                indent_var.set(fmt.get("indent", 0.0))
                highlight_var.set(fmt.get("highlight", "") or "(none)")
            _update_char_count()
            _refresh_preview()
            picker.destroy()

        tk.Button(picker, text="Duplicate", command=_do_duplicate).pack(pady=5)
        tk.Button(picker, text="Cancel", command=picker.destroy).pack(pady=5)

    def _apply_condition_from_string(cond: str):
        if not cond:
            cond_mode.set("always")
            for _v in data_cond_vars.values():
                _v.set(False)
            return
        data_keys = {v for v, _ in _data_conditions}
        data_parts: Optional[List[str]] = None
        if "&&" in cond:
            cand = [p.strip() for p in cond.split("&&") if p.strip()]
            if cand and all(p in data_keys or p.startswith("suspect_count") for p in cand):
                data_parts = cand
        elif cond in data_keys or cond.startswith("suspect_count"):
            data_parts = [cond]
        if data_parts is not None:
            cond_mode.set("data")
            for _k, _v in data_cond_vars.items():
                _v.set(_k in data_parts)
            return
        cond_mode.set("specific")
        for _v in data_cond_vars.values():
            _v.set(False)
        values = re.findall(r'"([^"]*)"', cond)
        for name, var in esp_vars.items():
            var.set(name in values)
        remaining = [v for v in values if v not in esp_vars]
        if remaining:
            custom_esp_entry.delete(0, tk.END)
            custom_esp_entry.insert(0, remaining[0])

    # ── Action Buttons ────────────────────────────────────────────
    btn_frame = ttk.Frame(add_frame)
    btn_frame.pack(fill="x", padx=10, pady=(5, 15))

    def _add_statement():
        name = key_entry.get().strip()
        if not name:
            messagebox.showwarning("Missing Name", "Please enter a statement name in Step 2.")
            key_entry.focus_set()
            return
        if not name.replace("_", "").replace("-", "").isalnum():
            messagebox.showwarning(
                "Invalid Name",
                "Statement name should only contain letters, numbers, underscores, or hyphens.",
            )
            key_entry.focus_set()
            return

        prefix = _add_all_prefixes.get(placement_var.get(), "")
        full_key = f"{prefix}{name}" if prefix else name
        if full_key in stmts.statements:
            messagebox.showwarning(
                "Duplicate Name",
                f"A statement named '{name}' already exists at '{placement_var.get()}'.\n"
                "Choose a different name or placement.",
            )
            return

        text = add_text.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Missing Text", "Please write your statement text in Step 4.")
            add_text.focus_set()
            return

        cond = _get_condition()
        if cond_mode.get() == "specific" and not cond:
            messagebox.showwarning(
                "No ESPs Selected",
                "You chose to limit this statement to specific ESPs but didn't select any.\n"
                "Either select ESP(s) in Step 3 or switch to 'Include in ALL reports'.",
            )
            return
        if cond_mode.get() == "data" and not cond:
            messagebox.showwarning(
                "No Data Conditions Selected",
                "You chose data conditions but didn't select any.\n"
                "Select one or more checkboxes in Step 3 or switch to another option.",
            )
            return

        stmt_data = {
            "text": text,
            "condition": cond,
            "formatting": _get_formatting(),
            "enabled": True,
        }
        stmts.statements[full_key] = stmt_data
        stmts.selected.add(full_key)

        if manage_refresh_hooks:
            for _hook in manage_refresh_hooks:
                _hook()

        placement_label = placement_var.get()
        cond_label = cond_summary_var.get()
        messagebox.showinfo(
            "Statement Added",
            f"Successfully added:\n\n"
            f"Name: {name}\n"
            f"Location: {placement_label}\n"
            f"{cond_label}",
        )
        _clear_form()

    def _clear_form():
        tmpl_var.set(template_names[0])
        tmpl_desc_var.set("")
        key_entry.delete(0, tk.END)
        placement_var.set("At End of Report")
        cond_mode.set("always")
        for var in esp_vars.values():
            var.set(False)
        for var in data_cond_vars.values():
            var.set(False)
        custom_esp_entry.delete(0, tk.END)
        add_text.delete(1.0, tk.END)
        font_size_var.set(12)
        bold_var.set(False)
        italic_var.set(False)
        indent_var.set(0.0)
        highlight_var.set("(none)")
        _update_char_count()
        _refresh_preview()

    ttk.Button(btn_frame, text="Add Statement", command=_add_statement).pack(side=tk.LEFT, padx=5)
    ttk.Button(btn_frame, text="Clear Form", command=_clear_form).pack(side=tk.LEFT, padx=5)

    _refresh_preview()


# ── Manage Custom Statements Tab (full overhaul) ─────────────────

def _build_manage_tab(
    notebook: ttk.Notebook,
    stmts: StatementManager,
    template=None,
    manage_refresh_hooks: Optional[List[Callable[[], None]]] = None,
    tab_title: str = "Manage custom statements",
) -> None:
    manage = ttk.Frame(notebook)
    notebook.add(manage, text=tab_title)

    # Top: search
    top = ttk.Frame(manage)
    top.pack(fill="x", padx=10, pady=5)
    tk.Label(top, text="Search:").pack(side=tk.LEFT, padx=(0, 5))
    search_var = tk.StringVar()
    tk.Entry(top, textvariable=search_var, width=40).pack(side=tk.LEFT)

    # Main pane
    pane = ttk.PanedWindow(manage, orient=tk.HORIZONTAL)
    pane.pack(fill="both", expand=True, padx=10, pady=5)

    # Left: statement list with columns
    left = ttk.Frame(pane)
    pane.add(left, weight=1)

    list_row = ttk.Frame(left)
    list_row.pack(fill="both", expand=True)

    columns = ("key", "placement", "enabled")
    tree = ttk.Treeview(list_row, columns=columns, show="headings", height=12)
    tree.heading("key", text="Name")
    tree.heading("placement", text="Placement")
    tree.heading("enabled", text="In reports?")
    tree.column("key", width=200)
    tree.column("placement", width=200)
    tree.column("enabled", width=72, anchor="center")
    tree.pack(side=tk.LEFT, fill="both", expand=True)

    tree_sb = ttk.Scrollbar(list_row, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=tree_sb.set)
    tree_sb.pack(side=tk.RIGHT, fill=tk.Y)

    tk.Label(
        left,
        text="Select a statement, then use Statement actions (top right) to include, save, or remove it.",
        font=("Arial", 8),
        fg="gray",
        wraplength=240,
        justify=tk.LEFT,
    ).pack(anchor="w", padx=2, pady=(4, 0))

    # Right: scrollable editing panel (content can exceed window height)
    right = ttk.Frame(pane)
    pane.add(right, weight=2)

    right_canvas = tk.Canvas(right, highlightthickness=0)
    right_scroll = ttk.Scrollbar(right, orient="vertical", command=right_canvas.yview)
    right_canvas.configure(yscrollcommand=right_scroll.set)
    right_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    right_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    right_inner = ttk.Frame(right_canvas)
    right_win = right_canvas.create_window((0, 0), window=right_inner, anchor="nw")

    def _manage_right_on_inner_configure(_event=None):
        right_canvas.configure(scrollregion=right_canvas.bbox("all"))

    def _manage_right_on_canvas_configure(event):
        # Match inner width to visible canvas so text wraps and horizontal scroll is not needed
        pad = 8
        right_canvas.itemconfig(right_win, width=max(1, event.width - pad))

    right_inner.bind("<Configure>", _manage_right_on_inner_configure)
    right_canvas.bind("<Configure>", _manage_right_on_canvas_configure)

    def _manage_right_mousewheel(event):
        right_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    right_canvas.bind("<MouseWheel>", _manage_right_mousewheel)

    def _bind_manage_panel_mousewheel(widget):
        # Do not steal wheel from text fields, spinboxes, or comboboxes.
        if isinstance(widget, (tk.Text, tk.Entry, tk.Spinbox, ttk.Combobox)):
            for ch in widget.winfo_children():
                _bind_manage_panel_mousewheel(ch)
            return
        widget.bind("<MouseWheel>", _manage_right_mousewheel)
        for ch in widget.winfo_children():
            _bind_manage_panel_mousewheel(ch)

    actions_top = ttk.LabelFrame(right_inner, text="  Statement actions  ", padding=8)
    actions_top.pack(fill="x", pady=(0, 8))

    tk.Label(right_inner, text="Statement text:", font=("Arial", 10, "bold")).pack(anchor="w", pady=(5, 2))
    edit_text = tk.Text(right_inner, width=55, height=8, wrap="word", font=("Consolas", 10))
    edit_text.pack(fill="x", pady=2)

    # Condition editor (ESP checkbox UI)
    cond_frame = ttk.LabelFrame(right_inner, text="  Condition  ", padding=5)
    cond_frame.pack(fill="x", pady=5)

    m_cond_mode = tk.StringVar(value="always")
    tk.Radiobutton(cond_frame, text="All reports", variable=m_cond_mode, value="always").pack(anchor="w")
    tk.Radiobutton(cond_frame, text="Specific ESP(s):", variable=m_cond_mode, value="specific").pack(anchor="w")

    m_esp_check_frame = ttk.Frame(cond_frame)
    m_esp_check_frame.pack(fill="x", padx=15, pady=2)
    m_esp_vars: Dict[str, tk.BooleanVar] = {}
    _col_count = 3
    for _i, _esp in enumerate(_ESP_LIST):
        _var = tk.BooleanVar(value=False)
        m_esp_vars[_esp] = _var
        _r, _c = divmod(_i, _col_count)
        ttk.Checkbutton(m_esp_check_frame, text=_esp, variable=_var).grid(row=_r, column=_c, sticky="w", padx=3, pady=1)

    m_custom_esp_row = ttk.Frame(cond_frame)
    m_custom_esp_row.pack(fill="x", padx=15)
    tk.Label(m_custom_esp_row, text="Other ESP:").pack(side=tk.LEFT, padx=(0, 3))
    m_custom_esp = tk.Entry(m_custom_esp_row, width=20)
    m_custom_esp.pack(side=tk.LEFT)

    tk.Radiobutton(
        cond_frame,
        text="Data condition(s) — all selected must match:",
        variable=m_cond_mode,
        value="data",
    ).pack(anchor="w", pady=(6, 0))
    m_data_cond_frame = ttk.Frame(cond_frame)
    m_data_cond_frame.pack(fill="x", padx=15, pady=2)
    _m_data_conditions = [
        ("has_evidence", "Report has evidence files"),
        ("has_ips", "Report has IP addresses"),
        ("is_multi_tip", "Analyzing multiple CyberTips"),
        ("suspect_count > 1", "More than one suspect"),
    ]
    m_data_cond_vars: Dict[str, tk.BooleanVar] = {}
    for _dc_val, _dc_label in _m_data_conditions:
        _mdv = tk.BooleanVar(value=False)
        m_data_cond_vars[_dc_val] = _mdv
        ttk.Checkbutton(m_data_cond_frame, text=_dc_label, variable=_mdv).pack(anchor="w")

    def _m_get_condition() -> str:
        mode = m_cond_mode.get()
        if mode == "always":
            return ""
        if mode == "data":
            parts = [val for val, _ in _m_data_conditions if m_data_cond_vars[val].get()]
            if not parts:
                return ""
            return "&&".join(parts)
        selected = [name for name, var in m_esp_vars.items() if var.get()]
        custom = m_custom_esp.get().strip()
        if custom:
            selected.append(custom)
        if not selected:
            return ""
        if len(selected) == 1:
            return f'esp_name == "{selected[0]}"'
        quoted = ", ".join('"' + v + '"' for v in selected)
        return f"esp_name in [{quoted}]"

    def _m_apply_condition(cond: str):
        if not cond:
            m_cond_mode.set("always")
            for var in m_esp_vars.values():
                var.set(False)
            for _v in m_data_cond_vars.values():
                _v.set(False)
            m_custom_esp.delete(0, tk.END)
            return
        _m_data_keys = {v for v, _ in _m_data_conditions}
        data_parts: Optional[List[str]] = None
        if "&&" in cond:
            cand = [p.strip() for p in cond.split("&&") if p.strip()]
            if cand and all(p in _m_data_keys or p.startswith("suspect_count") for p in cand):
                data_parts = cand
        elif cond in _m_data_keys or cond.startswith("suspect_count"):
            data_parts = [cond]
        if data_parts is not None:
            m_cond_mode.set("data")
            for _k, _v in m_data_cond_vars.items():
                _v.set(_k in data_parts)
            for var in m_esp_vars.values():
                var.set(False)
            m_custom_esp.delete(0, tk.END)
            return
        m_cond_mode.set("specific")
        for _v in m_data_cond_vars.values():
            _v.set(False)
        values = re.findall(r'"([^"]*)"', cond)
        for name, var in m_esp_vars.items():
            var.set(name in values)
        remaining = [v for v in values if v not in m_esp_vars]
        m_custom_esp.delete(0, tk.END)
        if remaining:
            m_custom_esp.insert(0, remaining[0])

    # Placement reassignment
    place_frame = ttk.LabelFrame(right_inner, text="  Placement in report  ", padding=5)
    place_frame.pack(fill="x", pady=5)
    _manage_all_prefixes = get_all_placement_prefixes(template)
    manage_place_var = tk.StringVar(value="At End of Report")
    ttk.Combobox(
        place_frame,
        textvariable=manage_place_var,
        values=list(_manage_all_prefixes.keys()),
        state="readonly",
        width=48,
    ).pack(fill="x", pady=2)
    manage_place_desc_var = tk.StringVar(value=_placement_description("At End of Report"))
    tk.Label(
        place_frame,
        textvariable=manage_place_desc_var,
        fg="gray",
        wraplength=520,
        justify=tk.LEFT,
        font=("Arial", 8),
    ).pack(anchor="w", pady=(2, 0))
    manage_place_var.trace_add("write", lambda *_: manage_place_desc_var.set(_placement_description(manage_place_var.get())))

    # Formatting
    fmt_frame = ttk.LabelFrame(right_inner, text="  Formatting  ", padding=5)
    fmt_frame.pack(fill="x", pady=5)
    fmt_r1 = ttk.Frame(fmt_frame)
    fmt_r1.pack(fill="x", pady=2)
    tk.Label(fmt_r1, text="Font Size:").pack(side=tk.LEFT, padx=(0, 3))
    m_fontsize = tk.IntVar(value=12)
    tk.Spinbox(fmt_r1, from_=8, to=28, textvariable=m_fontsize, width=4).pack(side=tk.LEFT, padx=(0, 10))
    m_bold = tk.BooleanVar(value=False)
    ttk.Checkbutton(fmt_r1, text="Bold", variable=m_bold).pack(side=tk.LEFT, padx=3)
    m_italic = tk.BooleanVar(value=False)
    ttk.Checkbutton(fmt_r1, text="Italic", variable=m_italic).pack(side=tk.LEFT, padx=3)

    fmt_r2 = ttk.Frame(fmt_frame)
    fmt_r2.pack(fill="x", pady=2)
    tk.Label(fmt_r2, text="Indent:").pack(side=tk.LEFT, padx=(0, 3))
    m_indent = tk.DoubleVar(value=0.0)
    tk.Spinbox(fmt_r2, from_=0.0, to=3.0, increment=0.25, textvariable=m_indent, width=5).pack(side=tk.LEFT, padx=(0, 10))
    tk.Label(fmt_r2, text="Highlight:").pack(side=tk.LEFT, padx=(0, 3))
    m_highlight = tk.StringVar(value="(none)")
    ttk.Combobox(fmt_r2, textvariable=m_highlight, values=HIGHLIGHT_OPTIONS, state="readonly", width=12).pack(side=tk.LEFT)

    _selected_key = tk.StringVar(value="")
    m_include_var = tk.BooleanVar(value=False)
    _syncing_include = False

    def _refresh_tree(q=""):
        tree.delete(*tree.get_children())
        custom = {k: v for k, v in stmts.statements.items() if k not in DEFAULT_STATEMENTS}
        for key in sorted(custom, key=lambda k: k.lower()):
            val = custom[key]
            text_val = val["text"] if isinstance(val, dict) else val
            if q and q.lower() not in key.lower() and q.lower() not in text_val.lower():
                continue
            placement = stmts.get_placement_label(key, template=template)
            enabled = key in stmts.selected
            tree.insert("", tk.END, iid=key, values=(key, placement, "Yes" if enabled else "No"))

    def _sync_include_from_selection():
        nonlocal _syncing_include
        key = _selected_key.get()
        _syncing_include = True
        try:
            m_include_var.set(bool(key and key in stmts.selected))
        finally:
            _syncing_include = False

    def _on_include_toggle():
        nonlocal _syncing_include
        if _syncing_include:
            return
        key = _selected_key.get()
        if not key or key not in stmts.statements:
            return
        if m_include_var.get():
            stmts.selected.add(key)
        else:
            stmts.selected.discard(key)
        val = stmts.statements.get(key)
        if isinstance(val, dict):
            val["enabled"] = m_include_var.get()
        _refresh_tree(search_var.get())

    def _on_tree_select(event=None):
        sel = tree.selection()
        if not sel:
            _selected_key.set("")
            _sync_include_from_selection()
            return
        key = sel[0]
        _selected_key.set(key)
        val = stmts.statements.get(key)
        if val is None:
            return

        edit_text.delete(1.0, tk.END)
        if isinstance(val, dict):
            edit_text.insert(tk.END, val.get("text", ""))
            _m_apply_condition(val.get("condition", ""))
            fmt = val.get("formatting", {})
            m_fontsize.set(fmt.get("font_size", 12))
            m_bold.set(fmt.get("bold", False))
            m_italic.set(fmt.get("italic", False))
            m_indent.set(fmt.get("indent", 0.0))
            m_highlight.set(fmt.get("highlight", "") or "(none)")
        else:
            edit_text.insert(tk.END, val)
            _m_apply_condition("")
            m_fontsize.set(12)
            m_bold.set(False)
            m_italic.set(False)
            m_indent.set(0.0)
            m_highlight.set("(none)")

        current_placement = stmts.get_placement_label(key, template=template)
        manage_place_var.set(current_placement)
        _sync_include_from_selection()

    tree.bind("<<TreeviewSelect>>", _on_tree_select)
    search_var.trace("w", lambda *_: _refresh_tree(search_var.get()))

    def _save_edits():
        key = _selected_key.get()
        if not key or key not in stmts.statements:
            messagebox.showwarning("Warning", "Select a statement first.")
            return

        new_text = edit_text.get(1.0, tk.END).strip()
        if not new_text:
            messagebox.showwarning("Warning", "Statement text cannot be empty.")
            return

        if m_cond_mode.get() == "specific" and not _m_get_condition():
            messagebox.showwarning(
                "Warning",
                "Specific ESP(s) is selected but no ESP is checked. "
                "Select ESP(s) or choose another condition mode.",
            )
            return
        if m_cond_mode.get() == "data" and not _m_get_condition():
            messagebox.showwarning(
                "Warning",
                "Data condition(s) is selected but none are checked. "
                "Select one or more conditions or choose another mode.",
            )
            return

        new_placement_label = manage_place_var.get()
        new_prefix = _manage_all_prefixes.get(new_placement_label, "")

        bare_name = key
        for pfx in sorted(_manage_all_prefixes.values(), key=len, reverse=True):
            if pfx and key.startswith(pfx):
                bare_name = key[len(pfx):]
                break

        new_key = f"{new_prefix}{bare_name}" if new_prefix else bare_name

        new_val = {
            "text": new_text,
            "condition": _m_get_condition(),
            "formatting": {
                "font_size": m_fontsize.get(),
                "bold": m_bold.get(),
                "italic": m_italic.get(),
                "indent": m_indent.get(),
                "highlight": "" if m_highlight.get() == "(none)" else m_highlight.get(),
            },
            "enabled": key in stmts.selected,
        }

        if new_key != key:
            del stmts.statements[key]
            stmts.selected.discard(key)
            stmts.statements[new_key] = new_val
            if new_val["enabled"]:
                stmts.selected.add(new_key)
            _selected_key.set(new_key)
        else:
            stmts.statements[key] = new_val

        messagebox.showinfo("Saved", f"Statement '{bare_name}' updated.")
        _refresh_tree(search_var.get())
        _sync_include_from_selection()

    def _delete_stmt():
        key = _selected_key.get()
        if not key:
            messagebox.showinfo("Delete", "Select a statement in the list first.")
            return
        if messagebox.askyesno("Delete statement", f"Permanently remove this custom statement?\n\n{key}"):
            del stmts.statements[key]
            stmts.selected.discard(key)
            _selected_key.set("")
            edit_text.delete(1.0, tk.END)
            _m_apply_condition("")
            _sync_include_from_selection()
            _refresh_tree(search_var.get())

    ttk.Checkbutton(
        actions_top,
        text="Include this statement in generated reports",
        variable=m_include_var,
        command=_on_include_toggle,
    ).pack(anchor="w")

    act_btn_row = ttk.Frame(actions_top)
    act_btn_row.pack(fill="x", pady=(6, 0))
    ttk.Button(act_btn_row, text="Save changes", command=_save_edits).pack(side=tk.LEFT, padx=(0, 6))
    ttk.Button(act_btn_row, text="Delete statement", command=_delete_stmt).pack(side=tk.LEFT, padx=6)

    tk.Label(
        actions_top,
        text="Uncheck to disable without deleting. Use Placement below to move this statement. "
             "If several share one placement, they are ordered A to Z by name. Scroll the right panel to see all options.",
        font=("Arial", 8),
        fg="gray",
        wraplength=480,
        justify=tk.LEFT,
    ).pack(anchor="w", pady=(8, 0))

    _bind_manage_panel_mousewheel(right_inner)

    if manage_refresh_hooks is not None:
        manage_refresh_hooks.append(lambda: _refresh_tree(search_var.get()))

    _refresh_tree()


# ── Choose Statements ────────────────────────────────────────────

def show_choose_statements_dialog(
    root: tk.Tk,
    stmts: StatementManager,
    json_path: str,
    generate_preview_fn=None,
) -> None:
    dialog = tk.Toplevel(root)
    dialog.title("Choose Statements for Report")
    dialog.geometry("800x600")
    dialog.transient(root)
    dialog.grab_set()

    tk.Label(dialog, text="Select statements to include in the report (all checked by default):").pack(pady=5)
    frame = ttk.Frame(dialog)
    frame.pack(fill="both", expand=True, padx=10, pady=10)
    canvas = tk.Canvas(frame)
    sb = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
    sf = ttk.Frame(canvas)
    sf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=sf, anchor="nw")
    canvas.configure(yscrollcommand=sb.set)
    canvas.pack(side="left", fill="both", expand=True)
    sb.pack(side="right", fill="y")

    check_vars: Dict[str, tk.BooleanVar] = {}
    for key in sorted(stmts.statements.keys()):
        var = tk.BooleanVar(value=key in stmts.selected)
        check_vars[key] = var
        placement = stmts.get_placement_label(key)
        val = stmts.statements[key]
        cond = val.get("condition", "") if isinstance(val, dict) else ""
        ttk.Checkbutton(sf, text=f"{key} ({placement}) [Condition: {cond or 'None'}]", variable=var).pack(anchor="w", pady=2)

    def save_selection():
        stmts.selected = {k for k, v in check_vars.items() if v.get()}
        messagebox.showinfo("Success", "Statement selection saved.")
        dialog.destroy()

    def preview():
        if not generate_preview_fn:
            return
        text = generate_preview_fn()
        if not text:
            return
        pw = tk.Toplevel(dialog)
        pw.title("Report Preview")
        pw.geometry("600x400")
        pw.transient(dialog)
        pw.grab_set()
        tk.Label(pw, text="Report Preview (first 1000 characters):").pack(pady=5)
        ta = scrolledtext.ScrolledText(pw, width=80, height=20, wrap="word")
        ta.pack(pady=5)
        ta.insert(tk.END, text[:1000])
        ta.config(state="disabled")
        tk.Button(pw, text="Close", command=pw.destroy).pack(pady=5)

    tk.Button(dialog, text="Preview Report", command=preview).pack(pady=5)
    tk.Button(dialog, text="Save Selection", command=save_selection).pack(pady=5)
    tk.Button(dialog, text="Cancel", command=dialog.destroy).pack(pady=5)


# ── Report Template Editor (full overhaul) ────────────────────────

SECTION_LABELS = {
    "incident_summary": "Incident Summary",
    "suspect_information": "Suspect Information",
    "evidence_summary": "Evidence Summary",
    "ip_analysis": "IP Address Analysis",
}


def show_template_dialog(root: tk.Tk, template: ReportTemplate) -> None:
    dialog = tk.Toplevel(root)
    dialog.title("Report Template Settings")
    dialog.geometry("750x750")
    dialog.transient(root)
    dialog.grab_set()

    # ── Profile Management Bar ────────────────────────────────────
    profile_bar = ttk.LabelFrame(
        dialog,
        text="  Default & custom report templates  ",
        padding=5,
    )
    profile_bar.pack(fill="x", padx=10, pady=(10, 5))

    tk.Label(
        profile_bar,
        text="Default is the standard layout (its own file). Custom templates are separate files—"
             "editing or saving one never changes another until you choose that template and save.",
        font=("Arial", 9),
        wraplength=680,
        justify=tk.LEFT,
    ).pack(anchor="w", padx=2, pady=(0, 6))

    profile_row = ttk.Frame(profile_bar)
    profile_row.pack(fill="x")
    tk.Label(profile_row, text="Template in editor:").pack(side=tk.LEFT, padx=(0, 5))

    active_profile: str = (template.name or "Default").strip() or "Default"
    template.name = active_profile

    suppress_profile_pick = False

    _profiles = ReportTemplate.combobox_profile_names()
    profile_var = tk.StringVar(value=active_profile)
    profile_cb = ttk.Combobox(profile_row, textvariable=profile_var,
                              values=_profiles, width=22)
    profile_cb.pack(side=tk.LEFT, padx=3)

    tk.Label(
        profile_bar,
        text="Choosing a template in the list loads it from disk into the editor. "
             "Use the tabs below to change section order, names, and custom sections. "
             "Save and Close writes only the template shown above. "
             "Create Custom Template saves a copy of what you see now under a new name and switches you to it—"
             "the previous template file is left unchanged on disk.",
        font=("Arial", 8),
        fg="gray",
        wraplength=680,
        justify=tk.LEFT,
    ).pack(anchor="w", padx=2, pady=(6, 4))

    profile_btn_row = ttk.Frame(profile_bar)
    profile_btn_row.pack(fill="x", pady=(0, 2))

    def _refresh_profile_list():
        nonlocal _profiles
        _profiles = ReportTemplate.combobox_profile_names()
        profile_cb["values"] = _profiles

    def _set_profile_var_quiet(value: str) -> None:
        nonlocal suppress_profile_pick
        suppress_profile_pick = True
        try:
            profile_var.set(value)
        finally:
            suppress_profile_pick = False

    def _load_template_by_name(name: str, *, revert_combo_on_error: bool = True) -> bool:
        """Load *name* from disk into the shared *template*. Returns True on success."""
        nonlocal active_profile
        name = (name or "").strip() or "Default"
        try:
            loaded = ReportTemplate.load_profile(name)
            template.name = loaded.name
            active_profile = template.name
            template.section_order = loaded.section_order
            template.section_visible = loaded.section_visible
            template.section_names = loaded.section_names
            template.custom_sections = loaded.custom_sections
            template.sync_custom_sections()
            _set_profile_var_quiet(active_profile)
            _refresh_all()
            return True
        except FileNotFoundError:
            messagebox.showwarning("Not found", f"Template '{name}' does not exist on disk.")
            if revert_combo_on_error:
                _set_profile_var_quiet(active_profile)
            return False
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load template: {e}")
            if revert_combo_on_error:
                _set_profile_var_quiet(active_profile)
            return False

    def _on_profile_combobox_selected(event=None):
        if suppress_profile_pick:
            return
        name = (profile_var.get() or "").strip() or "Default"
        if name.casefold() == active_profile.casefold():
            return
        _load_template_by_name(name, revert_combo_on_error=True)

    def _on_profile_combobox_return(event=None):
        _on_profile_combobox_selected()

    def _create_new_profile():
        nonlocal active_profile
        current_sel = (profile_var.get() or "").strip() or "Default"
        name = simpledialog.askstring(
            "Create custom report template",
            "Name for the new custom template (a copy of the layout you see now).\n"
            "The built-in Default template keeps its own name; pick another name here.",
            parent=dialog,
            initialvalue="",
        )
        if name is None:
            return
        name = name.strip()
        if not name:
            messagebox.showwarning("Warning", "Template name cannot be empty.")
            return
        if name.casefold() == "default":
            messagebox.showinfo(
                "Reserved name",
                "'Default' is the built-in standard template.\n\n"
                "To reset it to the original layout, use Restore Default to Original below.\n"
                "For a new layout, choose a different name for your custom template.",
            )
            return
        if name.casefold() == current_sel.casefold():
            messagebox.showinfo(
                "Same name as current template",
                f"Choose a different name than '{current_sel}', or use Save and Close to update that file.",
            )
            return
        if name in ReportTemplate.list_profiles():
            if not messagebox.askyesno(
                "Overwrite?",
                f"A profile named '{name}' already exists.\n\nOverwrite it with the current layout?",
            ):
                return
        for sid, ent in name_entries.items():
            template.section_names[sid] = ent.get().strip() or DEFAULT_SECTION_NAMES.get(sid, sid.upper().replace("_", " "))
        template.name = current_sel
        payload = template.to_dict()
        try:
            ReportTemplate.save_payload_as_named_profile(payload, name)
        except Exception as e:
            messagebox.showerror("Error", f"Could not create profile: {e}")
            return
        _refresh_profile_list()
        try:
            loaded = ReportTemplate.load_profile(name)
            template.name = loaded.name
            active_profile = template.name
            template.section_order = loaded.section_order
            template.section_visible = loaded.section_visible
            template.section_names = loaded.section_names
            template.custom_sections = loaded.custom_sections
            template.sync_custom_sections()
            _set_profile_var_quiet(active_profile)
            _refresh_all()
        except Exception as e:
            messagebox.showerror("Error", f"Profile file was written but could not be opened: {e}")
            return
        messagebox.showinfo(
            "Custom template created",
            f"Custom template '{name}' was saved. You are now editing '{name}'.\n\n"
            f"'{current_sel}' on disk was not changed.",
        )

    def _delete_profile():
        nonlocal active_profile
        name = profile_var.get().strip()
        if not name or name == "Default":
            messagebox.showwarning("Warning", "The built-in Default template cannot be deleted.")
            return
        if not messagebox.askyesno("Confirm", f"Delete custom template '{name}'?"):
            return
        ReportTemplate.delete_profile(name)
        _refresh_profile_list()
        if name.casefold() == active_profile.casefold():
            try:
                loaded = ReportTemplate.load_profile("Default")
                template.name = loaded.name
                active_profile = template.name
                template.section_order = loaded.section_order
                template.section_visible = loaded.section_visible
                template.section_names = loaded.section_names
                template.custom_sections = loaded.custom_sections
                template.sync_custom_sections()
            except FileNotFoundError:
                template.section_order = list(DEFAULT_SECTIONS)
                template.section_visible = {s: True for s in DEFAULT_SECTIONS}
                template.section_names = dict(DEFAULT_SECTION_NAMES)
                template.custom_sections = []
                template.sync_custom_sections()
                template.name = "Default"
                active_profile = "Default"
                try:
                    template.save_profile()
                except Exception:
                    pass
            _set_profile_var_quiet(active_profile)
            _refresh_all()
        else:
            _set_profile_var_quiet(active_profile)

    ttk.Button(profile_row, text="Create Custom Template", command=_create_new_profile).pack(side=tk.LEFT, padx=(8, 0))
    ttk.Button(profile_btn_row, text="Delete Template", command=_delete_profile).pack(side=tk.LEFT, padx=(0, 4))

    notebook = ttk.Notebook(dialog)
    notebook.pack(fill="both", expand=True, padx=10, pady=(5, 10))

    # ── Tab 1: Section Order & Visibility ─────────────────────────
    tab_order = ttk.Frame(notebook)
    notebook.add(tab_order, text="Section Order & Visibility")

    tk.Label(tab_order, text="Configure which sections appear and in what order.",
             wraplength=600).pack(pady=5)

    order_frame = ttk.Frame(tab_order)
    order_frame.pack(fill="both", expand=True, padx=10, pady=5)

    listbox = tk.Listbox(order_frame, width=45, height=8, font=("Arial", 11))
    listbox.pack(side=tk.LEFT, fill="both", expand=True)

    def _refresh_listbox():
        listbox.delete(0, tk.END)
        for s in template.section_order:
            label = template.get_section_name(s)
            prefix = "[x]" if template.section_visible.get(s, True) else "[ ]"
            listbox.insert(tk.END, f"{prefix}  {label}")

    _refresh_listbox()

    btn_frame = ttk.Frame(order_frame)
    btn_frame.pack(side=tk.RIGHT, padx=10)

    def move_up():
        sel = listbox.curselection()
        if not sel or sel[0] == 0:
            return
        idx = sel[0]
        template.section_order[idx], template.section_order[idx - 1] = (
            template.section_order[idx - 1], template.section_order[idx]
        )
        _refresh_listbox()
        listbox.selection_set(idx - 1)

    def move_down():
        sel = listbox.curselection()
        if not sel or sel[0] >= len(template.section_order) - 1:
            return
        idx = sel[0]
        template.section_order[idx], template.section_order[idx + 1] = (
            template.section_order[idx + 1], template.section_order[idx]
        )
        _refresh_listbox()
        listbox.selection_set(idx + 1)

    def toggle_visible():
        sel = listbox.curselection()
        if not sel:
            return
        section = template.section_order[sel[0]]
        template.section_visible[section] = not template.section_visible[section]
        _refresh_listbox()
        listbox.selection_set(sel[0])

    tk.Button(btn_frame, text="Move Up", command=move_up, width=14).pack(pady=3)
    tk.Button(btn_frame, text="Move Down", command=move_down, width=14).pack(pady=3)
    tk.Button(btn_frame, text="Toggle Visible", command=toggle_visible, width=14).pack(pady=3)

    # ── Tab 2: Section Renaming ───────────────────────────────────
    tab_rename = ttk.Frame(notebook)
    notebook.add(tab_rename, text="Section Names")

    tk.Label(tab_rename, text="Customize the header text for each report section.",
             wraplength=600).pack(pady=10)

    names_container = ttk.Frame(tab_rename)
    names_container.pack(fill="both", expand=True)

    name_entries: Dict[str, tk.Entry] = {}

    def _rebuild_name_entries():
        for w in names_container.winfo_children():
            w.destroy()
        name_entries.clear()
        for section_id in template.section_order:
            row = ttk.Frame(names_container)
            row.pack(fill="x", padx=20, pady=4)
            default_label = SECTION_LABELS.get(section_id, section_id.replace("_", " ").title())
            tk.Label(row, text=f"{default_label}:", width=22, anchor="w").pack(side=tk.LEFT)
            ent = tk.Entry(row, width=40)
            ent.pack(side=tk.LEFT, padx=5)
            ent.insert(0, template.section_names.get(section_id, DEFAULT_SECTION_NAMES.get(section_id, "")))
            name_entries[section_id] = ent

            default_val = DEFAULT_SECTION_NAMES.get(section_id, "")
            def _reset_name(sid=section_id, entry=ent, dv=default_val):
                entry.delete(0, tk.END)
                entry.insert(0, dv)
            tk.Button(row, text="Reset", command=_reset_name).pack(side=tk.LEFT, padx=5)

    _rebuild_name_entries()

    def _apply_names():
        for sid, ent in name_entries.items():
            template.section_names[sid] = ent.get().strip() or DEFAULT_SECTION_NAMES.get(sid, sid.upper().replace("_", " "))
        _refresh_listbox()
        messagebox.showinfo("Applied", "Section names updated.")

    tk.Button(tab_rename, text="Apply Names", command=_apply_names).pack(pady=10)

    # ── Tab 3: Custom Sections ────────────────────────────────────
    tab_custom = ttk.Frame(notebook)
    notebook.add(tab_custom, text="Custom Sections")

    tk.Label(tab_custom, text="Add custom named sections with free-form body text.\n"
             "These appear after the standard sections.",
             wraplength=600).pack(pady=10)

    cs_list = tk.Listbox(tab_custom, width=60, height=6)
    cs_list.pack(fill="x", padx=20, pady=5)

    def _refresh_cs_list():
        cs_list.delete(0, tk.END)
        for cs in template.custom_sections:
            cs_list.insert(tk.END, cs.get("name", "(unnamed)"))

    _refresh_cs_list()

    cs_name_row = ttk.Frame(tab_custom)
    cs_name_row.pack(fill="x", padx=20, pady=3)
    tk.Label(cs_name_row, text="Section Name:").pack(side=tk.LEFT, padx=(0, 5))
    cs_name_entry = tk.Entry(cs_name_row, width=30)
    cs_name_entry.pack(side=tk.LEFT)

    tk.Label(tab_custom, text="Section Body (supports placeholders):").pack(anchor="w", padx=20)
    tk.Label(tab_custom,
             text="Available: [CURRENT_DATE] [INVESTIGATOR_NAME] [INVESTIGATOR_TITLE] "
                  "[CYBERTIP_NUMBER] [ESP_NAME] [REPORT_DATE_RECEIVED] [SUSPECT_NAME] "
                  "[SUSPECT_EMAIL] [SUSPECT_PHONE] [SUSPECT_SCREEN_NAME] "
                  "[INCIDENT_DATE] [INCIDENT_TYPE] [EVIDENCE_COUNT] "
                  "[TOTAL_FILES] [TOTAL_IPS] [AGENCY_NAME] [CASE_NUMBER]",
             fg="gray", wraplength=550, font=("Arial", 8)).pack(anchor="w", padx=20)
    cs_body = tk.Text(tab_custom, width=60, height=6, wrap="word")
    cs_body.pack(fill="x", padx=20, pady=5)

    def _add_custom_section():
        name = cs_name_entry.get().strip()
        body = cs_body.get(1.0, tk.END).strip()
        if not name:
            messagebox.showwarning("Warning", "Enter a section name.")
            return
        template.custom_sections.append({"name": name, "body": body})
        template.sync_custom_sections()
        cs_name_entry.delete(0, tk.END)
        cs_body.delete(1.0, tk.END)
        _refresh_cs_list()
        _refresh_listbox()
        _rebuild_name_entries()

    def _edit_custom_section():
        sel = cs_list.curselection()
        if not sel:
            messagebox.showwarning("Warning", "Select a custom section to edit.")
            return
        cs = template.custom_sections[sel[0]]
        cs_name_entry.delete(0, tk.END)
        cs_name_entry.insert(0, cs.get("name", ""))
        cs_body.delete(1.0, tk.END)
        cs_body.insert(tk.END, cs.get("body", ""))
        template.custom_sections.pop(sel[0])
        template.sync_custom_sections()
        _refresh_cs_list()
        _refresh_listbox()
        _rebuild_name_entries()

    def _remove_custom_section():
        sel = cs_list.curselection()
        if not sel:
            return
        if not messagebox.askyesno("Confirm", f"Remove custom section '{template.custom_sections[sel[0]].get('name', '')}'?"):
            return
        template.custom_sections.pop(sel[0])
        template.sync_custom_sections()
        _refresh_cs_list()
        _refresh_listbox()
        _rebuild_name_entries()

    cs_btns = ttk.Frame(tab_custom)
    cs_btns.pack(fill="x", padx=20, pady=5)
    ttk.Button(cs_btns, text="Add Section", command=_add_custom_section).pack(side=tk.LEFT, padx=3)
    ttk.Button(cs_btns, text="Edit Selected", command=_edit_custom_section).pack(side=tk.LEFT, padx=3)
    ttk.Button(cs_btns, text="Remove Selected", command=_remove_custom_section).pack(side=tk.LEFT, padx=3)

    # ── Tab 4: Preview ────────────────────────────────────────────
    tab_preview = ttk.Frame(notebook)
    notebook.add(tab_preview, text="Preview")

    tk.Label(tab_preview, text="Skeleton preview of the report layout based on current settings.",
             wraplength=600).pack(pady=5)

    preview_text = scrolledtext.ScrolledText(tab_preview, width=70, height=20, wrap="word",
                                              font=("Consolas", 10), state="disabled", bg="#fafafa")
    preview_text.pack(fill="both", expand=True, padx=10, pady=5)

    def _refresh_preview():
        lines = []
        lines.append("=" * 50)
        lines.append("REPORT SKELETON PREVIEW")
        lines.append("=" * 50)
        lines.append("")
        lines.append("[Introduction Paragraph]")
        lines.append("On [CURRENT_DATE], I, [INVESTIGATOR_TITLE] [INVESTIGATOR_NAME], ...")
        lines.append("")

        for section in template.section_order:
            visible = template.section_visible.get(section, True)
            name = template.get_section_name(section)
            marker = "" if visible else "  [HIDDEN]"
            lines.append(f"{'─' * 40}")
            lines.append(f"{name}:{marker}")
            if section.startswith("custom_"):
                cs = template.get_custom_section_by_id(section)
                if cs:
                    body_preview = cs.get("body", "")[:80]
                    lines.append(f"  {body_preview}{'...' if len(cs.get('body', '')) > 80 else ''}")
                else:
                    lines.append(f"  (custom section content)")
            else:
                lines.append(f"  (content for {SECTION_LABELS.get(section, section)})")
            lines.append("")

        lines.append("─" * 40)
        lines.append("[Custom Statements at End of Report]")
        lines.append("=" * 50)

        preview_text.configure(state="normal")
        preview_text.delete(1.0, tk.END)
        preview_text.insert(tk.END, "\n".join(lines))
        preview_text.configure(state="disabled")

    ttk.Button(tab_preview, text="Refresh Preview", command=_refresh_preview).pack(pady=5)
    _refresh_preview()

    def _on_tab_changed(event=None):
        if notebook.index(notebook.select()) == notebook.index(tab_preview):
            _refresh_preview()
    notebook.bind("<<NotebookTabChanged>>", _on_tab_changed)

    def _refresh_all():
        _refresh_listbox()
        _rebuild_name_entries()
        _refresh_cs_list()
        _refresh_preview()

    profile_cb.bind("<<ComboboxSelected>>", _on_profile_combobox_selected)
    profile_cb.bind("<Return>", _on_profile_combobox_return)

    def _edit_template():
        sel = (profile_var.get() or "").strip() or "Default"
        if sel.casefold() != active_profile.casefold():
            if not _load_template_by_name(sel, revert_combo_on_error=True):
                return
        notebook.select(0)

    ttk.Button(profile_btn_row, text="Edit Template", command=_edit_template).pack(side=tk.LEFT, padx=4)

    # ── Bottom buttons ────────────────────────────────────────────
    bottom = ttk.Frame(dialog)
    bottom.pack(fill="x", padx=10, pady=10)
    tk.Label(
        bottom,
        text="Restore Default to Original resets only the built-in Default template file to factory section order and names. "
             "If you are editing Default, the screen updates; custom templates you have open are not changed.",
        font=("Arial", 8),
        fg="gray",
        wraplength=700,
        justify=tk.LEFT,
    ).pack(anchor="w", fill="x", pady=(0, 6))

    def _reset_default_profile():
        nonlocal active_profile
        sel = (profile_var.get() or "").strip() or "Default"
        try:
            ReportTemplate.write_factory_default_file()
        except Exception as e:
            messagebox.showerror("Error", f"Could not restore the Default template on disk: {e}")
            return
        if sel.casefold() == "default":
            try:
                loaded = ReportTemplate.load_profile("Default")
                template.name = loaded.name
                active_profile = template.name
                template.section_order = loaded.section_order
                template.section_visible = loaded.section_visible
                template.section_names = loaded.section_names
                template.custom_sections = loaded.custom_sections
                template.sync_custom_sections()
            except Exception as e:
                messagebox.showerror("Error", f"Default.json was written but could not be reloaded: {e}")
                return
            _set_profile_var_quiet(active_profile)
            _refresh_all()
            messagebox.showinfo(
                "Default restored",
                "The Default template was reset to the original built-in layout. The editor now matches.",
            )
        else:
            messagebox.showinfo(
                "Default restored",
                "The Default template file was reset to the original built-in layout.\n\n"
                f"Your custom template '{sel}' in the editor is unchanged.",
            )

    def save_and_close():
        nonlocal active_profile
        sel = (profile_var.get() or "").strip() or "Default"
        active_profile = sel
        template.name = active_profile
        for sid, ent in name_entries.items():
            template.section_names[sid] = ent.get().strip() or DEFAULT_SECTION_NAMES.get(sid, sid.upper().replace("_", " "))
        _refresh_listbox()
        template.sync_custom_sections()
        template.save_profile()
        _refresh_profile_list()
        messagebox.showinfo(f"Saved: {active_profile}", f"Template profile '{active_profile}' saved.")
        dialog.destroy()

    def _export_template():
        path = filedialog.asksaveasfilename(
            title="Export Template", defaultextension=".json",
            filetypes=[("JSON files", "*.json")],
        )
        if path:
            try:
                template.export_to_file(path)
                messagebox.showinfo("Exported", f"Template exported to:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"Export failed: {e}")

    def _import_template():
        nonlocal active_profile
        path = filedialog.askopenfilename(
            title="Import Template", filetypes=[("JSON files", "*.json")],
        )
        if not path:
            return
        try:
            loaded = ReportTemplate.import_from_file(path)
            template.name = loaded.name
            active_profile = template.name
            template.section_order = loaded.section_order
            template.section_visible = loaded.section_visible
            template.section_names = loaded.section_names
            template.custom_sections = loaded.custom_sections
            template.sync_custom_sections()
            _set_profile_var_quiet(active_profile)
            _refresh_all()
            messagebox.showinfo("Imported", "Template imported successfully.")
        except Exception as e:
            messagebox.showerror("Error", f"Import failed: {e}")

    ttk.Button(bottom, text="Restore Default to Original", command=_reset_default_profile).pack(side=tk.LEFT, padx=5)
    ttk.Button(bottom, text="Export", command=_export_template).pack(side=tk.LEFT, padx=5)
    ttk.Button(bottom, text="Import", command=_import_template).pack(side=tk.LEFT, padx=5)
    ttk.Button(bottom, text="Save and Close", command=save_and_close).pack(side=tk.RIGHT, padx=5)
    ttk.Button(bottom, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)


# ── Report Comparison ─────────────────────────────────────────────

_COMPARE_CATEGORIES = (
    "ips", "emails", "phones", "screen_names",
    "user_ids", "urls", "hashes",
)

_COMPARE_DND = False
try:
    import tkinterdnd2
    _COMPARE_DND = True
except (ImportError, RuntimeError):
    pass


def _parse_dnd_file_list(widget: tk.Misc, data: str) -> List[str]:
    """Parse tkdnd file drop *data* into a list of paths (handles spaces, Windows braces)."""
    if not data or not data.strip():
        return []
    try:
        return [str(p) for p in widget.tk.splitlist(data)]
    except tk.TclError:
        paths: List[str] = []
        for token in data.replace("{", " ").replace("}", " ").split():
            token = token.strip()
            if token:
                paths.append(token)
        return paths


def _comparison_tag_docx_style(tags: FrozenSet[str]) -> Dict[str, object]:
    """Map Tk Text tags to python-docx run options (first matching tag wins)."""
    from docx.shared import RGBColor

    priority: List[Tuple[str, Dict[str, object]]] = [
        ("error", {"bold": True, "color": RGBColor(0xCC, 0x00, 0x00)}),
        ("title", {"bold": True, "size_pt": 16, "color": RGBColor(0x1A, 0x3A, 0x5C)}),
        ("section_header", {"bold": True, "size_pt": 12, "color": RGBColor(0x2A, 0x64, 0x96)}),
        ("report_header", {"bold": True, "size_pt": 11, "color": RGBColor(0x1A, 0x3A, 0x5C)}),
        ("category_header", {"bold": True, "size_pt": 11, "color": RGBColor(0x00, 0x55, 0x00)}),
        ("shared_value", {"bold": True, "color": RGBColor(0xCC, 0x44, 0x00)}),
        ("legend", {"italic": True, "size_pt": 9, "color": RGBColor(0x88, 0x88, 0x88)}),
        ("no_data", {"italic": True, "color": RGBColor(0x99, 0x99, 0x99)}),
        ("tip_refs", {"size_pt": 9, "color": RGBColor(0x66, 0x66, 0x66)}),
        ("separator", {"color": RGBColor(0x88, 0x88, 0x88)}),
        ("tip_list", {"color": RGBColor(0x33, 0x55, 0x77)}),
        ("detail_category", {"bold": True, "color": RGBColor(0x33, 0x33, 0x33)}),
        ("detail_value", {"color": RGBColor(0x44, 0x44, 0x44)}),
        ("common_value", {}),
    ]
    for tag, style in priority:
        if tag in tags:
            return style
    return {}


def _export_comparison_to_docx(text_widget: tk.Text, filepath: str) -> None:
    """Write comparison results to *filepath* preserving tag colors and emphasis."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    active: Set[str] = set()
    para = None

    def apply_run(run, tags: Set[str]) -> None:
        st = _comparison_tag_docx_style(frozenset(tags))
        run.bold = bool(st.get("bold", False))
        run.italic = bool(st.get("italic", False))
        sz = st.get("size_pt")
        if sz:
            run.font.size = Pt(int(sz))
        col = st.get("color")
        if col is not None:
            run.font.color.rgb = col  # type: ignore[assignment]

    for key, value, *_ in text_widget.dump("1.0", tk.END, text=True, tag=True):
        if key == "tagon":
            active.add(value)
        elif key == "tagoff":
            active.discard(value)
        elif key == "text":
            if value == "":
                continue
            lines = value.split("\n")
            for li, line in enumerate(lines):
                if li > 0 or para is None:
                    para = doc.add_paragraph()
                if line:
                    apply_run(para.add_run(line), set(active))

    doc.save(filepath)


def _collect_json_paths(path: str) -> List[str]:
    """Return .json file paths from *path*.  If *path* is a directory,
    recursively walk it; otherwise return it as-is if it ends with .json."""
    p = Path(path)
    if p.is_dir():
        return sorted(str(f) for f in p.rglob("*.json") if f.is_file())
    if p.suffix.lower() == ".json" and p.is_file():
        return [str(p)]
    return []


def show_comparison_dialog(root: tk.Tk) -> None:
    dialog = tk.Toplevel(root)
    dialog.title("Compare CyberTip Reports")
    dialog.geometry("1000x800")
    dialog.transient(root)
    # Do not use grab_set(): it prevents native file drag-and-drop onto this window on Windows.

    compare_dnd_ok = False
    if _COMPARE_DND:
        try:
            tkinterdnd2.TkinterDnD._require(root)
            compare_dnd_ok = True
        except RuntimeError as exc:
            log.warning("tkdnd could not be loaded for comparison dialog: %s", exc)

    tk.Label(
        dialog,
        text="Load multiple CyberTip JSON files to find common identifiers.",
        font=("Arial", 11),
    ).pack(pady=10)

    # --- drop zone ---
    dnd_hint = (
        "Drag & drop .json files or folders here"
        if compare_dnd_ok else
        "Use the buttons to add .json files or folders"
    )
    drop_frame = tk.Frame(
        dialog, bd=2, relief="groove", bg="#f0f4f8",
        highlightbackground="#a0b4c8", highlightthickness=1,
    )
    drop_frame.pack(fill="x", padx=10, pady=(0, 5))
    drop_label = tk.Label(
        drop_frame, text=dnd_hint,
        font=("Arial", 10, "italic"), fg="#6688aa", bg="#f0f4f8", pady=8,
    )
    drop_label.pack(fill="both", expand=True, padx=4, pady=4)

    # --- file list with scrollbar ---
    files_frame = ttk.Frame(dialog)
    files_frame.pack(fill="x", padx=10, pady=5)

    list_frame = ttk.Frame(files_frame)
    list_frame.pack(side=tk.LEFT, fill="both", expand=True)

    files_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
    files_list = tk.Listbox(
        list_frame, width=80, height=8,
        yscrollcommand=files_scrollbar.set,
    )
    files_scrollbar.config(command=files_list.yview)
    files_list.pack(side=tk.LEFT, fill="both", expand=True)
    files_scrollbar.pack(side=tk.LEFT, fill=tk.Y)

    loaded_files: List[str] = []

    count_label = tk.Label(dialog, text="Files loaded: 0", font=("Arial", 9))
    count_label.pack(pady=(0, 5))

    def _update_count():
        count_label.config(text=f"Files loaded: {len(loaded_files)}")

    def _ingest_paths(paths) -> int:
        """Add JSON files (or folders of JSON files) to the list.
        Returns the number of new files actually added."""
        added = 0
        for raw_path in paths:
            for json_path in _collect_json_paths(raw_path):
                if json_path not in loaded_files:
                    loaded_files.append(json_path)
                    files_list.insert(tk.END, os.path.basename(json_path))
                    added += 1
        _update_count()
        return added

    # --- buttons ---
    btn_frame = ttk.Frame(files_frame)
    btn_frame.pack(side=tk.RIGHT, padx=5)

    def add_files():
        paths = filedialog.askopenfilenames(filetypes=[("JSON files", "*.json")])
        _ingest_paths(paths)

    def add_folder():
        folder = filedialog.askdirectory(title="Select folder to scan for .json files")
        if folder:
            added = _ingest_paths([folder])
            if added == 0:
                messagebox.showinfo("Add Folder", "No .json files found in the selected folder.")

    def remove_file():
        sel = files_list.curselection()
        if sel:
            loaded_files.pop(sel[0])
            files_list.delete(sel[0])
            _update_count()

    def clear_all():
        loaded_files.clear()
        files_list.delete(0, tk.END)
        _update_count()

    tk.Button(btn_frame, text="Add Files", command=add_files, width=12).pack(pady=3)
    tk.Button(btn_frame, text="Add Folder", command=add_folder, width=12).pack(pady=3)
    tk.Button(btn_frame, text="Remove", command=remove_file, width=12).pack(pady=3)
    tk.Button(btn_frame, text="Clear All", command=clear_all, width=12).pack(pady=3)

    # --- drag-and-drop setup ---
    def _on_drag_enter(event):
        drop_frame.config(bg="#d6e8f7", highlightbackground="#3388cc", highlightthickness=2)
        drop_label.config(bg="#d6e8f7", fg="#2266aa", text="Drop files or folders here...")

    def _on_drag_leave(event):
        drop_frame.config(bg="#f0f4f8", highlightbackground="#a0b4c8", highlightthickness=1)
        drop_label.config(bg="#f0f4f8", fg="#6688aa", text=dnd_hint)

    def _on_drop(event):
        _on_drag_leave(event)
        paths = _parse_dnd_file_list(dialog, getattr(event, "data", "") or "")
        if not paths:
            return
        added = _ingest_paths(paths)
        if added:
            drop_label.config(fg="#228833", text=f"{added} file(s) added")
            dialog.after(2500, lambda: drop_label.config(fg="#6688aa", text=dnd_hint))

    if compare_dnd_ok:
        try:
            # Register the label (fills the drop zone); child widgets do not inherit drop targets.
            drop_label.drop_target_register(tkinterdnd2.DND_FILES)
            drop_label.dnd_bind("<<Drop>>", _on_drop)
            drop_label.dnd_bind("<<DragEnter>>", _on_drag_enter)
            drop_label.dnd_bind("<<DragLeave>>", _on_drag_leave)
        except Exception as exc:
            log.warning("Could not enable drag-and-drop on comparison dialog: %s", exc)

    # --- results area ---
    results_text = scrolledtext.ScrolledText(
        dialog, width=110, height=30, wrap="word",
        font=("Consolas", 10), bg="#fdfdfd",
    )
    results_text.pack(fill="both", expand=True, padx=10, pady=10)

    # Text tags for styled output
    results_text.tag_configure("title", font=("Arial", 13, "bold"), foreground="#1a3a5c")
    results_text.tag_configure("separator", foreground="#888888")
    results_text.tag_configure("tip_list", font=("Consolas", 10), foreground="#335577", lmargin1=20)
    results_text.tag_configure("section_header", font=("Arial", 12, "bold"), foreground="#2a6496",
                               spacing3=4)
    results_text.tag_configure("category_header", font=("Consolas", 11, "bold"), foreground="#005500",
                               lmargin1=10)
    results_text.tag_configure("common_value", font=("Consolas", 10), lmargin1=30)
    results_text.tag_configure("tip_refs", font=("Consolas", 9), foreground="#666666")
    results_text.tag_configure("report_header", font=("Arial", 11, "bold"), foreground="#1a3a5c",
                               spacing1=6)
    results_text.tag_configure("detail_category", font=("Consolas", 10, "bold"), foreground="#333333",
                               lmargin1=20)
    results_text.tag_configure("detail_value", font=("Consolas", 10), foreground="#444444", lmargin1=40)
    results_text.tag_configure("shared_value", font=("Consolas", 10, "bold"), foreground="#cc4400",
                               lmargin1=40)
    results_text.tag_configure("no_data", font=("Consolas", 10, "italic"), foreground="#999999",
                               lmargin1=20)
    results_text.tag_configure("legend", font=("Consolas", 9, "italic"), foreground="#888888")
    results_text.tag_configure("error", font=("Consolas", 10, "bold"), foreground="#cc0000")

    def _ins(text: str, *tags: str):
        results_text.insert(tk.END, text, tags)

    # --- comparison logic (pairwise overlap: values in 2+ reports) ---
    def run_comparison():
        if len(loaded_files) < 2:
            messagebox.showwarning("Warning", "Load at least 2 files to compare.")
            return

        results_text.delete(1.0, tk.END)
        all_data: Dict[str, dict] = {}

        for path in loaded_files:
            data = load_json(path)
            if data is None:
                _ins(f"ERROR: Could not load {os.path.basename(path)}\n", "error")
                continue
            comparison = extract_comparison_data(data)
            report_id = str(data.get("reportId", os.path.basename(path)))
            all_data[report_id] = comparison

        if len(all_data) < 2:
            _ins("Not enough valid files to compare.\n", "error")
            return

        _ins(f"Comparing {len(all_data)} CyberTip Reports\n", "title")
        for rid in sorted(all_data.keys()):
            _ins(f"  - CyberTip #{rid}\n", "tip_list")
        _ins("=" * 60 + "\n\n", "separator")

        found_any = False

        for category in _COMPARE_CATEGORIES:
            label = category.replace("_", " ").title()

            value_to_tips: Dict[str, Set[str]] = defaultdict(set)
            for rid, comp in all_data.items():
                for val in comp.get(category, set()):
                    value_to_tips[val].add(rid)

            common = {v: tips for v, tips in value_to_tips.items() if len(tips) >= 2}
            if not common:
                continue

            found_any = True
            sorted_vals = sorted(common.items(), key=lambda x: len(x[1]), reverse=True)
            _ins(f"COMMON {label.upper()} ", "category_header")
            _ins(f"({len(common)} values)\n", "tip_refs")
            for val, tips in sorted_vals:
                tip_list = ", ".join(sorted(tips))
                _ins(f"  {val}  ", "common_value")
                _ins(f"-- Found in {len(tips)} tips: {tip_list}\n", "tip_refs")
            _ins("\n")

        if not found_any:
            _ins("No common identifiers found across the loaded reports.\n\n", "no_data")

        # Build a set of values that appeared in 2+ reports so we can
        # highlight them in the per-report detail.
        shared_values: Dict[str, set] = defaultdict(set)
        for category in _COMPARE_CATEGORIES:
            for rid, comp in all_data.items():
                for val in comp.get(category, set()):
                    shared_values[(category, val)].add(rid)

        _ins("\nPER-REPORT DETAIL\n", "section_header")
        _ins("=" * 60 + "\n", "separator")

        for rid, comp in all_data.items():
            _ins(f"\nCyberTip #{rid}\n", "report_header")
            has_data = False
            for cat in _COMPARE_CATEGORIES:
                vals = comp.get(cat, set())
                if not vals:
                    continue
                has_data = True
                label = cat.replace("_", " ").title()
                _ins(f"  {label} ({len(vals)}):\n", "detail_category")
                for val in sorted(vals):
                    is_shared = len(shared_values.get((cat, val), set())) >= 2
                    if is_shared:
                        _ins(f"    {val} *\n", "shared_value")
                    else:
                        _ins(f"    {val}\n", "detail_value")
            if not has_data:
                _ins("  (no identifiers extracted)\n", "no_data")

        _ins("\n" + "-" * 60 + "\n", "separator")
        _ins("* = value also found in another CyberTip report\n", "legend")

    # --- export results ---
    def export_results():
        content = results_text.get(1.0, tk.END).strip()
        if not content:
            messagebox.showinfo("Export", "No results to export. Run a comparison first.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[
                ("Word document", "*.docx"),
                ("Plain text", "*.txt"),
                ("All files", "*.*"),
            ],
            title="Export Comparison Results",
        )
        if not path:
            return
        try:
            lower = path.lower()
            if lower.endswith(".docx"):
                _export_comparison_to_docx(results_text, path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
            messagebox.showinfo("Export", f"Results exported to:\n{path}")
        except Exception as exc:
            messagebox.showerror("Export Error", str(exc))

    # --- bottom buttons ---
    bottom_frame = ttk.Frame(dialog)
    bottom_frame.pack(pady=5)
    tk.Button(bottom_frame, text="Compare", command=run_comparison, width=14).pack(side=tk.LEFT, padx=5)
    tk.Button(bottom_frame, text="Export Results", command=export_results, width=14).pack(side=tk.LEFT, padx=5)
    tk.Button(bottom_frame, text="Close", command=dialog.destroy, width=14).pack(side=tk.LEFT, padx=5)
